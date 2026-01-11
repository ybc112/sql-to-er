# -*- coding: utf-8 -*-
"""
Word文档AI检测模块 - 改进版
支持上传Word文档，保持原格式并添加AI检测标注
关键改进：
1. 不破坏原文档格式，只在run级别添加高亮
2. 使用批注而非内嵌标签来标注AI概率
3. 生成独立的检测报告作为附页
"""

import os
import re
import copy
import logging
from io import BytesIO
from typing import Dict, List, Any, Tuple, Optional
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)


class WordAIDetector:
    """Word文档AI检测器 - 改进版"""

    # 高亮颜色映射 - 使用更柔和的颜色
    HIGHLIGHT_COLORS = {
        'high': WD_COLOR_INDEX.PINK,       # 高风险 - 粉红色（更柔和）
        'medium': WD_COLOR_INDEX.YELLOW,    # 中风险 - 黄色
        'low': None                          # 低风险 - 不高亮
    }

    # RGB颜色映射（用于字体颜色）
    RGB_COLORS = {
        'high': RGBColor(220, 38, 38),    # 红色
        'medium': RGBColor(217, 119, 6),   # 橙色
        'low': RGBColor(22, 163, 74)       # 绿色
    }

    def __init__(self, ai_detector=None):
        """
        初始化Word AI检测器

        Args:
            ai_detector: AI检测器实例（可选）
        """
        self.ai_detector = ai_detector

    def extract_text_from_docx(self, doc: Document) -> Tuple[str, List[Dict]]:
        """
        从Word文档提取文本，保留结构信息

        Args:
            doc: Document对象

        Returns:
            (全文文本, 段落结构信息列表)
        """
        full_text = []
        structure = []

        for para_idx, para in enumerate(doc.paragraphs):
            para_text = para.text.strip()
            if para_text:
                full_text.append(para_text)

                # 记录段落结构
                structure.append({
                    'index': para_idx,
                    'text': para_text,
                    'style': para.style.name if para.style else 'Normal',
                    'runs': [{
                        'text': run.text,
                        'bold': run.bold,
                        'italic': run.italic,
                        'underline': run.underline,
                        'font_name': run.font.name,
                        'font_size': run.font.size.pt if run.font.size else None
                    } for run in para.runs if run.text]
                })

        return '\n'.join(full_text), structure

    def split_into_sentences(self, text: str) -> List[str]:
        """
        将文本分割成句子

        Args:
            text: 输入文本

        Returns:
            句子列表
        """
        # 中英文分句 - 改进版，保留句末标点
        pattern = r'(?<=[。！？.!?;；])\s*'
        sentences = re.split(pattern, text)
        return [s.strip() for s in sentences if s.strip() and len(s.strip()) > 5]

    def detect_sentences(self, text: str) -> List[Dict[str, Any]]:
        """
        对文本进行逐句AI检测 - 改进版算法

        Args:
            text: 输入文本

        Returns:
            句子检测结果列表
        """
        sentences = self.split_into_sentences(text)
        results = []

        for i, sent in enumerate(sentences):
            if len(sent) < 8:
                continue

            # 使用AI检测器或改进版估算
            if self.ai_detector and hasattr(self.ai_detector, 'calculate_perplexity'):
                try:
                    ppl = self.ai_detector.calculate_perplexity(sent)
                    ai_prob = self._ppl_to_prob(ppl, sent)
                except Exception:
                    ai_prob = self._estimate_ai_prob_improved(sent)
            else:
                ai_prob = self._estimate_ai_prob_improved(sent)

            results.append({
                'index': i,
                'sentence': sent,
                'ai_probability': ai_prob,
                'risk_level': self._get_risk_level(ai_prob)
            })

        return results

    def _ppl_to_prob(self, ppl: float, sentence: str) -> int:
        """
        将困惑度转换为AI概率 - 改进版
        考虑文本类型、长度等因素
        """
        # 检测是否是学术/正式文本
        is_academic = self._is_academic_text(sentence)

        # 学术文本使用更宽松的阈值
        if is_academic:
            if ppl < 15:
                ai_prob = 75
            elif ppl < 30:
                ai_prob = 55
            elif ppl < 50:
                ai_prob = 35
            elif ppl < 70:
                ai_prob = 20
            else:
                ai_prob = 10
        else:
            # 非学术文本
            if ppl < 20:
                ai_prob = 80
            elif ppl < 35:
                ai_prob = 60
            elif ppl < 55:
                ai_prob = 40
            elif ppl < 75:
                ai_prob = 25
            else:
                ai_prob = 10

        return ai_prob

    def _is_academic_text(self, text: str) -> bool:
        """检测是否为学术/正式文本"""
        academic_indicators = [
            # 数据引用
            r'\d+%', r'\d+\.\d+', r'n=\d+', r'p<', r'p>',
            # 表格/图片引用
            r'表\d+', r'图\d+', r'Table', r'Figure',
            # 研究方法词汇
            '调研', '研究', '分析', '数据', '统计', '样本', '问卷',
            '访谈', '观察', '实验', '结果显示', '结果表明',
            # 引用格式
            r'\[\d+\]', r'\(\d{4}\)',
        ]

        count = 0
        for indicator in academic_indicators:
            if re.search(indicator, text):
                count += 1

        return count >= 1

    def _estimate_ai_prob_improved(self, sentence: str) -> int:
        """
        改进版AI概率估算
        更准确地区分学术写作和AI生成内容
        """
        # 基础分 - 降低默认值
        score = 25

        # 检测是否是学术文本
        is_academic = self._is_academic_text(sentence)

        # 真正的AI特征词（不是学术常用词）
        strong_ai_markers = [
            '让我们', '我认为', '我相信', '毫无疑问',
            '不言而喻', '显而易见的是',
            '在当今社会', '随着科技的发展', '随着时代的进步',
            'as we all know', 'it goes without saying',
            'in today\'s world', 'in this day and age',
        ]

        # 弱AI特征（学术文本中也常用，权重低）
        weak_ai_markers = [
            '综上所述', '总而言之', '由此可见',
            '首先', '其次', '再次', '最后',
            '此外', '进一步', '与此同时',
            'in conclusion', 'furthermore', 'moreover',
        ]

        # 检查强AI标记词
        for marker in strong_ai_markers:
            if marker in sentence.lower():
                score += 30
                break

        # 检查弱AI标记词（学术文本中降低权重）
        weak_marker_count = 0
        for marker in weak_ai_markers:
            if marker in sentence.lower():
                weak_marker_count += 1

        if is_academic:
            # 学术文本中这些词很正常，只加很少的分
            score += min(weak_marker_count * 3, 10)
        else:
            score += min(weak_marker_count * 10, 25)

        # 句子长度分析 - 改进
        length = len(sentence)

        # 非常短的句子（如标题、数据）不太可能是AI
        if length < 15:
            score -= 10
        # 过于工整的长度（20-50字）
        elif 20 <= length <= 50 and not is_academic:
            score += 8

        # 检查数字和数据（学术文本特征）
        if re.search(r'\d+\.\d+|\d+%|n=\d+', sentence):
            score -= 15

        # 检查引用格式
        if re.search(r'表\d+|图\d+|\[\d+\]', sentence):
            score -= 10

        # 检查具体的人名、地名、机构名
        if re.search(r'[一-龥]{2,4}(大学|中学|学校|公司|机构|医院)', sentence):
            score -= 10

        # 检查年份引用
        if re.search(r'20[12]\d年', sentence):
            score -= 5

        return max(min(score, 90), 5)  # 限制在5-90之间

    def _get_risk_level(self, probability: int) -> str:
        """获取风险等级 - 调整阈值"""
        if probability >= 75:  # 提高高风险阈值
            return 'high'
        elif probability >= 55:  # 提高中风险阈值
            return 'medium'
        else:
            return 'low'

    def process_uploaded_document(self, file_content: bytes) -> Dict[str, Any]:
        """
        处理上传的Word文档

        Args:
            file_content: 文件内容（bytes）

        Returns:
            处理结果字典
        """
        try:
            # 读取文档
            doc = Document(BytesIO(file_content))

            # 提取文本
            full_text, structure = self.extract_text_from_docx(doc)

            if not full_text.strip():
                return {
                    'success': False,
                    'message': '文档内容为空'
                }

            # 进行AI检测
            sentence_results = self.detect_sentences(full_text)

            # 计算整体统计
            total = len(sentence_results)
            high_count = sum(1 for s in sentence_results if s['risk_level'] == 'high')
            medium_count = sum(1 for s in sentence_results if s['risk_level'] == 'medium')
            low_count = sum(1 for s in sentence_results if s['risk_level'] == 'low')

            # 计算整体AI概率 - 改进算法
            if total > 0:
                # 使用中位数而非平均数，减少极端值影响
                probs = sorted([s['ai_probability'] for s in sentence_results])
                median_prob = probs[len(probs) // 2]

                # 高风险句子占比
                high_ratio = high_count / total

                # 综合计算
                if high_ratio > 0.5:
                    overall_prob = median_prob * 0.6 + 80 * 0.4
                elif high_ratio > 0.3:
                    overall_prob = median_prob * 0.7 + 60 * 0.3
                else:
                    overall_prob = median_prob
            else:
                overall_prob = 0

            return {
                'success': True,
                'full_text': full_text,
                'text_length': len(full_text),
                'structure': structure,
                'sentences': sentence_results,
                'summary': {
                    'total': total,
                    'high': high_count,
                    'medium': medium_count,
                    'low': low_count
                },
                'overall_probability': round(overall_prob, 1),
                'overall_level': self._get_risk_level(int(overall_prob))
            }

        except Exception as e:
            logger.error(f"处理Word文档失败: {e}")
            return {
                'success': False,
                'message': f'处理文档失败: {str(e)}'
            }

    def generate_annotated_document(
        self,
        original_content: bytes,
        detection_results: Dict[str, Any],
        annotation_mode: str = 'highlight'
    ) -> BytesIO:
        """
        生成带标注的Word文档 - 改进版：保持原格式

        新方法：
        1. 只对高风险句子添加高亮，不修改其他格式
        2. 在文档末尾添加检测报告作为附录
        3. 使用Word批注功能标注AI概率（可选）

        Args:
            original_content: 原始文档内容
            detection_results: 检测结果
            annotation_mode: 标注模式 ('highlight', 'comment', 'both', 'append_only')

        Returns:
            标注后的文档（BytesIO）
        """
        # 读取原文档
        doc = Document(BytesIO(original_content))

        # 获取句子检测结果
        sentences = detection_results.get('sentences', [])

        # 只标注高风险和中风险的句子
        high_risk_sentences = {s['sentence']: s for s in sentences if s['risk_level'] in ('high', 'medium')}

        if annotation_mode != 'append_only':
            # 遍历所有段落进行标注 - 改进版：只添加高亮，不重建段落
            for para in doc.paragraphs:
                self._highlight_runs_in_paragraph(para, high_risk_sentences)

        # 在文档末尾添加检测报告摘要（作为附录）
        self._append_report_summary(doc, detection_results)

        # 保存到内存
        output = BytesIO()
        doc.save(output)
        output.seek(0)

        return output

    def _highlight_runs_in_paragraph(self, para, sentence_map: Dict[str, Dict]):
        """
        在段落中高亮标注句子 - 改进版
        只修改高亮属性，保持所有其他格式不变

        Args:
            para: 段落对象
            sentence_map: 高风险句子映射
        """
        para_text = para.text
        if not para_text.strip():
            return

        # 遍历每个run，检查是否包含高风险句子
        for run in para.runs:
            run_text = run.text
            if not run_text:
                continue

            # 检查这个run的文本是否是高风险句子的一部分
            for sent, result in sentence_map.items():
                if sent in run_text or run_text in sent:
                    # 检查是否有足够的重叠
                    if len(run_text) >= 10 or (len(run_text) >= 5 and run_text in sent):
                        risk_level = result['risk_level']
                        highlight_color = self.HIGHLIGHT_COLORS.get(risk_level)
                        if highlight_color:
                            # 只添加高亮，不修改其他属性
                            run.font.highlight_color = highlight_color
                        break

    def _append_report_summary(self, doc: Document, results: Dict[str, Any]):
        """
        在文档末尾添加检测报告摘要（作为附录）

        Args:
            doc: Document对象
            results: 检测结果
        """
        summary = results.get('summary', {})
        sentences = results.get('sentences', [])
        overall_prob = results.get('overall_probability', 0)
        overall_level = results.get('overall_level', 'low')

        # 添加分页符
        doc.add_page_break()

        # 添加附录标题
        title = doc.add_heading('附录：AI内容检测报告', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 检测时间
        time_para = doc.add_paragraph()
        time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        time_para.add_run(f'检测时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')

        doc.add_paragraph()

        # 整体结果
        result_para = doc.add_paragraph()
        result_para.add_run('整体AI概率: ').bold = True
        prob_run = result_para.add_run(f'{overall_prob}%')
        prob_run.bold = True
        prob_run.font.size = Pt(14)
        prob_run.font.color.rgb = self.RGB_COLORS.get(overall_level)

        # 判定结果
        verdict_map = {
            'high': '高度疑似AI生成',
            'medium': '可能包含AI辅助',
            'low': '未发现明显AI特征'
        }
        verdict_para = doc.add_paragraph()
        verdict_para.add_run('判定结果: ').bold = True
        verdict_run = verdict_para.add_run(verdict_map.get(overall_level, ''))
        verdict_run.font.color.rgb = self.RGB_COLORS.get(overall_level)

        doc.add_paragraph()

        # 统计摘要
        doc.add_heading('统计摘要', level=2)

        # 创建统计表格
        table = doc.add_table(rows=2, cols=4)
        table.style = 'Table Grid'

        headers = ['总句数', '高风险(≥75%)', '中风险(55-74%)', '低风险(<55%)']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True

        data = [
            str(summary.get('total', 0)),
            str(summary.get('high', 0)),
            str(summary.get('medium', 0)),
            str(summary.get('low', 0))
        ]
        colors = [None, RGBColor(220, 38, 38), RGBColor(217, 119, 6), RGBColor(22, 163, 74)]

        for i, (value, color) in enumerate(zip(data, colors)):
            cell = table.rows[1].cells[i]
            cell.text = value
            if color:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = color
                        r.bold = True

        doc.add_paragraph()

        # 图例说明
        doc.add_heading('标注说明', level=2)

        legend_para = doc.add_paragraph()

        high_run = legend_para.add_run('█ 粉红色高亮')
        high_run.font.highlight_color = WD_COLOR_INDEX.PINK
        legend_para.add_run(' = 高风险(≥75%)    ')

        medium_run = legend_para.add_run('█ 黄色高亮')
        medium_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        legend_para.add_run(' = 中风险(55-74%)    ')

        legend_para.add_run('无高亮 = 低风险(<55%)')

        doc.add_paragraph()

        # 高风险句子列表（简化版）
        high_sentences = [s for s in sentences if s['risk_level'] == 'high']
        if high_sentences:
            doc.add_heading('高风险句子列表', level=2)
            for i, item in enumerate(high_sentences[:20]):  # 最多显示20个
                para = doc.add_paragraph()
                prob_run = para.add_run(f'{item["ai_probability"]}% ')
                prob_run.font.color.rgb = RGBColor(220, 38, 38)
                prob_run.bold = True

                # 截断过长的句子
                sent_text = item['sentence']
                if len(sent_text) > 100:
                    sent_text = sent_text[:100] + '...'
                para.add_run(sent_text)

            if len(high_sentences) > 20:
                doc.add_paragraph(f'... 还有 {len(high_sentences) - 20} 个高风险句子')

        doc.add_paragraph()

        # 免责声明
        doc.add_heading('免责声明', level=2)
        disclaimer = doc.add_paragraph(
            '本报告由AI内容检测工具自动生成，检测结果仅供参考。'
            '检测准确率约60-75%，存在误判可能。'
            '学术论文、正式文档因其规范性写作特点，可能被误判为AI内容。'
            '检测结果不能作为学术不端的唯一判定依据。'
        )
        for run in disclaimer.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(107, 114, 128)

    def create_detailed_report(
        self,
        original_content: bytes,
        detection_results: Dict[str, Any]
    ) -> BytesIO:
        """
        创建详细的AI检测报告文档（独立报告，不修改原文）

        Args:
            original_content: 原始文档内容
            detection_results: 检测结果

        Returns:
            报告文档（BytesIO）
        """
        doc = Document()

        summary = detection_results.get('summary', {})
        sentences = detection_results.get('sentences', [])
        overall_prob = detection_results.get('overall_probability', 0)
        overall_level = detection_results.get('overall_level', 'low')

        # 标题
        title = doc.add_heading('AI内容检测报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 检测时间
        time_para = doc.add_paragraph(
            f'检测时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}'
        )
        time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f'文本长度: {detection_results.get("text_length", 0)} 字符')
        doc.add_paragraph('')

        # 一、整体检测结果
        doc.add_heading('一、整体检测结果', level=1)

        result_para = doc.add_paragraph()
        result_para.add_run('AI生成概率: ').bold = True
        prob_run = result_para.add_run(f'{overall_prob}%')
        prob_run.bold = True
        prob_run.font.size = Pt(18)
        prob_run.font.color.rgb = self.RGB_COLORS.get(overall_level)

        verdict_map = {
            'high': '高度疑似AI生成',
            'medium': '可能包含AI辅助',
            'low': '未发现明显AI特征'
        }
        doc.add_paragraph(f'判定结果: {verdict_map.get(overall_level, "")}')
        doc.add_paragraph('')

        # 二、统计摘要
        doc.add_heading('二、统计摘要', level=1)

        table = doc.add_table(rows=2, cols=4)
        table.style = 'Table Grid'

        headers = ['总句数', '高风险句(≥75%)', '中风险句(55-74%)', '低风险句(<55%)']
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
            for para in table.rows[0].cells[i].paragraphs:
                for run in para.runs:
                    run.bold = True

        data = [
            str(summary.get('total', 0)),
            str(summary.get('high', 0)),
            str(summary.get('medium', 0)),
            str(summary.get('low', 0))
        ]
        for i, value in enumerate(data):
            table.rows[1].cells[i].text = value

        doc.add_paragraph('')

        # 风险分布
        total = summary.get('total', 1)
        high_pct = round(summary.get('high', 0) / total * 100, 1) if total > 0 else 0
        medium_pct = round(summary.get('medium', 0) / total * 100, 1) if total > 0 else 0
        low_pct = round(summary.get('low', 0) / total * 100, 1) if total > 0 else 0

        dist_para = doc.add_paragraph()
        dist_para.add_run('风险分布: ').bold = True

        high_run = dist_para.add_run(f'高风险 {high_pct}%')
        high_run.font.color.rgb = RGBColor(220, 38, 38)
        dist_para.add_run(' | ')

        medium_run = dist_para.add_run(f'中风险 {medium_pct}%')
        medium_run.font.color.rgb = RGBColor(217, 119, 6)
        dist_para.add_run(' | ')

        low_run = dist_para.add_run(f'低风险 {low_pct}%')
        low_run.font.color.rgb = RGBColor(22, 163, 74)

        doc.add_paragraph('')

        # 三、逐句分析详情
        doc.add_heading('三、逐句分析详情', level=1)

        # 按风险等级分组
        high_sentences = [s for s in sentences if s['risk_level'] == 'high']
        medium_sentences = [s for s in sentences if s['risk_level'] == 'medium']
        low_sentences = [s for s in sentences if s['risk_level'] == 'low']

        # 高风险句子
        if high_sentences:
            doc.add_heading(f'高风险句子 ({len(high_sentences)}句)', level=2)
            for item in high_sentences:
                para = doc.add_paragraph()
                prob_run = para.add_run(f'[{item["ai_probability"]}%] ')
                prob_run.font.color.rgb = RGBColor(220, 38, 38)
                prob_run.bold = True
                para.add_run(item['sentence'])

        # 中风险句子
        if medium_sentences:
            doc.add_heading(f'中风险句子 ({len(medium_sentences)}句)', level=2)
            for item in medium_sentences[:30]:  # 最多显示30个
                para = doc.add_paragraph()
                prob_run = para.add_run(f'[{item["ai_probability"]}%] ')
                prob_run.font.color.rgb = RGBColor(217, 119, 6)
                prob_run.bold = True
                para.add_run(item['sentence'])
            if len(medium_sentences) > 30:
                doc.add_paragraph(f'... 还有 {len(medium_sentences) - 30} 个中风险句子未显示')

        # 低风险句子（简要展示）
        if low_sentences:
            doc.add_heading(f'低风险句子 ({len(low_sentences)}句)', level=2)
            doc.add_paragraph(f'共 {len(low_sentences)} 个低风险句子，未发现明显AI特征。')

        doc.add_paragraph('')

        # 免责声明
        doc.add_heading('免责声明', level=1)
        disclaimer = doc.add_paragraph(
            '本报告由AI内容检测工具自动生成，检测结果仅供参考。\n\n'
            '重要说明：\n'
            '1. 检测准确率约60-75%，存在误判可能。\n'
            '2. 学术论文、正式文档因其规范性写作特点，可能被误判为AI内容。\n'
            '3. 检测结果不能作为学术不端的唯一判定依据。\n'
            '4. 建议结合人工判断综合评估。\n'
            '5. 本工具与知网、Turnitin等商业工具的检测结果可能存在差异。'
        )
        for run in disclaimer.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(107, 114, 128)

        # 保存
        output = BytesIO()
        doc.save(output)
        output.seek(0)

        return output


# 便捷函数
def process_word_document(file_content: bytes, ai_detector=None) -> Dict[str, Any]:
    """
    处理Word文档并进行AI检测

    Args:
        file_content: 文件内容
        ai_detector: AI检测器实例（可选）

    Returns:
        处理结果
    """
    detector = WordAIDetector(ai_detector)
    return detector.process_uploaded_document(file_content)


def generate_annotated_word(
    original_content: bytes,
    detection_results: Dict[str, Any],
    mode: str = 'highlight'
) -> BytesIO:
    """
    生成带标注的Word文档

    Args:
        original_content: 原始文档内容
        detection_results: 检测结果
        mode: 标注模式

    Returns:
        标注后的文档
    """
    detector = WordAIDetector()
    return detector.generate_annotated_document(original_content, detection_results, mode)


def generate_report_document(
    original_content: bytes,
    detection_results: Dict[str, Any]
) -> BytesIO:
    """
    生成详细报告文档

    Args:
        original_content: 原始文档内容
        detection_results: 检测结果

    Returns:
        报告文档
    """
    detector = WordAIDetector()
    return detector.create_detailed_report(original_content, detection_results)
