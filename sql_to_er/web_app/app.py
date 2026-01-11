# -*- coding: utf-8 -*-
"""
ER Diagram Web Application - Flask Backend
"""
from flask import Flask, render_template, request, jsonify, send_file, session, redirect, url_for, flash
from flask_cors import CORS
import sys
import os
import io
import base64
import json

import uuid
import time
import secrets
import re
from datetime import datetime
import requests
import threading

# 添加父目录到系统路径
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
# 添加虎皮椒支付模块路径
sys.path.insert(0, os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))), 'hupijiao-v3-python'))

from src import parse_sql
from src.er_model import build_er_model
from src.doc_generator import generate_html, generate_docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from user_manager import UserManager, login_required

# 导入配置模块
from app_config import config

# 导入虎皮椒支付类
try:
    # 添加虎皮椒支付模块路径
    hupi_path = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))), 'hupijiao-v3-python')
    if hupi_path not in sys.path:
        sys.path.append(hupi_path)

    # 动态导入模块
    import importlib.util
    spec = importlib.util.spec_from_file_location("hupijiao_module",
                                                os.path.join(hupi_path, 'hupijiao-v3-python.py'))
    hupijiao_module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(hupijiao_module)

    # 获取Hupi类
    Hupi = hupijiao_module.Hupi
    print("虎皮椒支付模块导入成功")

except Exception as e:
    print(f"虎皮椒支付模块导入失败: {e}")
    # 创建一个备用的Hupi类以防止错误
    hupi_config = config.get_hupi_config()
    class Hupi:
        def __init__(self):
            self.appid = hupi_config['appid']
            self.AppSecret = hupi_config['appsecret']

        def Pay(self, *args, **kwargs):
            raise Exception("虎皮椒支付模块未正确导入")

app = Flask(__name__)
CORS(app)
app.secret_key = config.SECRET_KEY  # 从环境变量加载

# 数据库配置 - 从环境变量加载
DB_CONFIG = config.get_db_config()

# 初始化用户管理器
user_manager = UserManager(DB_CONFIG)

# 初始化稳定的邮件服务
from stable_email_service import StableEmailService, PasswordResetManager

# 邮箱配置 - 从环境变量加载
EMAIL_CONFIG = config.get_email_config()

# 创建邮件服务和密码重置管理器
email_service = StableEmailService(**EMAIL_CONFIG)
password_reset_manager = PasswordResetManager(email_service, user_manager)

# 将用户管理器添加到应用上下文中
app.user_manager = user_manager

# 管理员模块
from admin_auth import AdminAuth
from admin_routes import create_admin_blueprint
from admin_stats import AdminStats

# 初始化管理员认证和统计
admin_auth = AdminAuth(DB_CONFIG)
admin_auth.init_admin_table()  # 初始化管理员表

# 初始化管理员统计模块
admin_stats = AdminStats(DB_CONFIG)

# 注册管理员蓝图
admin_blueprint = create_admin_blueprint(admin_auth, admin_stats, user_manager)
app.register_blueprint(admin_blueprint)

# DeepSeek API配置 - 从环境变量加载
DEEPSEEK_API_KEY = config.DEEPSEEK_API_KEY
DEEPSEEK_API_URL = config.DEEPSEEK_API_URL

# 内存中存储项目（实际应用中应使用数据库）
projects = {}

# 论文生成进度存储
paper_generation_tasks = {}


def get_user_info(user_id):
    """获取用户信息"""
    return user_manager.get_user_info(user_id)


def translate_database_terms_with_ai(entities_data, relationships_data=None):
    """
    使用DeepSeek AI智能翻译数据库表名、字段名和关系名
    """
    try:
        # 构建翻译请求的数据
        translation_request = {
            "entities": [],
            "relationships": []
        }

        for entity_name, entity in entities_data.items():
            entity_info = {
                "table_name": entity_name,
                "attributes": []
            }

            for attr in entity.attributes:
                entity_info["attributes"].append({
                    "name": attr.name,
                    "type": attr.data_type,
                    "comment": attr.comment if attr.comment else ""
                })

            translation_request["entities"].append(entity_info)

        # 添加关系信息
        if relationships_data:
            for rel in relationships_data:
                translation_request["relationships"].append({
                    "name": rel.name,
                    "from_entity": rel.from_entity,
                    "to_entity": rel.to_entity,
                    "type": rel.rel_type
                })

        # 构建AI提示词
        prompt = f"""
你是一个专业的数据库设计专家，请帮我将以下英文数据库表名、字段名和关系名翻译成合适的中文名称。

要求：
1. 表名翻译要准确反映业务含义，使用常见的中文术语
2. 字段名翻译要简洁明了，符合中文表达习惯
3. 关系名翻译要体现实体间的业务关系
4. 保持专业性，使用标准的数据库术语
5. 所有表名、字段名和关系名都必须翻译，不能遗漏
6. 翻译要符合中国人的表达习惯
7. 返回JSON格式，保持原有结构
8. 【重要】不要添加任何额外的字段，如"FOREIGN"、"PRIMARY"等标签
9. 【重要】只翻译提供的字段名，不要创造新的字段

数据库结构：
{json.dumps(translation_request, ensure_ascii=False, indent=2)}

常见翻译参考：
- Department → 部门
- Employee → 员工
- Project → 项目
- Employee_Project → 员工项目关联
- dept_id → 部门ID
- emp_id → 员工ID
- project_id → 项目ID
- name/emp_name/dept_name/project_name → 名称
- email → 邮箱
- hire_date → 入职日期
- salary → 薪资
- location → 位置
- start_date → 开始日期
- end_date → 结束日期
- budget → 预算
- role → 角色
- hours_worked → 工作小时数
- Employee_to_Department → 员工所属部门
- Employee_Project_to_Employee → 员工参与项目
- Employee_Project_to_Project → 项目包含员工

请返回以下格式的JSON：
{{
  "translations": {{
    "Department": {{
      "chinese_name": "部门",
      "attributes": {{
        "dept_id": "部门ID",
        "dept_name": "部门名称",
        "location": "位置"
      }}
    }},
    "Employee": {{
      "chinese_name": "员工",
      "attributes": {{
        "emp_id": "员工ID",
        "emp_name": "员工姓名",
        "email": "邮箱",
        "hire_date": "入职日期",
        "salary": "薪资",
        "dept_id": "部门ID"
      }}
    }}
  }},
  "relationship_translations": {{
    "Employee_to_Department": "员工所属部门",
    "Employee_Project_to_Employee": "员工参与项目",
    "Employee_Project_to_Project": "项目包含员工"
  }}
}}

注意：
1. 必须翻译所有提供的表名、字段名和关系名
2. 只返回JSON，不要包含其他文字说明
3. 确保JSON格式正确
4. 【严禁】添加任何额外的字段，如"FOREIGN"、"PRIMARY"、"外键"、"主键"等标签
5. 【严禁】创造不存在的字段名，只翻译提供的字段
"""

        # 调用DeepSeek API
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {DEEPSEEK_API_KEY}"
        }

        data = {
            "model": "deepseek-chat",
            "messages": [
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            "temperature": 0.3,
            "max_tokens": 2000
        }

        response = requests.post(DEEPSEEK_API_URL, headers=headers, json=data, timeout=90)

        if response.status_code == 200:
            result = response.json()
            ai_response = result['choices'][0]['message']['content']

            # 尝试解析AI返回的JSON
            try:
                # 提取JSON部分
                start_idx = ai_response.find('{')
                end_idx = ai_response.rfind('}') + 1

                if start_idx != -1 and end_idx != 0:
                    json_str = ai_response[start_idx:end_idx]
                    translations = json.loads(json_str)

                    # 应用翻译到实体数据
                    for entity_name, entity in entities_data.items():
                        if entity_name in translations.get('translations', {}):
                            translation_info = translations['translations'][entity_name]

                            # 设置表的中文名 - 使用comment字段，因为get_display_name()方法读取的是comment
                            if 'chinese_name' in translation_info:
                                entity.comment = translation_info['chinese_name']

                            # 设置字段的中文名
                            if 'attributes' in translation_info:
                                for attr in entity.attributes:
                                    if attr.name in translation_info['attributes']:
                                        attr.display_name = translation_info['attributes'][attr.name]

                    # 应用翻译到关系数据
                    if relationships_data and 'relationship_translations' in translations:
                        rel_translations = translations['relationship_translations']
                        for rel in relationships_data:
                            if rel.name in rel_translations:
                                rel.comment = rel_translations[rel.name]

                    app.logger.info("AI翻译成功应用")
                    return True
                else:
                    app.logger.warning("AI返回的JSON格式无效")
                    return False

            except json.JSONDecodeError as e:
                app.logger.error(f"AI翻译结果JSON解析失败: {e}")
                app.logger.error(f"AI响应内容: {ai_response}")
                return False
        else:
            app.logger.error(f"DeepSeek API调用失败: {response.status_code} - {response.text}")
            return False

    except Exception as e:
        app.logger.error(f"AI翻译过程中出错: {e}")
        return False

@app.route('/')
def homepage():
    """工具集首页"""
    user_info = None
    if 'user_id' in session:
        user_info = get_user_info(session['user_id'])
    return render_template('homepage.html', user_info=user_info)

@app.route('/announcements')
def announcements_page():
    """公告页面"""
    user_info = None
    if 'user_id' in session:
        user_info = get_user_info(session['user_id'])
    return render_template('announcements.html', user_info=user_info)

@app.route('/services')
def services_page():
    """技术服务推广页面"""
    user_info = None
    if 'user_id' in session:
        user_info = get_user_info(session['user_id'])
    return render_template('services.html', user_info=user_info)

@app.route('/api/announcements')
def api_get_announcements():
    """获取有效公告列表API"""
    try:
        conn = user_manager.get_db_connection()
        with conn.cursor() as cursor:
            # 获取当前有效的公告
            current_time = datetime.now()
            cursor.execute("""
                SELECT id, title, content, type, is_sticky, start_time, end_time, view_count, comment_count, created_at
                FROM announcements
                WHERE is_active = 1
                AND (start_time IS NULL OR start_time <= %s)
                AND (end_time IS NULL OR end_time >= %s)
                ORDER BY is_sticky DESC, created_at DESC
                LIMIT 10
            """, (current_time, current_time))

            announcements = cursor.fetchall()

            # 格式化时间
            for announcement in announcements:
                if announcement['created_at']:
                    announcement['created_at'] = announcement['created_at'].strftime('%Y-%m-%d %H:%M:%S')
                if announcement['start_time']:
                    announcement['start_time'] = announcement['start_time'].strftime('%Y-%m-%d %H:%M:%S')
                if announcement['end_time']:
                    announcement['end_time'] = announcement['end_time'].strftime('%Y-%m-%d %H:%M:%S')

            return jsonify({'success': True, 'data': announcements})

    except Exception as e:
        app.logger.error(f"获取公告列表失败: {e}")
        return jsonify({'success': False, 'message': '获取公告失败'})
    finally:
        if 'conn' in locals():
            conn.close()

@app.route('/api/announcements/<int:announcement_id>/view', methods=['POST'])
def api_view_announcement(announcement_id):
    """增加公告查看次数"""
    try:
        conn = user_manager.get_db_connection()
        with conn.cursor() as cursor:
            cursor.execute("""
                UPDATE announcements
                SET view_count = view_count + 1
                WHERE id = %s AND is_active = 1
            """, (announcement_id,))
            conn.commit()

            return jsonify({'success': True})

    except Exception as e:
        app.logger.error(f"更新公告查看次数失败: {e}")
        return jsonify({'success': False, 'message': '操作失败'})
    finally:
        if 'conn' in locals():
            conn.close()


# ============ 公告评论相关API ============

@app.route('/api/announcements/<int:announcement_id>/comments', methods=['GET'])
def api_get_comments(announcement_id):
    """获取公告评论列表"""
    try:
        conn = user_manager.get_db_connection()
        with conn.cursor() as cursor:
            cursor.execute("""
                SELECT id, announcement_id, user_id, username, content, created_at
                FROM announcement_comments
                WHERE announcement_id = %s AND is_deleted = 0
                ORDER BY created_at DESC
            """, (announcement_id,))

            comments = cursor.fetchall()

            return jsonify({
                'success': True,
                'comments': comments
            })

    except Exception as e:
        app.logger.error(f"获取评论失败: {e}")
        return jsonify({'success': False, 'message': '获取评论失败'})
    finally:
        if 'conn' in locals():
            conn.close()


@app.route('/api/announcements/<int:announcement_id>/comments', methods=['POST'])
@login_required
def api_add_comment(announcement_id):
    """发表评论"""
    try:
        data = request.get_json()
        content = data.get('content', '').strip()

        if not content:
            return jsonify({'success': False, 'message': '评论内容不能为空'})

        if len(content) > 500:
            return jsonify({'success': False, 'message': '评论内容不能超过500字'})

        user_id = session.get('user_id')
        username = session.get('username')

        conn = user_manager.get_db_connection()
        with conn.cursor() as cursor:
            # 检查公告是否存在
            cursor.execute("SELECT id FROM announcements WHERE id = %s AND is_active = 1", (announcement_id,))
            if not cursor.fetchone():
                return jsonify({'success': False, 'message': '公告不存在'})

            # 插入评论
            cursor.execute("""
                INSERT INTO announcement_comments (announcement_id, user_id, username, content)
                VALUES (%s, %s, %s, %s)
            """, (announcement_id, user_id, username, content))

            # 更新公告的评论数
            cursor.execute("""
                UPDATE announcements
                SET comment_count = comment_count + 1
                WHERE id = %s
            """, (announcement_id,))

            conn.commit()

            return jsonify({'success': True, 'message': '评论发表成功'})

    except Exception as e:
        app.logger.error(f"发表评论失败: {e}")
        return jsonify({'success': False, 'message': '发表评论失败'})
    finally:
        if 'conn' in locals():
            conn.close()


@app.route('/api/announcements/comments/<int:comment_id>', methods=['DELETE'])
@login_required
def api_delete_comment(comment_id):
    """删除评论"""
    try:
        user_id = session.get('user_id')

        conn = user_manager.get_db_connection()
        with conn.cursor() as cursor:
            # 检查评论是否属于当前用户
            cursor.execute("""
                SELECT id, announcement_id, user_id
                FROM announcement_comments
                WHERE id = %s AND is_deleted = 0
            """, (comment_id,))

            comment = cursor.fetchone()
            if not comment:
                return jsonify({'success': False, 'message': '评论不存在'})

            if comment['user_id'] != user_id:
                return jsonify({'success': False, 'message': '无权删除此评论'})

            # 软删除评论
            cursor.execute("""
                UPDATE announcement_comments
                SET is_deleted = 1
                WHERE id = %s
            """, (comment_id,))

            # 更新公告的评论数
            cursor.execute("""
                UPDATE announcements
                SET comment_count = GREATEST(comment_count - 1, 0)
                WHERE id = %s
            """, (comment['announcement_id'],))

            conn.commit()

            return jsonify({'success': True, 'message': '评论删除成功'})

    except Exception as e:
        app.logger.error(f"删除评论失败: {e}")
        return jsonify({'success': False, 'message': '删除评论失败'})
    finally:
        if 'conn' in locals():
            conn.close()


@app.route('/sql-to-er')
def sql_to_er():
    """SQL转ER图工具"""
    return render_template('index.html')


@app.route('/paper-structure')
def paper_structure():
    """系统功能结构图生成器"""
    return render_template('paper-structure.html')


@app.route('/test-case-generator')
def test_case_generator():
    """AI测试用例生成器"""
    return render_template('test-case-generator.html')


@app.route('/thesis-defense')
def thesis_defense():
    """论文答辩问题生成器"""
    return render_template('thesis-defense-new.html')


@app.route('/text-optimizer')
def text_optimizer_page():
    """文本优化器页面"""
    user_info = None
    if 'user_id' in session:
        user_info = get_user_info(session['user_id'])
    return render_template('text-optimizer.html', user_info=user_info)


@app.route('/progress-test')
def progress_test():
    """进度测试页面"""
    return render_template('progress_test.html')


# ==================== AI内容检测功能 ====================

# 导入AI检测模块
from ai_detector import get_detector, detect_ai_content

@app.route('/ai-detector')
def ai_detector_page():
    """AI内容检测页面"""
    user_info = None
    if 'user_id' in session:
        user_info = get_user_info(session['user_id'])
    return render_template('ai-detector.html', user_info=user_info)


@app.route('/api/detect-ai', methods=['POST'])
def api_detect_ai():
    """AI内容检测API"""
    try:
        data = request.get_json()
        text = data.get('text', '').strip()
        detailed = data.get('detailed', False)

        if not text:
            return jsonify({'success': False, 'message': '请输入待检测的文本'})

        if len(text) < 50:
            return jsonify({'success': False, 'message': '文本太短，至少需要50个字符'})

        if len(text) > 50000:
            return jsonify({'success': False, 'message': '文本太长，最多支持50000个字符'})

        # 检查用户是否登录并扣费
        if 'user_id' in session:
            # 获取检测费用配置
            cost = user_manager.get_system_config('ai_detect_cost', 0.5)

            # 检查余额
            user_info = user_manager.get_user_info(session['user_id'])
            if user_info and user_info['balance'] < cost:
                return jsonify({
                    'success': False,
                    'message': f'余额不足，AI检测需要 {cost} 元，当前余额 {user_info["balance"]} 元'
                })

            # 扣费
            consume_result = user_manager.consume_balance(
                session['user_id'],
                cost,
                'ai_detect',
                f'AI内容检测 - {len(text)}字'
            )
            if not consume_result['success']:
                return jsonify(consume_result)

        # 执行检测
        result = detect_ai_content(text, detailed=detailed)

        if result.get('success'):
            app.logger.info(f"AI检测完成: 文本长度={len(text)}, AI概率={result.get('ai_probability')}%")

        return jsonify(result)

    except Exception as e:
        app.logger.error(f"AI检测失败: {e}")
        return jsonify({'success': False, 'message': f'检测失败: {str(e)}'})


@app.route('/api/detect-ai-sentences', methods=['POST'])
def api_detect_ai_sentences():
    """逐句AI检测API - 标注每个句子的AI概率"""
    try:
        data = request.get_json()
        text = data.get('text', '').strip()

        if not text:
            return jsonify({'success': False, 'message': '请输入待检测的文本'})

        if len(text) < 50:
            return jsonify({'success': False, 'message': '文本太短，至少需要50个字符'})

        if len(text) > 20000:
            return jsonify({'success': False, 'message': '逐句检测最多支持20000个字符'})

        # 检查用户是否登录并扣费（逐句检测费用更高）
        if 'user_id' in session:
            cost = user_manager.get_system_config('ai_detect_sentence_cost', 1.0)

            user_info = user_manager.get_user_info(session['user_id'])
            if user_info and user_info['balance'] < cost:
                return jsonify({
                    'success': False,
                    'message': f'余额不足，逐句检测需要 {cost} 元'
                })

            consume_result = user_manager.consume_balance(
                session['user_id'],
                cost,
                'ai_detect_sentence',
                f'AI逐句检测 - {len(text)}字'
            )
            if not consume_result['success']:
                return jsonify(consume_result)

        # 执行逐句检测
        detector = get_detector()
        sentences = detector.detect_sentences(text)

        return jsonify({
            'success': True,
            'sentences': sentences,
            'total_sentences': len(sentences)
        })

    except Exception as e:
        app.logger.error(f"逐句AI检测失败: {e}")
        return jsonify({'success': False, 'message': f'检测失败: {str(e)}'})


@app.route('/api/detect-ai-free', methods=['POST'])
def api_detect_ai_free():
    """AI内容检测API - 免费版（有限制）"""
    try:
        data = request.get_json()
        text = data.get('text', '').strip()

        if not text:
            return jsonify({'success': False, 'message': '请输入待检测的文本'})

        if len(text) < 50:
            return jsonify({'success': False, 'message': '文本太短，至少需要50个字符'})

        # 免费版限制1000字
        if len(text) > 1000:
            return jsonify({
                'success': False,
                'message': '免费版最多支持1000字，请登录后使用完整版',
                'need_login': True
            })

        # 执行简化检测
        result = detect_ai_content(text, detailed=False)

        # 免费版不返回详细信息
        if result.get('success'):
            return jsonify({
                'success': True,
                'ai_probability': result.get('ai_probability'),
                'verdict': result.get('verdict'),
                'verdict_level': result.get('verdict_level'),
                'language': result.get('language'),
                'is_free': True,
                'message': '免费版检测完成，登录后可使用详细分析功能'
            })

        return jsonify(result)

    except Exception as e:
        app.logger.error(f"免费AI检测失败: {e}")
        return jsonify({'success': False, 'message': f'检测失败: {str(e)}'})


@app.route('/api/detect-ai-report', methods=['POST'])
def api_detect_ai_report():
    """生成AI检测报告 - 逐句分析并返回统计摘要"""
    try:
        data = request.get_json()
        text = data.get('text', '').strip()

        if not text:
            return jsonify({'success': False, 'message': '请输入待检测的文本'})

        if len(text) < 50:
            return jsonify({'success': False, 'message': '文本太短，至少需要50个字符'})

        max_len = 50000 if 'user_id' in session else 1000
        if len(text) > max_len:
            return jsonify({
                'success': False,
                'message': f'文本太长，最多支持{max_len}字符',
                'need_login': 'user_id' not in session
            })

        # 执行逐句检测
        detector = get_detector()
        sentences = detector.detect_sentences(text)

        # 如果逐句检测失败（可能模型未加载），使用简化版
        if not sentences:
            # 使用简化方式分句和估算
            import re
            raw_sentences = re.split(r'[。！？.!?;；]+', text)
            sentences = []
            for i, sent in enumerate(raw_sentences):
                sent = sent.strip()
                if len(sent) > 5:
                    # 基于句子特征简单估算AI概率
                    ai_prob = estimate_sentence_ai_prob(sent)
                    sentences.append({
                        'index': i,
                        'sentence': sent,
                        'ai_probability': ai_prob
                    })

        # 计算统计摘要
        total = len(sentences)
        high_count = sum(1 for s in sentences if s['ai_probability'] >= 80)
        medium_count = sum(1 for s in sentences if 60 <= s['ai_probability'] < 80)
        low_count = sum(1 for s in sentences if s['ai_probability'] < 60)

        return jsonify({
            'success': True,
            'sentences': sentences,
            'summary': {
                'total': total,
                'high': high_count,
                'medium': medium_count,
                'low': low_count
            }
        })

    except Exception as e:
        app.logger.error(f"生成AI检测报告失败: {e}")
        return jsonify({'success': False, 'message': f'生成报告失败: {str(e)}'})


def estimate_sentence_ai_prob(sentence):
    """简化版句子AI概率估算（当模型不可用时）"""
    import re

    ai_markers = [
        '综上所述', '总而言之', '由此可见', '基于以上', '值得注意',
        '首先', '其次', '再次', '最后', '此外', '进一步',
        '具有重要意义', '具有深远影响', '不可否认',
        'in conclusion', 'furthermore', 'moreover', 'additionally',
        'it is worth noting', 'it should be noted'
    ]

    score = 30  # 基础分

    # 检查AI标记词
    for marker in ai_markers:
        if marker in sentence.lower():
            score += 15
            break

    # 句子长度均匀性（20-40字的句子更可能是AI）
    length = len(sentence)
    if 20 <= length <= 40:
        score += 10

    # 检查过于工整的结构
    if re.search(r'[，,].*[，,].*[，,]', sentence):  # 多个逗号，结构工整
        score += 10

    return min(score, 95)


@app.route('/api/export-ai-report', methods=['POST'])
def api_export_ai_report():
    """导出AI检测报告为Word文档"""
    try:
        # 检查登录
        if 'user_id' not in session:
            return jsonify({'success': False, 'message': '请登录后使用导出功能'}), 401

        data = request.get_json()
        text = data.get('text', '')
        report_data = data.get('report_data', {})
        export_format = data.get('format', 'word')

        if not report_data:
            return jsonify({'success': False, 'message': '报告数据不能为空'})

        # 生成Word文档
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from io import BytesIO

        doc = Document()

        # 标题
        title = doc.add_heading('AI内容检测报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 检测时间
        from datetime import datetime
        doc.add_paragraph(f'检测时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph(f'文本长度: {len(text)} 字符')
        doc.add_paragraph('')

        # 整体结果
        overall = report_data.get('overall', {})
        doc.add_heading('一、整体检测结果', level=1)

        result_para = doc.add_paragraph()
        result_para.add_run(f'AI概率: ').bold = True
        prob = overall.get('ai_probability', 0)
        prob_run = result_para.add_run(f'{prob}%')
        prob_run.bold = True
        if prob >= 80:
            prob_run.font.color.rgb = RGBColor(239, 68, 68)  # 红色
        elif prob >= 60:
            prob_run.font.color.rgb = RGBColor(245, 158, 11)  # 黄色
        else:
            prob_run.font.color.rgb = RGBColor(16, 185, 129)  # 绿色

        doc.add_paragraph(f'判定结果: {overall.get("verdict", "")}')
        doc.add_paragraph('')

        # 统计摘要
        summary = report_data.get('summary', {})
        doc.add_heading('二、统计摘要', level=1)

        table = doc.add_table(rows=2, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '总句数'
        hdr_cells[1].text = '高风险句(≥80%)'
        hdr_cells[2].text = '中风险句(60-79%)'
        hdr_cells[3].text = '低风险句(<60%)'

        data_cells = table.rows[1].cells
        data_cells[0].text = str(summary.get('total', 0))
        data_cells[1].text = str(summary.get('high', 0))
        data_cells[2].text = str(summary.get('medium', 0))
        data_cells[3].text = str(summary.get('low', 0))

        doc.add_paragraph('')

        # 逐句分析
        doc.add_heading('三、逐句分析详情', level=1)

        sentences = report_data.get('sentences', [])
        for i, item in enumerate(sentences):
            prob = item.get('ai_probability', 0)
            sent = item.get('sentence', '')

            # 确定风险等级标记
            if prob >= 80:
                level_mark = '[高风险]'
                color = RGBColor(239, 68, 68)
            elif prob >= 60:
                level_mark = '[中风险]'
                color = RGBColor(245, 158, 11)
            else:
                level_mark = '[低风险]'
                color = RGBColor(16, 185, 129)

            para = doc.add_paragraph()
            run = para.add_run(f'{i+1}. {level_mark} {prob}% ')
            run.font.color.rgb = color
            run.bold = True
            para.add_run(sent)

        doc.add_paragraph('')

        # 免责声明
        doc.add_heading('免责声明', level=1)
        disclaimer = doc.add_paragraph(
            '本报告由AI内容检测工具自动生成，检测准确率约60-75%，存在误判可能。'
            '检测结果不能作为学术不端的唯一判定依据，建议结合人工判断综合评估。'
        )
        disclaimer.runs[0].font.size = Pt(10)
        disclaimer.runs[0].font.color.rgb = RGBColor(107, 114, 128)

        # 保存到内存
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        return send_file(
            file_stream,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'AI检测报告_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        )

    except Exception as e:
        app.logger.error(f"导出AI检测报告失败: {e}")
        return jsonify({'success': False, 'message': f'导出失败: {str(e)}'}), 500


# Word文档上传检测功能
from word_ai_detector import (
    WordAIDetector,
    process_word_document,
    generate_annotated_word,
    generate_report_document
)

# 存储上传的文档（实际应用中应使用数据库或文件系统）
uploaded_documents = {}


@app.route('/api/upload-word-detect', methods=['POST'])
def api_upload_word_detect():
    """
    上传Word文档进行AI检测
    返回检测结果，不直接返回文件
    """
    try:
        # 检查文件
        if 'file' not in request.files:
            return jsonify({'success': False, 'message': '请选择要上传的Word文档'})

        file = request.files['file']
        if file.filename == '':
            return jsonify({'success': False, 'message': '请选择要上传的Word文档'})

        # 检查文件类型
        if not file.filename.lower().endswith(('.docx', '.doc')):
            return jsonify({'success': False, 'message': '只支持.docx或.doc格式的Word文档'})

        # 检查文件大小（最大10MB）
        file_content = file.read()
        if len(file_content) > 10 * 1024 * 1024:
            return jsonify({'success': False, 'message': '文件太大，最大支持10MB'})

        # 检查用户登录状态
        is_logged_in = 'user_id' in session
        if not is_logged_in:
            # 免费用户限制
            if len(file_content) > 100 * 1024:  # 100KB限制
                return jsonify({
                    'success': False,
                    'message': '免费用户文件大小限制100KB，请登录后使用完整版',
                    'need_login': True
                })

        # 登录用户扣费
        if is_logged_in:
            cost = user_manager.get_system_config('word_detect_cost', 2.0)
            user_info = user_manager.get_user_info(session['user_id'])
            if user_info and user_info['balance'] < cost:
                return jsonify({
                    'success': False,
                    'message': f'余额不足，Word文档检测需要 {cost} 元，当前余额 {user_info["balance"]} 元'
                })

        # 获取AI检测器
        try:
            detector = get_detector()
        except Exception:
            detector = None

        # 处理文档
        word_detector = WordAIDetector(detector)
        result = word_detector.process_uploaded_document(file_content)

        if not result['success']:
            return jsonify(result)

        # 登录用户扣费
        if is_logged_in:
            consume_result = user_manager.consume_balance(
                session['user_id'],
                cost,
                'word_detect',
                f'Word文档AI检测 - {file.filename}'
            )
            if not consume_result['success']:
                return jsonify(consume_result)

        # 生成文档ID，用于后续下载
        doc_id = str(uuid.uuid4())
        uploaded_documents[doc_id] = {
            'original_content': file_content,
            'filename': file.filename,
            'detection_result': result,
            'upload_time': datetime.now(),
            'user_id': session.get('user_id')
        }

        # 清理过期文档（超过1小时）
        cleanup_old_documents()

        # 返回结果
        return jsonify({
            'success': True,
            'doc_id': doc_id,
            'filename': file.filename,
            'text_length': result.get('text_length', 0),
            'overall_probability': result.get('overall_probability', 0),
            'overall_level': result.get('overall_level', 'low'),
            'summary': result.get('summary', {}),
            'sentences': result.get('sentences', [])
        })

    except Exception as e:
        app.logger.error(f"Word文档上传检测失败: {e}")
        return jsonify({'success': False, 'message': f'处理失败: {str(e)}'})


@app.route('/api/download-annotated-word/<doc_id>')
def api_download_annotated_word(doc_id):
    """
    下载带标注的Word文档（保持原格式）
    """
    try:
        # 检查登录
        if 'user_id' not in session:
            return jsonify({'success': False, 'message': '请登录后下载'}), 401

        # 检查文档是否存在
        if doc_id not in uploaded_documents:
            return jsonify({'success': False, 'message': '文档不存在或已过期'}), 404

        doc_data = uploaded_documents[doc_id]

        # 检查权限
        if doc_data.get('user_id') != session['user_id']:
            return jsonify({'success': False, 'message': '无权访问此文档'}), 403

        # 生成带标注的文档
        annotated_doc = generate_annotated_word(
            doc_data['original_content'],
            doc_data['detection_result'],
            mode='highlight'
        )

        # 生成文件名
        original_name = doc_data['filename']
        base_name = original_name.rsplit('.', 1)[0]
        download_name = f'{base_name}_AI检测标注.docx'

        return send_file(
            annotated_doc,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=download_name
        )

    except Exception as e:
        app.logger.error(f"下载标注文档失败: {e}")
        return jsonify({'success': False, 'message': f'下载失败: {str(e)}'}), 500


@app.route('/api/download-detection-report/<doc_id>')
def api_download_detection_report(doc_id):
    """
    下载详细检测报告（独立报告文档）
    """
    try:
        # 检查登录
        if 'user_id' not in session:
            return jsonify({'success': False, 'message': '请登录后下载'}), 401

        # 检查文档是否存在
        if doc_id not in uploaded_documents:
            return jsonify({'success': False, 'message': '文档不存在或已过期'}), 404

        doc_data = uploaded_documents[doc_id]

        # 检查权限
        if doc_data.get('user_id') != session['user_id']:
            return jsonify({'success': False, 'message': '无权访问此文档'}), 403

        # 生成报告文档
        report_doc = generate_report_document(
            doc_data['original_content'],
            doc_data['detection_result']
        )

        # 生成文件名
        original_name = doc_data['filename']
        base_name = original_name.rsplit('.', 1)[0]
        download_name = f'{base_name}_AI检测报告.docx'

        return send_file(
            report_doc,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=download_name
        )

    except Exception as e:
        app.logger.error(f"下载检测报告失败: {e}")
        return jsonify({'success': False, 'message': f'下载失败: {str(e)}'}), 500


def cleanup_old_documents():
    """清理超过1小时的临时文档"""
    current_time = datetime.now()
    expired_ids = []

    for doc_id, doc_data in uploaded_documents.items():
        upload_time = doc_data.get('upload_time')
        if upload_time and (current_time - upload_time).total_seconds() > 3600:
            expired_ids.append(doc_id)

    for doc_id in expired_ids:
        del uploaded_documents[doc_id]

    if expired_ids:
        app.logger.info(f"清理了 {len(expired_ids)} 个过期文档")


















# ==================== 用户系统路由 ====================

@app.route('/login')
def login():
    """登录页面"""
    return render_template('login.html')

@app.route('/register')
def register():
    """注册页面"""
    return render_template('register.html')

@app.route('/profile')
@login_required
def profile():
    """个人中心"""
    user_info = user_manager.get_user_info(session['user_id'])
    consumption_data = user_manager.get_consumption_records(session['user_id'], page=1, per_page=10)
    recharge_data = user_manager.get_recharge_records(session['user_id'], page=1, per_page=10)
    return render_template('profile.html',
                        user_info=user_info,
                        records=consumption_data['records'],
                        consumption_total=consumption_data['total'],
                        consumption_pages=consumption_data['total_pages'],
                        recharge_records=recharge_data['records'],
                        recharge_total=recharge_data['total'],
                        recharge_pages=recharge_data['total_pages'])

@app.route('/api/consumption-records')
@login_required
def api_consumption_records():
    """获取消费记录API（分页）"""
    try:
        page = request.args.get('page', 1, type=int)
        per_page = request.args.get('per_page', 10, type=int)
        per_page = min(per_page, 50)

        data = user_manager.get_consumption_records(session['user_id'], page=page, per_page=per_page)

        records = []
        for record in data['records']:
            records.append({
                'created_at': record['created_at'].strftime('%Y-%m-%d %H:%M:%S'),
                'service_type': record['service_type'],
                'amount': float(record['amount']),
                'description': record['description'] or '-'
            })

        return jsonify({
            'success': True,
            'records': records,
            'total': data['total'],
            'page': data['page'],
            'per_page': data['per_page'],
            'total_pages': data['total_pages']
        })
    except Exception as e:
        logger.error(f"获取消费记录失败: {e}")
        return jsonify({'success': False, 'message': '获取记录失败'})

@app.route('/api/recharge-records')
@login_required
def api_recharge_records():
    """获取充值记录API（分页）"""
    try:
        page = request.args.get('page', 1, type=int)
        per_page = request.args.get('per_page', 10, type=int)
        per_page = min(per_page, 50)

        data = user_manager.get_recharge_records(session['user_id'], page=page, per_page=per_page)

        records = []
        for record in data['records']:
            records.append({
                'created_at': record['created_at'].strftime('%Y-%m-%d %H:%M:%S'),
                'payment_method': record['payment_method'],
                'amount': float(record['amount']),
                'description': record['description'] or '-'
            })

        return jsonify({
            'success': True,
            'records': records,
            'total': data['total'],
            'page': data['page'],
            'per_page': data['per_page'],
            'total_pages': data['total_pages']
        })
    except Exception as e:
        logger.error(f"获取充值记录失败: {e}")
        return jsonify({'success': False, 'message': '获取记录失败'})

@app.route('/api/register', methods=['POST'])
def api_register():
    """用户注册API - 支持用户名或邮箱注册"""
    try:
        data = request.get_json()
        username = data.get('username', '').strip()
        email = data.get('email', '').strip()
        password = data.get('password', '').strip()
        confirm_password = data.get('confirm_password', '').strip()
        invite_code = data.get('invite_code', '').strip()

        # 基本验证
        if not password:
            return jsonify({'success': False, 'message': '密码不能为空'})

        if len(password) < 6:
            return jsonify({'success': False, 'message': '密码长度不能少于6个字符'})

        if password != confirm_password:
            return jsonify({'success': False, 'message': '两次输入的密码不一致'})

        # 必须提供用户名或邮箱中的一个
        if not username and not email:
            return jsonify({'success': False, 'message': '请输入用户名或邮箱'})

        # 用户名验证（如果提供）
        if username:
            if len(username) < 3 or len(username) > 20:
                return jsonify({'success': False, 'message': '用户名长度应在3-20个字符之间'})

        # 邮箱验证（如果提供）
        if email:
            if not user_manager.is_valid_email(email):
                return jsonify({'success': False, 'message': '邮箱格式不正确'})

        result = user_manager.register_user(
            username=username if username else None,
            email=email if email else None,
            password=password,
            invite_code=invite_code if invite_code else None,
            ip_address=request.remote_addr  # 传递IP用于防刷检查
        )
        return jsonify(result)

    except Exception as e:
        app.logger.error(f"注册失败: {e}")
        return jsonify({'success': False, 'message': '注册失败，请稍后重试'})

@app.route('/api/login', methods=['POST'])
def api_login():
    """用户登录API - 支持用户名或邮箱登录"""
    try:
        data = request.get_json()
        username_or_email = data.get('username', '').strip()  # 可以是用户名或邮箱
        password = data.get('password', '').strip()
        captcha = data.get('captcha', '').strip()  # 验证码

        if not username_or_email or not password:
            return jsonify({'success': False, 'message': '用户名/邮箱和密码不能为空'})

        result = user_manager.login_user(username_or_email, password, captcha)

        if result['success']:
            # 设置session
            session['user_id'] = result['user']['id']
            session['username'] = result['user']['username']
            session['user_role'] = result['user']['role']
            session['login_ip'] = request.environ.get('REMOTE_ADDR', 'unknown')
            session.permanent = True

        return jsonify(result)

    except Exception as e:
        app.logger.error(f"登录失败: {e}")
        return jsonify({'success': False, 'message': '登录失败，请稍后重试'})

@app.route('/api/logout', methods=['POST'])
def api_logout():
    """用户登出API"""
    session.clear()
    return jsonify({'success': True, 'message': '已退出登录'})

@app.route('/api/csrf-token', methods=['GET'])
def api_csrf_token():
    """获取CSRF令牌API"""
    try:
        # 生成或获取CSRF令牌
        if 'csrf_token' not in session:
            session['csrf_token'] = secrets.token_hex(32)
        return jsonify({'csrf_token': session['csrf_token']})
    except Exception as e:
        app.logger.error(f"获取CSRF令牌失败: {e}")
        return jsonify({'csrf_token': ''})

@app.route('/api/login-captcha', methods=['GET'])
def api_login_captcha():
    """获取登录验证码API"""
    try:
        captcha_text, captcha_image = user_manager.login_security.generate_captcha()
        user_manager.login_security.store_captcha(captcha_text)

        if captcha_image:
            return jsonify({
                'success': True,
                'captcha_text': captcha_text,  # 仅用于调试，生产环境应移除
                'captcha_image': captcha_image
            })
        else:
            # 降级到文本验证码
            return jsonify({
                'success': True,
                'captcha_text': captcha_text,
                'captcha_image': None
            })
    except Exception as e:
        app.logger.error(f"生成验证码失败: {e}")
        return jsonify({'success': False, 'message': '验证码生成失败'})

# ==================== 稳定的密码重置功能 ====================

@app.route('/forgot-password')
def forgot_password():
    """忘记密码页面"""
    return render_template('forgot_password.html')

@app.route('/api/send-reset-code', methods=['POST'])
def api_send_reset_code():
    """发送密码重置验证码API"""
    try:
        data = request.get_json()
        email = data.get('email', '').strip()

        if not email:
            return jsonify({'success': False, 'message': '请输入邮箱地址'})

        result = password_reset_manager.send_reset_code(email)
        return jsonify(result)

    except Exception as e:
        app.logger.error(f"发送验证码失败: {e}")
        return jsonify({'success': False, 'message': '发送失败，请稍后重试'})

@app.route('/api/reset-password', methods=['POST'])
def api_reset_password():
    """重置密码API"""
    try:
        data = request.get_json()
        email = data.get('email', '').strip()
        code = data.get('code', '').strip()
        new_password = data.get('new_password', '').strip()

        if not all([email, code, new_password]):
            return jsonify({'success': False, 'message': '请填写完整信息'})

        result = password_reset_manager.verify_and_reset(email, code, new_password)
        return jsonify(result)

    except Exception as e:
        app.logger.error(f"重置密码失败: {e}")
        return jsonify({'success': False, 'message': '重置失败，请稍后重试'})

@app.route('/api/recharge', methods=['POST'])
@login_required
def api_recharge():
    """充值API - 使用虎皮椒支付（仅支持支付宝）"""
    try:
        amount = float(request.form.get('amount', 0))
        payment_method = request.form.get('payment_method', 'alipay')

        # 只支持支付宝
        if payment_method != 'alipay':
            flash('当前仅支持支付宝支付', 'danger')
            return redirect('/profile')

        # 金额验证：最小0.01元，最大10000元
        if amount < 0.01:
            flash('充值金额不能小于0.01元', 'danger')
            return redirect('/profile')

        if amount > 10000:
            flash('单次充值金额不能超过10000元', 'danger')
            return redirect('/profile')

        # 金额精度处理：保留两位小数
        amount = round(amount, 2)

        # 1. 先创建待支付订单
        order_result = user_manager.create_pending_order(
            user_id=session['user_id'],
            amount=amount,
            payment_method=payment_method
        )

        if not order_result['success']:
            app.logger.error(f"创建订单失败: {order_result['message']}")
            flash('订单创建失败，请稍后重试', 'danger')
            return redirect('/profile')

        trade_order_id = order_result['trade_order_id']
        order_no = order_result['order_no']

        app.logger.info(f"创建待支付订单成功: order_no={order_no}, trade_order_id={trade_order_id}")

        # 2. 创建虎皮椒支付实例
        try:
            hupi = Hupi()
        except Exception as e:
            app.logger.error(f"虎皮椒支付实例创建失败: {e}")
            flash('支付系统初始化失败，请稍后重试', 'danger')
            return redirect('/profile')

        # 3. 调用虎皮椒支付接口
        try:
            app.logger.info(f"发起支付请求: 用户{session['user_id']}, 金额{amount}, 订单号{trade_order_id}")

            response = hupi.Pay(
                trade_order_id=trade_order_id,
                payment=payment_method,
                total_fee=amount,
                title=f"账户充值 - ¥{amount}",
                attach=f"user_{session['user_id']}"
            )

            if response.status_code == 200:
                try:
                    result = response.json()
                    app.logger.info(f"支付接口返回: {result}")

                    if result.get('errcode') == 0:
                        # 支付请求成功，保存订单号到session用于成功页面展示
                        session['last_order_no'] = order_no
                        session['last_order_amount'] = amount

                        pay_url = result.get('url')
                        if pay_url:
                            app.logger.info(f"跳转到支付页面: {pay_url}")
                            return redirect(pay_url)
                        else:
                            flash('支付链接获取失败', 'danger')
                    else:
                        error_msg = result.get('errmsg', '未知错误')
                        app.logger.error(f'支付请求失败: {error_msg}')
                        flash(f'支付请求失败: {error_msg}', 'danger')

                except json.JSONDecodeError as e:
                    app.logger.error(f"支付接口返回数据解析失败: {e}")
                    flash('支付接口返回数据格式错误', 'danger')
            else:
                app.logger.error(f'支付接口调用失败: HTTP {response.status_code}')
                flash(f'支付接口调用失败: HTTP {response.status_code}', 'danger')

        except Exception as e:
            app.logger.error(f"虎皮椒支付调用异常: {e}")
            flash(f'支付系统异常: {str(e)}', 'danger')

    except ValueError:
        flash('请输入有效的充值金额', 'danger')
    except Exception as e:
        app.logger.error(f"充值处理失败: {e}")
        flash('充值处理失败，请稍后重试', 'danger')

    return redirect('/profile')

@app.route('/payment/success')
@app.route('/payment/success/')
def payment_success():
    """支付成功回调页面"""
    app.logger.info("用户访问支付成功页面")
    try:
        # 从session获取订单信息
        order_no = session.pop('last_order_no', None)
        order_amount = session.pop('last_order_amount', None)

        order_info = None
        if order_no:
            order_info = user_manager.get_order_by_order_no(order_no)

        return render_template('payment_success.html',
                             order_info=order_info,
                             order_no=order_no,
                             order_amount=order_amount)
    except Exception as e:
        app.logger.error(f"渲染支付成功页面失败: {e}")
        return f"<h1>支付成功！</h1><p>恭喜您，充值已成功完成！</p><a href='/profile'>查看余额</a>", 200

@app.route('/payment/callback')
@app.route('/payment/callback/')
def payment_callback():
    """支付失败/取消回调页面"""
    app.logger.info("用户访问支付失败/取消页面")
    try:
        return render_template('payment_callback.html')
    except Exception as e:
        app.logger.error(f"渲染支付失败页面失败: {e}")
        return f"<h1>支付未完成</h1><p>您取消了支付或支付未能成功完成。</p><a href='/profile'>返回个人中心</a>", 200

@app.route('/notify_url/', methods=['POST'])
def payment_notify():
    """虎皮椒支付异步通知处理"""
    try:
        # 获取通知数据
        data = request.form.to_dict()
        app.logger.info(f"收到支付通知: {data}")

        # 创建虎皮椒实例验证签名
        try:
            hupi = Hupi()
            # 验证签名
            if not hupi.verify_notify(data.copy()):
                app.logger.error("支付通知签名验证失败")
                return "fail"
        except Exception as e:
            app.logger.error(f"支付通知验证异常: {e}")
            return "fail"

        # 处理支付成功逻辑
        trade_order_id = data.get('trade_order_id', '')
        total_fee = float(data.get('total_fee', 0))
        status = data.get('status', '')
        transaction_id = data.get('transaction_id', '')

        app.logger.info(f"支付通知详情: 订单号={trade_order_id}, 金额={total_fee}, 状态={status}, 交易号={transaction_id}")

        if status == 'OD':  # 已支付
            # 使用新的complete_order方法，包含订单验证和金额核对
            result = user_manager.complete_order(
                trade_order_id=trade_order_id,
                transaction_id=transaction_id,
                callback_amount=total_fee
            )

            if result['success']:
                app.logger.info(f"支付处理成功: {trade_order_id}, {result['message']}")
                return "success"
            else:
                app.logger.error(f"支付处理失败: {trade_order_id}, {result['message']}")
                return "fail"
        else:
            app.logger.info(f"支付状态非成功: {status}")
            return "fail"

    except Exception as e:
        app.logger.error(f"处理支付通知失败: {e}")
        return "fail"

@app.route('/api/user/info')
@login_required
def api_user_info():
    """获取用户信息API"""
    try:
        user_info = user_manager.get_user_info(session['user_id'])
        if user_info:
            return jsonify({'success': True, 'user': user_info})
        else:
            return jsonify({'success': False, 'message': '用户不存在'})
    except Exception as e:
        app.logger.error(f"获取用户信息失败: {e}")
        return jsonify({'success': False, 'message': '获取用户信息失败'})


@app.route('/api/get_translation_price', methods=['GET'])
def api_get_translation_price():
    """获取AI翻译的价格和用户余额信息"""
    try:
        # 获取AI翻译费用配置
        ai_translation_cost = float(user_manager.get_system_config('ai_translation_cost', 2.00))
        
        # 检查用户是否登录
        if 'user_id' not in session:
            return jsonify({
                'success': True,
                'cost': ai_translation_cost,
                'logged_in': False,
                'balance': 0,
                'message': '请先登录后使用AI翻译功能'
            })
        
        # 获取用户余额
        user_info = user_manager.get_user_info(session['user_id'])
        balance = float(user_info['balance']) if user_info else 0
        
        return jsonify({
            'success': True,
            'cost': ai_translation_cost,
            'logged_in': True,
            'balance': balance,
            'sufficient': balance >= ai_translation_cost
        })
        
    except Exception as e:
        app.logger.error(f"获取翻译价格失败: {e}")
        return jsonify({
            'success': False,
            'message': '获取价格信息失败'
        }), 500


@app.route('/api/parse_sql', methods=['POST'])
def api_parse_sql():
    """解析SQL并生成ER图数据"""
    try:
        data = request.get_json()
        sql = data.get('sql', '')
        enable_translation = data.get('enableTranslation', True)  # 默认启用翻译

        # 解析SQL
        parsed_result, error_message = parse_sql(sql)

        if error_message:
            # 提供更友好的错误信息
            user_friendly_msg = error_message
            if "Expecting )" in error_message:
                user_friendly_msg = "SQL 语法错误：括号不匹配。请检查 CREATE TABLE 语句中的括号是否正确配对。"
            elif "No CREATE TABLE" in error_message:
                user_friendly_msg = "未找到有效的 CREATE TABLE 语句。请确保 SQL 中包含表定义。"

            return jsonify({
                'error': user_friendly_msg,
                'details': error_message
            }), 400

        if not parsed_result:
            return jsonify({'error': 'No CREATE TABLE statements found in the provided SQL.'}), 400

        tables = parsed_result
        # 构建ER模型
        entities, relationships = build_er_model(tables)

        # 如果启用翻译，需要检查用户登录状态和余额
        translation_success = False
        translation_charged = False
        if enable_translation:
            # 检查用户是否登录
            if 'user_id' not in session:
                return jsonify({
                    'error': 'AI翻译功能需要登录后使用',
                    'need_login': True
                }), 401
            
            # 获取AI翻译费用配置
            ai_translation_cost = float(user_manager.get_system_config('ai_translation_cost', 2.00))
            
            # 获取用户信息并检查余额
            user_info = user_manager.get_user_info(session['user_id'])
            if not user_info:
                return jsonify({
                    'error': '用户信息获取失败，请重新登录',
                    'need_login': True
                }), 401
            
            if user_info['balance'] < ai_translation_cost:
                return jsonify({
                    'error': f'余额不足，AI翻译需要 {ai_translation_cost:.2f} 元，当前余额 {user_info["balance"]:.2f} 元',
                    'need_recharge': True,
                    'cost': ai_translation_cost,
                    'balance': float(user_info['balance'])
                }), 400
            
            # 执行AI翻译
            try:
                translation_success = translate_database_terms_with_ai(entities, relationships)
                if translation_success:
                    # 翻译成功，扣除费用
                    consume_result = user_manager.consume_balance(
                        session['user_id'],
                        ai_translation_cost,
                        'ai_translation',
                        f'SQL转ER图 - AI智能翻译（{len(tables)}个表）'
                    )
                    if consume_result['success']:
                        translation_charged = True
                        app.logger.info(f"AI翻译成功，已扣费 {ai_translation_cost} 元")
                    else:
                        app.logger.warning(f"AI翻译扣费失败: {consume_result['message']}")
                else:
                    app.logger.warning("AI翻译失败，使用原始名称，不扣费")
            except Exception as e:
                app.logger.error(f"翻译过程中出错: {e}")

        # 转换为前端需要的格式
        result = {
            'entities': [],
            'relationships': [],
            'translationApplied': translation_success,  # 告知前端是否应用了翻译
            'translationCharged': translation_charged   # 告知前端是否已扣费
        }

        # 处理实体
        for entity_name, entity in entities.items():
            entity_data = {
                'name': entity_name,
                'displayName': entity.get_display_name(), # 添加实体显示名称
                'attributes': []
            }

            for attr in entity.attributes:
                entity_data['attributes'].append({
                    'name': attr.name,
                    'type': attr.data_type,
                    'isPK': attr.is_pk,
                    'isFK': attr.is_fk,                 # 新增: 外键标识
                    'comment': attr.comment,
                    'displayName': attr.display_name,  # 修复: 使用属性而非方法
                    'nullable': attr.nullable,          # 新增: 可空属性
                    'default': attr.default             # 新增: 默认值属性
                })

            result['entities'].append(entity_data)

        # 处理关系
        for rel in relationships:
            result['relationships'].append({
                'from': rel.from_entity,
                'to': rel.to_entity,
                'fromAttr': rel.from_attribute,
                'toAttr': rel.to_attribute,
                'type': rel.rel_type,
                'name': rel.name,
                'displayName': rel.get_display_name() # 添加关系显示名称
            })

        return jsonify(result)

    except Exception as e:
        app.logger.error(f"An unexpected error occurred in api_parse_sql: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/translate_terms', methods=['POST'])
def api_translate_terms():
    """专门用于翻译数据库术语的API"""
    try:
        data = request.get_json()
        terms = data.get('terms', [])

        if not terms:
            return jsonify({'error': '请提供需要翻译的术语'}), 400

        # 构建翻译提示词
        terms_text = '\n'.join([f"- {term}" for term in terms])

        prompt = f"""
请将以下英文数据库表名和字段名翻译成合适的中文名称：

{terms_text}

要求：
1. 翻译要准确反映业务含义
2. 使用常见的中文术语
3. 保持简洁明了
4. 返回JSON格式

请返回以下格式的JSON：
{{
  "translations": {{
    "英文术语1": "中文翻译1",
    "英文术语2": "中文翻译2"
  }}
}}

只返回JSON，不要包含其他文字。
"""

        # 调用DeepSeek API
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {DEEPSEEK_API_KEY}"
        }

        api_data = {
            "model": "deepseek-chat",
            "messages": [
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            "temperature": 0.3,
            "max_tokens": 1000
        }

        response = requests.post(DEEPSEEK_API_URL, headers=headers, json=api_data, timeout=90)

        if response.status_code == 200:
            result = response.json()
            ai_response = result['choices'][0]['message']['content']

            # 解析AI返回的JSON
            try:
                start_idx = ai_response.find('{')
                end_idx = ai_response.rfind('}') + 1

                if start_idx != -1 and end_idx != 0:
                    json_str = ai_response[start_idx:end_idx]
                    translations = json.loads(json_str)

                    return jsonify({
                        'success': True,
                        'translations': translations.get('translations', {}),
                        'message': '翻译成功'
                    })
                else:
                    return jsonify({'error': 'AI返回格式无效'}), 500

            except json.JSONDecodeError as e:
                app.logger.error(f"翻译结果JSON解析失败: {e}")
                return jsonify({'error': '翻译结果解析失败'}), 500
        else:
            app.logger.error(f"DeepSeek API调用失败: {response.status_code}")
            return jsonify({'error': 'AI翻译服务暂时不可用'}), 500

    except Exception as e:
        app.logger.error(f"翻译API出错: {e}")
        return jsonify({'error': f'翻译失败: {str(e)}'}), 500


@app.route('/api/generate_sql', methods=['POST'])
def api_generate_sql():
    """从ER图数据生成SQL"""
    try:
        data = request.get_json()
        entities = data.get('entities', [])
        relationships = data.get('relationships', [])
        db_type = data.get('dbType', 'mysql')

        sql = f"-- Generated by ER Diagram Editor\n"
        sql += f"-- Database: {db_type.upper()}\n\n"

        # 生成CREATE TABLE语句
        for entity in entities:
            sql += f"CREATE TABLE {entity['name']} (\n"

            # 添加列
            primary_keys = []
            for attr in entity['attributes']:
                col_type = attr.get('type', 'VARCHAR(255)')
                col_def = f"    {attr['name']} {col_type}"

                if attr.get('isPK'):
                    primary_keys.append(attr['name'])
                    if db_type == 'mysql':
                        col_def += " NOT NULL"

                sql += col_def + ",\n"

            # 添加主键约束
            if primary_keys:
                if len(primary_keys) == 1:
                    # 单主键
                    sql = sql.rstrip(",\n")
                    for i, line in enumerate(sql.split('\n')):
                        if primary_keys[0] in line and 'CREATE TABLE' not in line:
                            sql = sql.replace(line, line.rstrip(",") + " PRIMARY KEY,")
                            break
                else:
                    # 复合主键
                    pk_def = ", ".join(primary_keys)
                    sql += f"    PRIMARY KEY ({pk_def}),\n"

            # 处理外键
            for rel in relationships:
                if rel['from'] == entity['name']:
                    sql += f"    FOREIGN KEY ({rel['fromAttr']}) REFERENCES {rel['to']}({rel['toAttr']}),\n"

            # 移除最后的逗号
            sql = sql.rstrip(",\n") + "\n"
            sql += ");\n\n"

        return jsonify({'sql': sql})

    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/generate_doc', methods=['POST'])
def api_generate_doc():
    """
    解析SQL并生成指定格式的数据库结构文档
    """
    try:
        data = request.get_json()
        sql = data.get('sql', '')
        output_format = data.get('format', 'html')  # 'html' or 'docx'

        if not sql.strip():
            return jsonify({'error': 'SQL内容不能为空'}), 400

        # 1. 解析SQL获取表结构元数据
        tables_data, error = parse_sql(sql)
        if error:
            return jsonify({'error': f'SQL解析错误: {error}'}), 400

        if not tables_data:
            return jsonify({'error': '未找到有效的数据表定义'}), 400

        # 2. 根据请求的格式生成文档
        if output_format == 'html':
            html_content = generate_html(tables_data)
            return jsonify({'html': html_content})

        elif output_format == 'docx':
            import tempfile
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
                file_path = tmp_file.name

            generate_docx(tables_data, file_path)

            return send_file(
                file_path,
                as_attachment=True,
                download_name='数据库结构设计文档.docx',
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

        else:
            return jsonify({'error': '无效的格式，请使用 "html" 或 "docx"'}), 400

    except Exception as e:
        app.logger.error(f"生成文档时出错: {e}")
        return jsonify({'error': f'生成文档失败: {str(e)}'}), 500


    @app.route('/api/export_svg', methods=['POST'])
    def api_export_svg():
        """导出SVG格式的ER图"""
        try:
            data = request.get_json()
            svg_content = data.get('svg', '')
            
            # 创建SVG文件
            svg_file = io.BytesIO()
            svg_file.write(svg_content.encode('utf-8'))
            svg_file.seek(0)
            
            return send_file(svg_file, 
                            mimetype='image/svg+xml',
                            as_attachment=True,
                            download_name='er_diagram.svg')
            
        except Exception as e:
            return jsonify({'error': str(e)}), 500


    @app.route('/api/save_project', methods=['POST'])
    def api_save_project():
        """保存项目数据"""
        try:
            data = request.get_json()
            
            # 生成项目ID（如果没有）
            project_id = data.get('project_id') or str(uuid.uuid4())
            
            # 准备项目数据
            project_data = {
                'id': project_id,
                'name': data.get('name', f'项目_{datetime.now().strftime("%Y%m%d_%H%M%S")}'),
                'sql': data.get('sql', ''),
                'entities': data.get('entities', []),
                'relationships': data.get('relationships', []),
                'created_at': datetime.now().isoformat(),
                'updated_at': datetime.now().isoformat()
            }
            
            # 保存到内存存储（实际应用中应使用数据库）
            projects[project_id] = project_data
            
            # 保存到session
            if 'recent_projects' not in session:
                session['recent_projects'] = []
            
            # 更新最近项目列表
            recent = session['recent_projects']
            # 移除重复项
            recent = [p for p in recent if p['id'] != project_id]
            # 添加到开头
            recent.insert(0, {
                'id': project_id,
                'name': project_data['name'],
                'updated_at': project_data['updated_at']
            })
            # 只保留最近10个
            session['recent_projects'] = recent[:10]
            
            return jsonify({
                'success': True,
                'project_id': project_id,
                'message': '项目保存成功'
            })
            
        except Exception as e:
            app.logger.error(f"Error saving project: {e}")
            return jsonify({'error': str(e)}), 500


    @app.route('/api/load_project/<project_id>', methods=['GET'])
    def api_load_project(project_id):
        """加载项目数据"""
        try:
            # 从内存存储获取项目（实际应用中应从数据库获取）
            if project_id not in projects:
                return jsonify({'error': '项目不存在'}), 404
            
            project_data = projects[project_id]
            
            # 更新访问时间
            project_data['updated_at'] = datetime.now().isoformat()
            
            return jsonify({
                'success': True,
                'project': project_data
            })
            
        except Exception as e:
            app.logger.error(f"Error loading project: {e}")
            return jsonify({'error': str(e)}), 500


    @app.route('/api/list_projects', methods=['GET'])
    def api_list_projects():
        """列出所有项目"""
        try:
            # 获取所有项目列表
            project_list = []
            for pid, project in projects.items():
                project_list.append({
                    'id': pid,
                    'name': project['name'],
                    'created_at': project['created_at'],
                    'updated_at': project['updated_at']
                })
            
            # 按更新时间排序
            project_list.sort(key=lambda x: x['updated_at'], reverse=True)
            
            return jsonify({
                'success': True,
                'projects': project_list,
                'recent': session.get('recent_projects', [])
            })
            
        except Exception as e:
            app.logger.error(f"Error listing projects: {e}")
            return jsonify({'error': str(e)}), 500


@app.route('/api/delete_project/<project_id>', methods=['DELETE'])
def api_delete_project(project_id):
    """删除项目"""
    try:
        if project_id not in projects:
            return jsonify({'error': '项目不存在'}), 404

        # 删除项目
        del projects[project_id]

        # 从最近项目列表中移除
        if 'recent_projects' in session:
            session['recent_projects'] = [
                p for p in session['recent_projects']
                if p['id'] != project_id
            ]

        return jsonify({
            'success': True,
            'message': '项目删除成功'
        })

    except Exception as e:
        app.logger.error(f"Error deleting project: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/test-deepseek', methods=['GET'])
def api_test_deepseek():
    """测试DeepSeek API连接"""
    try:
        headers = {
            'Authorization': f'Bearer {DEEPSEEK_API_KEY}',
            'Content-Type': 'application/json'
        }

        test_payload = {
            'model': 'deepseek-chat',
            'messages': [
                {
                    'role': 'user',
                    'content': '请回复"API连接正常"'
                }
            ],
            'temperature': 0.1,
            'max_tokens': 50
        }

        response = requests.post(DEEPSEEK_API_URL, headers=headers, json=test_payload, timeout=30)

        if response.status_code == 200:
            result = response.json()
            content = result['choices'][0]['message']['content']
            return jsonify({
                'success': True,
                'message': 'DeepSeek API连接正常',
                'response': content,
                'api_key_prefix': DEEPSEEK_API_KEY[:10] + '...' if DEEPSEEK_API_KEY else 'None'
            })
        else:
            return jsonify({
                'success': False,
                'error': f'API调用失败: {response.status_code}',
                'response_text': response.text,
                'api_key_prefix': DEEPSEEK_API_KEY[:10] + '...' if DEEPSEEK_API_KEY else 'None'
            }), response.status_code

    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'API测试失败: {str(e)}',
            'api_key_prefix': DEEPSEEK_API_KEY[:10] + '...' if DEEPSEEK_API_KEY else 'None'
        }), 500





@app.route('/api/get_structure_cost', methods=['GET'])
def api_get_structure_cost():
    """获取AI生成结构图的费用信息"""
    try:
        cost = float(user_manager.get_system_config('paper_structure_cost', 4.00))
        
        if 'user_id' not in session:
            return jsonify({
                'is_logged_in': False,
                'cost': cost,
                'balance': 0,
                'sufficient': False
            })
        
        user_info = user_manager.get_user_info(session['user_id'])
        balance = float(user_info['balance']) if user_info else 0
        
        return jsonify({
            'is_logged_in': True,
            'cost': cost,
            'balance': balance,
            'sufficient': balance >= cost
        })
    except Exception as e:
        app.logger.error(f"获取结构图费用失败: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/generate-structure', methods=['POST'])
def api_generate_structure():
    """AI生成系统结构图API"""
    try:
        data = request.get_json()
        description = data.get('description', '')

        if not description:
            return jsonify({'success': False, 'message': '请提供系统描述'}), 400

        # 检查用户登录状态
        if 'user_id' not in session:
            return jsonify({
                'success': False,
                'message': 'AI生成功能需要登录后使用',
                'need_login': True
            }), 401
        
        # 获取费用并检查余额
        cost = float(user_manager.get_system_config('paper_structure_cost', 4.00))
        user_info = user_manager.get_user_info(session['user_id'])
        
        if not user_info or user_info['balance'] < cost:
            return jsonify({
                'success': False,
                'message': f'余额不足，AI生成需要 {cost:.2f} 元，当前余额 {user_info["balance"]:.2f} 元',
                'need_recharge': True,
                'cost': cost,
                'balance': float(user_info['balance']) if user_info else 0
            }), 400

        # 调用AI生成系统结构
        structure_data, is_ai_success = generate_system_structure_with_ai(description)

        # 只有AI真正成功才扣费
        charged = False
        if is_ai_success:
            consume_result = user_manager.consume_balance(
                session['user_id'],
                cost,
                'paper_structure',
                f'AI生成系统结构图'
            )
            charged = consume_result['success']
            if charged:
                app.logger.info(f"AI生成结构图成功，已扣费 {cost} 元")
            else:
                app.logger.warning(f"AI生成结构图扣费失败: {consume_result['message']}")
        else:
            app.logger.warning("AI调用失败，使用备用方案，不扣费")

        return jsonify({
            'success': True,
            'data': structure_data,
            'message': '系统结构图生成成功' + ('' if is_ai_success else '（AI服务暂时不可用，已使用备用方案，未扣费）'),
            'charged': charged,
            'cost': cost,
            'ai_success': is_ai_success
        })

    except Exception as e:
        app.logger.error(f"Error generating system structure: {e}")
        return jsonify({'success': False, 'message': f'生成失败: {str(e)}'}), 500


@app.route('/api/generate_simplified_er', methods=['POST'])
def api_generate_simplified_er():
    """生成简化ER图API"""
    try:
        data = request.get_json()
        sql = data.get('sql', '')
        options = data.get('options', {})
        
        if not sql.strip():
            return jsonify({'error': 'SQL内容不能为空'}), 400
        
        # 调用AI生成简化ER图数据
        simplified_data = generate_simplified_er_with_ai(sql, options)
        
        return jsonify(simplified_data)
        
    except Exception as e:
        app.logger.error(f"生成简化ER图失败: {e}")
        return jsonify({'error': f'生成失败: {str(e)}'}), 500


@app.route('/api/get_test_case_cost', methods=['GET'])
def api_get_test_case_cost():
    """获取AI测试用例生成的费用信息"""
    try:
        cost = float(user_manager.get_system_config('ai_test_case_cost', 3.00))
        
        if 'user_id' not in session:
            return jsonify({
                'is_logged_in': False,
                'cost': cost,
                'balance': 0,
                'sufficient': False
            })
        
        user_info = user_manager.get_user_info(session['user_id'])
        balance = float(user_info['balance']) if user_info else 0
        
        return jsonify({
            'is_logged_in': True,
            'cost': cost,
            'balance': balance,
            'sufficient': balance >= cost
        })
    except Exception as e:
        app.logger.error(f"获取测试用例费用失败: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/generate-test-cases', methods=['POST'])
def api_generate_test_cases():
    """生成测试用例API"""
    try:
        data = request.get_json()
        system_name = data.get('systemName', '')
        test_type = data.get('testType', '')
        system_description = data.get('systemDescription', '')

        if not all([system_name, test_type, system_description]):
            return jsonify({'error': '请提供完整的系统信息'}), 400

        # 检查用户登录状态
        if 'user_id' not in session:
            return jsonify({
                'error': 'AI生成功能需要登录后使用',
                'need_login': True
            }), 401
        
        # 获取费用并检查余额
        cost = float(user_manager.get_system_config('ai_test_case_cost', 3.00))
        user_info = user_manager.get_user_info(session['user_id'])
        
        if not user_info or user_info['balance'] < cost:
            return jsonify({
                'error': f'余额不足，AI生成需要 {cost:.2f} 元，当前余额 {user_info["balance"]:.2f} 元',
                'need_recharge': True,
                'cost': cost,
                'balance': float(user_info['balance']) if user_info else 0
            }), 400

        # 调用AI生成测试用例
        test_cases, is_ai_success = generate_test_cases_with_ai(system_name, test_type, system_description)

        # 只有AI真正成功才扣费
        charged = False
        if is_ai_success:
            consume_result = user_manager.consume_balance(
                session['user_id'],
                cost,
                'ai_test_case',
                f'AI生成测试用例 - {system_name}'
            )
            charged = consume_result['success']
            if charged:
                app.logger.info(f"AI生成测试用例成功，已扣费 {cost} 元")
            else:
                app.logger.warning(f"AI生成测试用例扣费失败: {consume_result['message']}")
        else:
            app.logger.warning(f"AI调用失败，使用备用方案，不扣费")

        return jsonify({
            'success': True,
            'testCases': test_cases,
            'message': f'成功生成 {len(test_cases)} 个测试用例' + ('' if is_ai_success else '（AI服务暂时不可用，已使用备用方案，未扣费）'),
            'charged': charged,
            'cost': cost,
            'ai_success': is_ai_success
        })

    except Exception as e:
        app.logger.error(f"Error generating test cases: {e}")
        return jsonify({'error': f'生成测试用例失败: {str(e)}'}), 500


@app.route('/api/export-test-cases-word', methods=['POST'])
def api_export_test_cases_word():
    """导出测试用例到Word文档"""
    try:
        data = request.get_json()
        test_cases = data.get('testCases', [])
        system_name = data.get('systemName', '系统')
        test_type = data.get('testType', '功能测试')

        if not test_cases:
            return jsonify({'error': '没有测试用例数据'}), 400

        # 生成Word文档
        doc_buffer = generate_test_cases_word(test_cases, system_name, test_type)

        # 返回文件
        return send_file(
            doc_buffer,
            as_attachment=True,
            download_name=f'测试用例_{system_name}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        app.logger.error(f"Error exporting test cases to Word: {e}")
        return jsonify({'error': f'导出Word文档失败: {str(e)}'}), 500


def generate_system_structure_with_ai(description):
    """使用AI生成系统功能结构图
    返回: (structure_data, is_ai_success) - 结构数据和是否AI成功的标志
    """
    try:
        # 构建提示词
        prompt = f"""
你是一个专业的系统架构师，请根据以下系统描述生成层次化的功能结构图数据：

系统描述：{description}

请分析系统需求，设计合理的功能模块层次结构。要求：

1. 生成一个根节点作为系统主体
2. 设计2-4个主要功能模块（一级子节点）
3. 每个主要模块下设计2-5个具体功能（二级子节点）
4. 根据需要可以有三级子节点
5. 节点名称要简洁明确，体现功能特点
6. 整体结构要逻辑清晰，层次分明

请严格按照以下JSON格式返回，不要包含任何其他文字：
{{
  "title": "系统名称",
  "nodes": [
    {{
      "id": 1,
      "text": "系统主体名称",
      "level": 0,
      "x": 0,
      "y": 0,
      "children": []
    }},
    {{
      "id": 2,
      "text": "主要功能模块1",
      "level": 1,
      "x": 0,
      "y": 0,
      "children": []
    }},
    {{
      "id": 3,
      "text": "具体功能1-1",
      "level": 2,
      "x": 0,
      "y": 0,
      "children": []
    }}
  ],
  "links": [
    {{
      "source": 1,
      "target": 2
    }},
    {{
      "source": 2,
      "target": 3
    }}
  ]
}}

注意：
- id必须是唯一的数字
- level表示层级：0=根节点，1=一级子节点，2=二级子节点
- links中的source和target必须对应nodes中的id
- 确保JSON格式完全正确
"""

        # 调用DeepSeek API
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {DEEPSEEK_API_KEY}"
        }

        data = {
            "model": "deepseek-chat",
            "messages": [
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            "temperature": 0.7,
            "max_tokens": 3000
        }

        response = requests.post(DEEPSEEK_API_URL, headers=headers, json=data, timeout=120)

        if response.status_code == 200:
            result = response.json()
            ai_response = result['choices'][0]['message']['content']

            # 尝试解析AI返回的JSON
            try:
                # 提取JSON部分
                start_idx = ai_response.find('{')
                end_idx = ai_response.rfind('}') + 1

                if start_idx != -1 and end_idx != 0:
                    json_str = ai_response[start_idx:end_idx]
                    structure_data = json.loads(json_str)

                    # 验证数据结构
                    if 'nodes' in structure_data and 'links' in structure_data:
                        return (structure_data, True)  # AI成功
                    else:
                        raise ValueError("AI返回的数据结构不完整")
                else:
                    raise ValueError("AI返回格式无效")

            except (json.JSONDecodeError, ValueError) as e:
                app.logger.error(f"AI返回数据解析失败: {e}")
                # 返回默认结构
                return (generate_default_structure(description), False)
        else:
            app.logger.error(f"DeepSeek API调用失败: {response.status_code}")
            return (generate_default_structure(description), False)

    except Exception as e:
        app.logger.error(f"AI生成系统结构失败: {e}")
        return (generate_default_structure(description), False)


def generate_default_structure(description):
    """生成默认的系统结构"""
    # 简单的关键词分析来生成默认结构
    system_name = "信息管理系统"
    if "教育" in description or "学生" in description or "课程" in description:
        system_name = "教育管理系统"
    elif "医疗" in description or "医院" in description or "病人" in description:
        system_name = "医疗管理系统"
    elif "电商" in description or "购物" in description or "商品" in description:
        system_name = "电商管理系统"
    elif "图书" in description or "借阅" in description:
        system_name = "图书管理系统"

    return {
        "title": system_name,
        "nodes": [
            {"id": 1, "text": system_name, "level": 0, "x": 0, "y": 0, "children": []},
            {"id": 2, "text": "用户管理", "level": 1, "x": 0, "y": 0, "children": []},
            {"id": 3, "text": "数据管理", "level": 1, "x": 0, "y": 0, "children": []},
            {"id": 4, "text": "系统管理", "level": 1, "x": 0, "y": 0, "children": []},
            {"id": 5, "text": "用户注册", "level": 2, "x": 0, "y": 0, "children": []},
            {"id": 6, "text": "用户登录", "level": 2, "x": 0, "y": 0, "children": []},
            {"id": 7, "text": "数据录入", "level": 2, "x": 0, "y": 0, "children": []},
            {"id": 8, "text": "数据查询", "level": 2, "x": 0, "y": 0, "children": []},
            {"id": 9, "text": "权限管理", "level": 2, "x": 0, "y": 0, "children": []},
            {"id": 10, "text": "系统配置", "level": 2, "x": 0, "y": 0, "children": []}
        ],
        "links": [
            {"source": 1, "target": 2},
            {"source": 1, "target": 3},
            {"source": 1, "target": 4},
            {"source": 2, "target": 5},
            {"source": 2, "target": 6},
            {"source": 3, "target": 7},
            {"source": 3, "target": 8},
            {"source": 4, "target": 9},
            {"source": 4, "target": 10}
        ]
    }


def identify_system_type(sql):
    """识别SQL系统类型"""
    sql_lower = sql.lower()
    
    # 校友/社区系统特征
    if any(keyword in sql_lower for keyword in ['alumni', 'xiaoyou', 'posts', 'comments', 'forum']):
        if 'alumni' in sql_lower or 'xiaoyou' in sql_lower:
            return "校友社区系统"
        else:
            return "社区论坛系统"
    
    # 电商系统特征
    elif any(keyword in sql_lower for keyword in ['order', 'product', 'cart', 'payment', 'shop']):
        return "电子商务系统"
    
    # 教育系统特征
    elif any(keyword in sql_lower for keyword in ['student', 'course', 'teacher', 'class', 'education']):
        return "教育管理系统"
    
    # 医疗系统特征
    elif any(keyword in sql_lower for keyword in ['patient', 'doctor', 'hospital', 'medical', 'treatment']):
        return "医疗管理系统"
    
    # 企业管理系统特征
    elif any(keyword in sql_lower for keyword in ['employee', 'department', 'project', 'company', 'hr']):
        return "企业管理系统"
    
    # 内容管理系统特征
    elif any(keyword in sql_lower for keyword in ['article', 'content', 'category', 'tag', 'publish']):
        return "内容管理系统"
    
    # 默认
    else:
        return "信息管理系统"


def get_system_specific_guidance(system_type):
    """根据系统类型提供简化指导"""
    guidance_map = {
        "校友社区系统": """
核心实体：用户（users）、帖子（posts）、评论（comments）、新闻（news）、活动（activities）
功能模块分组建议：
- 用户管理：users（用户）
- 内容管理：posts（论坛帖子）、comments（帖子评论）、news（新闻资讯）
- 活动管理：activities（校友活动）、alumni_associations（校友会）
- 招聘信息：jobs（招聘信息）
- 辅助功能：donations（捐赠记录）、favorites（用户收藏）、carousel（轮播图）等可合并为"系统功能"
关系重点：用户发表帖子、用户评论帖子、用户参与活动、用户浏览新闻等核心业务关系
""",
        "社区论坛系统": """
核心实体：用户（users）、帖子（posts）、评论（comments）、分类（categories）
功能模块：用户管理、内容发布、互动交流、系统管理
关系重点：用户-帖子-评论的核心交互链
""",
        "电子商务系统": """
核心实体：用户（users）、商品（products）、订单（orders）、购物车（cart）
功能模块：用户管理、商品管理、订单管理、支付系统
关系重点：用户下单、订单包含商品等交易关系
""",
        "企业管理系统": """
核心实体：员工（employees）、部门（departments）、项目（projects）
功能模块：人员管理、部门管理、项目管理、权限管理
关系重点：员工属于部门、员工参与项目等组织关系
""",
        "信息管理系统": """
核心实体：用户（users）、数据（data）、管理（management）
功能模块：用户管理、数据管理、系统管理
关系重点：用户操作数据的基本关系
"""
    }
    return guidance_map.get(system_type, guidance_map["信息管理系统"])


def generate_simplified_er_with_ai(sql, options):
    """使用AI生成简化的ER图数据"""
    try:
        # 分析SQL内容，识别系统类型
        system_type = identify_system_type(sql)
        
        # 构建针对简化ER图的提示词
        prompt = f"""
请分析以下SQL建表语句，这是一个{system_type}系统，生成适合{options.get('usagePurpose', 'thesis')}的简化ER图数据。

SQL语句：
```sql
{sql}
```

系统分析：
- 系统类型：{system_type}
- 简化要求：{'只显示核心实体（用户、管理员等主要角色）' if options.get('showMainEntities') else '显示所有实体'}
- 模块分组：{'将功能相关的表归类为功能模块' if options.get('groupFunctions') else '保持独立实体'}
- 属性显示：{'隐藏详细属性，只显示实体间关系' if options.get('hideAttributes') else '显示关键属性（主键、外键）'}
- 风格优化：{'优化为论文风格，高层次抽象' if options.get('paperStyle') else '保持技术细节'}

针对{system_type}系统的简化指导：
{get_system_specific_guidance(system_type)}

请返回JSON格式的简化ER图数据，包含：
1. entities: 简化后的实体列表，每个实体包含 name, displayName（中文名称）
2. relationships: 实体间的关系，包含 from, to, name, displayName, type

返回格式示例：
{{
  "entities": [
    {{
      "name": "users",
      "displayName": "用户",
      "attributes": [
        {{"name": "id", "displayName": "用户ID", "isPK": true, "type": "BIGINT"}},
        {{"name": "username", "displayName": "用户名", "isPK": false, "type": "VARCHAR"}}
      ]
    }}
  ],
  "relationships": [
    {{
      "from": "users",
      "to": "posts",
      "name": "user_posts",
      "displayName": "发表",
      "type": "1:N"
    }}
  ]
}}

重要：请根据系统类型合理简化，突出核心业务实体和关系，使用准确的中文名称。
请只返回JSON数据，不要包含其他说明文字。
"""

        # 调用DeepSeek API
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {DEEPSEEK_API_KEY}"
        }

        data = {
            "model": "deepseek-chat",
            "messages": [
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            "temperature": 0.3,  # 较低的温度获得更一致的结果
            "max_tokens": 4000
        }

        # 重试机制：最多尝试3次
        max_retries = 3
        for attempt in range(max_retries):
            try:
                app.logger.info(f"开始调用DeepSeek API生成简化ER图... (尝试 {attempt + 1}/{max_retries})")
                response = requests.post(DEEPSEEK_API_URL, headers=headers, json=data, timeout=120)

                if response.status_code == 200:
                    result = response.json()
                    ai_response = result['choices'][0]['message']['content']
                    app.logger.info(f"AI响应长度: {len(ai_response)}")

                    # 尝试解析AI返回的JSON
                    try:
                        # 提取JSON部分
                        start_idx = ai_response.find('{')
                        end_idx = ai_response.rfind('}') + 1

                        if start_idx != -1 and end_idx != 0:
                            json_str = ai_response[start_idx:end_idx]
                            simplified_data = json.loads(json_str)

                            # 验证数据结构
                            if 'entities' in simplified_data and 'relationships' in simplified_data:
                                app.logger.info("AI简化ER图生成成功")
                                simplified_data['ai_generated'] = True
                                return simplified_data
                            else:
                                raise ValueError("AI返回的数据结构不完整")
                        else:
                            raise ValueError("AI返回格式无效")

                    except (json.JSONDecodeError, ValueError) as e:
                        app.logger.error(f"AI返回数据解析失败: {e}")
                        app.logger.error(f"原始响应: {ai_response[:500]}...")
                        # 如果是最后一次尝试，返回默认结果
                        if attempt == max_retries - 1:
                            return generate_default_simplified_er(sql, options)
                        # 否则继续重试
                        continue
                else:
                    app.logger.error(f"DeepSeek API调用失败: {response.status_code}")
                    if attempt == max_retries - 1:
                        return generate_default_simplified_er(sql, options)
                    continue
                    
            except requests.exceptions.Timeout:
                app.logger.error(f"DeepSeek API调用超时 (尝试 {attempt + 1}/{max_retries})")
                if attempt == max_retries - 1:
                    app.logger.error("所有重试都超时，返回默认简化结果")
                    return generate_default_simplified_er(sql, options)
                # 等待后重试
                import time
                time.sleep(2)
                continue
            except Exception as e:
                app.logger.error(f"DeepSeek API调用异常 (尝试 {attempt + 1}/{max_retries}): {e}")
                if attempt == max_retries - 1:
                    return generate_default_simplified_er(sql, options)
                continue
        
        # 如果所有重试都失败，返回默认结果
        return generate_default_simplified_er(sql, options)

    except Exception as e:
        app.logger.error(f"AI生成简化ER图失败: {e}")
        return generate_default_simplified_er(sql, options)


def generate_default_simplified_er(sql, options):
    """基于SQL解析生成默认的简化ER图"""
    try:
        # 使用现有的SQL解析功能
        parsed_result, error_message = parse_sql(sql)
        
        if error_message or not parsed_result:
            # 如果SQL解析失败，返回最基本的结构
            return {
                "entities": [
                    {
                        "name": "user",
                        "displayName": "用户",
                        "attributes": [
                            {"name": "id", "displayName": "ID", "isPK": True}
                        ] if not options.get('hideAttributes') else []
                    },
                    {
                        "name": "admin",
                        "displayName": "管理员",
                        "attributes": [
                            {"name": "id", "displayName": "ID", "isPK": True}
                        ] if not options.get('hideAttributes') else []
                    }
                ],
                "relationships": [
                    {
                        "from": "user",
                        "to": "admin",
                        "name": "manages",
                        "displayName": "管理",
                        "type": "1:N"
                    }
                ],
                "ai_generated": False
            }

        # 简化表结构
        simplified_entities = []
        simplified_relationships = []
        
        # 识别核心表和功能表
        core_tables = []
        function_tables = []
        
        for table_name, table_data in parsed_result.items():
            table_lower = table_name.lower()
            
            # 判断是否为核心表（用户、内容、交互相关）
            if any(keyword in table_lower for keyword in [
                'user', 'admin', 'member', 'account', 'person',  # 用户相关
                'post', 'comment', 'news', 'article', 'content',  # 内容相关
                'activity', 'event', 'job', 'alumni'  # 核心业务相关
            ]):
                core_tables.append(table_name)
            else:
                function_tables.append(table_name)
        
        # 处理核心表
        for table_name in core_tables:
            table_data = parsed_result[table_name]
            entity = {
                "name": table_name,
                "displayName": get_chinese_name(table_name),
            }
            
            # 如果不隐藏属性，添加关键属性
            if not options.get('hideAttributes'):
                key_attrs = []
                for col in table_data["columns"]:
                    if col.get("pk") or col["name"] in table_data.get("primary_keys", []):
                        key_attrs.append({
                            "name": col["name"],
                            "displayName": col.get("comment") or get_chinese_name(col["name"]),
                            "isPK": True,
                            "type": col.get("type", "")
                        })
                entity["attributes"] = key_attrs
            else:
                entity["attributes"] = []
                
            simplified_entities.append(entity)
        
        # 如果启用了功能分组，将功能表作为一个整体
        if options.get('groupFunctions') and function_tables:
            entity = {
                "name": "function_modules",
                "displayName": "功能模块",
                "attributes": []
            }
            simplified_entities.append(entity)
        else:
            # 选择部分重要的功能表
            important_function_tables = function_tables[:3]  # 只取前3个
            for table_name in important_function_tables:
                entity = {
                    "name": table_name,
                    "displayName": get_chinese_name(table_name),
                    "attributes": []
                }
                simplified_entities.append(entity)
        
        # 生成简化的关系
        entities, relationships = build_er_model(parsed_result)
        
        # 只保留核心实体间的关系
        for rel in relationships:
            if rel.from_entity in core_tables and rel.to_entity in core_tables:
                simplified_relationships.append({
                    "from": rel.from_entity,
                    "to": rel.to_entity,
                    "name": rel.name,
                    "displayName": get_chinese_name(rel.name),
                    "type": rel.rel_type
                })
        
        return {
            "entities": simplified_entities,
            "relationships": simplified_relationships,
            "ai_generated": False
        }
        
    except Exception as e:
        app.logger.error(f"生成默认简化ER图失败: {e}")
        # 最后的fallback
        return {
            "entities": [
                {"name": "user", "displayName": "用户", "attributes": []},
                {"name": "admin", "displayName": "管理员", "attributes": []}
            ],
            "relationships": [
                {"from": "user", "to": "admin", "name": "manages", "displayName": "管理", "type": "1:N"}
            ],
            "ai_generated": False
        }


def get_chinese_name(english_name):
    """将英文名称转换为中文（简单映射）"""
    name_mapping = {
        # 通用实体
        'user': '用户', 'users': '用户',
        'admin': '管理员', 'administrator': '管理员',
        'member': '会员', 'members': '会员',
        
        # 校友系统专用
        'posts': '帖子', 'post': '帖子',
        'comments': '评论', 'comment': '评论',
        'news': '新闻',
        'activities': '活动', 'activity': '活动',
        'jobs': '招聘信息', 'job': '工作',
        'alumni_associations': '校友会', 'alumni_association': '校友会',
        'donations': '捐赠记录', 'donation': '捐赠',
        'favorites': '收藏', 'favorite': '收藏',
        'carousel': '轮播图',
        'post_likes': '帖子点赞', 'like': '点赞',
        
        # 电商系统
        'product': '商品', 'products': '商品',
        'order': '订单', 'orders': '订单',
        'cart': '购物车', 'shopping_cart': '购物车',
        
        # 内容管理
        'category': '分类', 'categories': '分类',
        'article': '文章', 'articles': '文章',
        'tag': '标签', 'tags': '标签',
        
        # 系统功能
        'log': '日志', 'logs': '日志',
        'config': '配置', 'configuration': '配置',
        'setting': '设置', 'settings': '设置',
        'permission': '权限', 'permissions': '权限',
        'role': '角色', 'roles': '角色',
        
        # 企业管理
        'department': '部门', 'departments': '部门',
        'employee': '员工', 'employees': '员工',
        'project': '项目', 'projects': '项目',
        'task': '任务', 'tasks': '任务',
        
        # 媒体文件
        'file': '文件', 'files': '文件',
        'image': '图片', 'images': '图片',
        'video': '视频', 'videos': '视频',
        'upload': '上传', 'uploads': '上传文件',
        
        # 通讯
        'message': '消息', 'messages': '消息',
        'notification': '通知', 'notifications': '通知',
        
        # 关系名称翻译
        'manages': '管理',
        'belongs_to': '属于',
        'has_many': '拥有',
        'creates': '创建',
        'likes': '点赞',
        'comments_on': '评论',
        'participates': '参与',
        'publishes': '发布',
        'reads': '阅读',
        'favorites': '收藏',
    }
    
    name_lower = english_name.lower().strip('_').strip()
    return name_mapping.get(name_lower, english_name.replace('_', ' ').title())


def generate_test_cases_with_ai(system_name, test_type, system_description):
    """使用DeepSeek AI生成测试用例
    返回: (test_cases, is_ai_success) - 测试用例列表和是否AI成功的标志
    """
    try:
        # 构建AI提示词
        prompt = f"""
你是一个专业的软件测试工程师，请仔细分析以下系统信息，生成针对性的测试用例。

系统名称：{system_name}
测试类型：{test_type}
系统功能描述：
{system_description}

重要要求：
1. 请仔细阅读系统功能描述，理解系统的具体业务场景和功能特点
2. 根据系统的实际功能生成测试用例，不要使用通用模板
3. 测试用例要覆盖系统描述中提到的具体功能点
4. 测试步骤要具体可执行，符合实际操作流程
5. 实际结果要模拟真实的测试执行结果，描述功能是否正常工作

请生成5-8个测试用例，每个测试用例必须包含以下字段：
- caseId: 测试用例编号（格式：TC-001, TC-002...）
- module: 测试模块（根据系统功能确定）
- function: 测试功能点（具体的功能名称）
- precondition: 前置条件（测试前需要满足的条件）
- steps: 测试步骤（详细的操作步骤，用\\n分隔）
- expectedResult: 预期结果（期望的结果，用\\n分隔）
- actualResult: 实际结果（模拟真实测试结果，描述功能执行情况和结果）
- remark: 备注信息（可以为空或填写注意事项）

请以JSON数组格式返回，确保JSON格式正确。示例格式：
[
  {{
    "caseId": "TC-001",
    "module": "具体模块名",
    "function": "具体功能名",
    "precondition": "1. 具体前置条件1\\n2. 具体前置条件2",
    "steps": "1. 具体操作步骤1\\n2. 具体操作步骤2\\n3. 具体操作步骤3",
    "expectedResult": "1. 具体预期结果1\\n2. 具体预期结果2",
    "actualResult": "功能正常执行，结果符合预期，具体描述测试结果",
    "remark": "针对性备注"
  }}
]

请确保生成的测试用例与系统描述高度相关，体现系统的具体功能特点。
"""

        # 调用DeepSeek API
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {DEEPSEEK_API_KEY}"
        }

        data = {
            "model": "deepseek-chat",
            "messages": [
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            "temperature": 0.7,
            "max_tokens": 4000
        }

        response = requests.post(DEEPSEEK_API_URL, headers=headers, json=data, timeout=120)

        if response.status_code == 200:
            result = response.json()
            ai_response = result['choices'][0]['message']['content']

            # 尝试解析AI返回的JSON
            try:
                # 提取JSON部分（可能包含其他文本）
                start_idx = ai_response.find('[')
                end_idx = ai_response.rfind(']') + 1

                if start_idx != -1 and end_idx != 0:
                    json_str = ai_response[start_idx:end_idx]
                    test_cases = json.loads(json_str)

                    # 验证数据格式
                    for case in test_cases:
                        required_fields = ['caseId', 'module', 'function', 'precondition', 'steps', 'expectedResult', 'actualResult', 'remark']
                        for field in required_fields:
                            if field not in case:
                                case[field] = ""

                    return (test_cases, True)  # AI成功
                else:
                    raise ValueError("无法找到有效的JSON格式")

            except json.JSONDecodeError as e:
                app.logger.error(f"AI返回的JSON格式错误: {e}")
                app.logger.error(f"AI响应内容: {ai_response}")
                # 如果AI返回格式有问题，使用备用方案
                return (generate_fallback_test_cases(system_name, test_type, system_description), False)
        else:
            app.logger.error(f"DeepSeek API调用失败: {response.status_code} - {response.text}")
            return (generate_fallback_test_cases(system_name, test_type, system_description), False)

    except Exception as e:
        app.logger.error(f"AI生成测试用例时出错: {e}")
        # 出错时使用备用方案
        return (generate_fallback_test_cases(system_name, test_type, system_description), False)


def generate_fallback_test_cases(system_name, test_type, system_description):
    """备用测试用例生成方案（当AI API不可用时）"""
    if test_type == "功能测试":
        return generate_functional_test_cases(system_name, system_description)
    elif test_type == "接口测试":
        return generate_api_test_cases(system_name, system_description)
    elif test_type == "性能测试":
        return generate_performance_test_cases(system_name, system_description)
    elif test_type == "安全测试":
        return generate_security_test_cases(system_name, system_description)
    else:
        return generate_comprehensive_test_cases(system_name, system_description)


def generate_functional_test_cases(system_name, description):
    """生成功能测试用例"""
    return [
        {
            "caseId": "TC-001",
            "module": f"{system_name}用户登录",
            "function": "用户名密码验证",
            "precondition": "1. 系统已启动\n2. 数据库连接正常",
            "steps": "1. 打开登录页面\n2. 输入用户名\n3. 输入密码\n4. 点击'登录'按钮",
            "expectedResult": "1. 系统跳转到主页\n2. 页面显示用户信息",
            "actualResult": "登录成功，系统正确跳转",
            "remark": "基础功能测试"
        }
    ]


def generate_api_test_cases(system_name, description):
    """生成接口测试用例"""
    # 简化版本，具体可以根据需要扩展
    return [
        {
            "caseId": "API-001",
            "module": "用户接口",
            "function": "GET /api/users/{id}",
            "precondition": "1. 认证Token有效\n2. 用户ID存在",
            "steps": "发送请求：\nGET /api/users/123\nHeader: Authorization: Bearer <token>",
            "expectedResult": "HTTP 200 OK\nBody: {id:123, name:'Alice'}",
            "actualResult": "HTTP 200 OK，返回用户信息正确，响应时间150ms",
            "remark": "Token权限配置错误"
        }
    ]

def generate_performance_test_cases(system_name, description):
    """生成性能测试用例"""
    # 简化版本
    return [
        {
            "caseId": "PERF-001",
            "module": "系统性能",
            "function": "并发用户登录",
            "precondition": "1. 系统正常运行\n2. 测试数据准备完成",
            "steps": "JMeter脚本：\n- 100个并发用户\n- 持续时间：5分钟\n- 登录接口压测",
            "expectedResult": "1. 响应时间 < 2秒\n2. 成功率 > 99%\n3. 系统稳定运行",
            "actualResult": "平均响应时间1.2秒，成功率99.8%，系统运行稳定，CPU使用率65%",
            "remark": "监控CPU、内存使用率"
        }
    ]

def generate_security_test_cases(system_name, description):
    """生成安全测试用例"""
    # 简化版本
    return [
        {
            "caseId": "SEC-001",
            "module": "身份认证",
            "function": "SQL注入防护",
            "precondition": "1. 系统正常运行",
            "steps": "1. 在登录框输入SQL注入代码\n2. 尝试绕过认证\n3. 检查系统响应",
            "expectedResult": "1. 系统拒绝恶意输入\n2. 记录安全日志\n3. 不泄露敏感信息",
            "actualResult": "系统成功拦截SQL注入攻击，参数化查询生效，安全日志已记录",
            "remark": "使用OWASP测试用例"
        }
    ]

def generate_comprehensive_test_cases(system_name, description):
    """生成综合测试用例"""
    functional_cases = generate_functional_test_cases(system_name, description)
    api_cases = generate_api_test_cases(system_name, description)

    # 合并不同类型的测试用例
    all_cases = functional_cases + api_cases

    # 重新编号
    for i, case in enumerate(all_cases, 1):
        case['caseId'] = f"TC-{i:03d}"

    return all_cases


def _set_triple_line_style(table):
    """
    应用标准三线表样式 - 只有三条线
    - 表格顶部粗线
    - 标题行下方细线
    - 表格底部粗线
    - 无任何垂直线或其他横线
    """
    # 清除表格所有默认边框
    tbl = table._tbl

    # 移除表格样式，避免样式冲突
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl.insert(0, tbl_pr)

    # 清除所有现有边框
    old_borders = tbl_pr.find(qn('w:tblBorders'))
    if old_borders is not None:
        tbl_pr.remove(old_borders)

    # 设置表格边框为无
    tbl_borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tbl_borders.append(border)
    tbl_pr.append(tbl_borders)

    # 清除所有单元格边框
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            # 移除旧的边框
            old_borders = tc_pr.find(qn('w:tcBorders'))
            if old_borders is not None:
                tc_pr.remove(old_borders)

            # 设置无边框
            tc_borders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'nil')
                tc_borders.append(border)
            tc_pr.append(tc_borders)

    # 只添加三条线
    # 1. 顶线（第一行的顶部）- 粗线
    for cell in table.rows[0].cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = tc_pr.find(qn('w:tcBorders'))
        if tc_borders is None:
            tc_borders = OxmlElement('w:tcBorders')
            tc_pr.append(tc_borders)

        # 添加顶部边框
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), '12')  # 1.5pt 粗线
        top.set(qn('w:space'), '0')
        top.set(qn('w:color'), '000000')
        tc_borders.append(top)

    # 2. 栏目线（第一行的底部）- 细线
    for cell in table.rows[0].cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = tc_pr.find(qn('w:tcBorders'))
        if tc_borders is None:
            tc_borders = OxmlElement('w:tcBorders')
            tc_pr.append(tc_borders)

        # 添加底部边框
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')  # 0.75pt 细线
        bottom.set(qn('w:space'), '0')
        bottom.set(qn('w:color'), '000000')
        tc_borders.append(bottom)

    # 3. 底线（最后一行的底部）- 粗线
    for cell in table.rows[-1].cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = tc_pr.find(qn('w:tcBorders'))
        if tc_borders is None:
            tc_borders = OxmlElement('w:tcBorders')
            tc_pr.append(tc_borders)

        # 添加底部边框
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')  # 1.5pt 粗线
        bottom.set(qn('w:space'), '0')
        bottom.set(qn('w:color'), '000000')
        tc_borders.append(bottom)


def generate_test_cases_word(test_cases, system_name, test_type):
    """生成测试用例Word文档"""
    doc = Document()

    # 设置文档标题
    title = doc.add_heading(f'{system_name} - {test_type}用例', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加基本信息
    info_para = doc.add_paragraph()
    info_para.add_run('项目名称：').bold = True
    info_para.add_run(system_name)
    info_para.add_run('\n测试类型：').bold = True
    info_para.add_run(test_type)
    info_para.add_run('\n生成时间：').bold = True
    info_para.add_run(datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
    info_para.add_run('\n测试用例数量：').bold = True
    info_para.add_run(str(len(test_cases)))

    # 添加空行
    doc.add_paragraph()

    # 创建表格（三线表格式）
    table = doc.add_table(rows=1, cols=8)
    # 不使用任何预定义样式，后面会应用三线表样式
    table.style = None
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 设置表头
    header_cells = table.rows[0].cells
    headers = ['用例编号', '测试模块', '测试功能点', '前置条件', '测试步骤', '预期结果', '实际结果', '备注']

    for i, header in enumerate(headers):
        header_cells[i].text = header
        # 设置表头样式
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加测试用例数据
    for test_case in test_cases:
        row_cells = table.add_row().cells

        # 填充数据
        row_cells[0].text = test_case.get('caseId', '')
        row_cells[1].text = test_case.get('module', '')
        row_cells[2].text = test_case.get('function', '')
        row_cells[3].text = test_case.get('precondition', '')
        row_cells[4].text = test_case.get('steps', '')
        row_cells[5].text = test_case.get('expectedResult', '')
        row_cells[6].text = test_case.get('actualResult', '')
        row_cells[7].text = test_case.get('remark', '')

        # 设置单元格样式
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            # 设置单元格内边距
            cell.vertical_alignment = 1  # 垂直居中

    # 设置表格列宽
    for i, width in enumerate([0.8, 1.0, 1.2, 1.5, 2.0, 1.5, 1.0, 1.2]):
        for row in table.rows:
            row.cells[i].width = Inches(width)

    # 应用三线表样式
    _set_triple_line_style(table)

    # 保存到内存
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)

    return doc_buffer


# ============== AI检测率预估功能 ==============

def calculate_ai_detection_score(text):
    """
    计算文本的AI检测率预估分数
    基于Perplexity（困惑度）和Burstiness（突发性）原理

    返回：
    - ai_score: AI检测率预估 (0-100)
    - details: 详细分析结果
    """
    import math
    import statistics

    if not text or len(text) < 50:
        return {
            'ai_score': 0,
            'confidence': 'low',
            'details': {
                'text_length': len(text) if text else 0,
                'error': '文本太短，无法进行准确检测（至少需要50字符）'
            }
        }

    # 1. 句子分割
    sentences = re.split(r'[。！？；\n]+', text)
    sentences = [s.strip() for s in sentences if s.strip() and len(s.strip()) > 2]

    if len(sentences) < 3:
        return {
            'ai_score': 0,
            'confidence': 'low',
            'details': {
                'sentence_count': len(sentences),
                'error': '句子数量太少，无法进行准确检测（至少需要3个句子）'
            }
        }

    # 2. 计算Burstiness（句子长度突发性）
    sentence_lengths = [len(s) for s in sentences]
    avg_length = statistics.mean(sentence_lengths)
    std_length = statistics.stdev(sentence_lengths) if len(sentence_lengths) > 1 else 0

    # 突发性分数：标准差越大，越像人类写作
    # AI文本通常std/avg < 0.3，人类文本通常 > 0.5
    burstiness_ratio = std_length / avg_length if avg_length > 0 else 0

    # 3. 计算句子长度分布的均匀性
    # 人类写作：长短句交替，有变化
    # AI写作：句子长度趋于均匀
    length_variance = statistics.variance(sentence_lengths) if len(sentence_lengths) > 1 else 0

    # 4. 检测AI特征词频率
    ai_markers = [
        # 结构性标记词
        '首先', '其次', '再次', '最后', '另外',
        '此外', '与此同时', '综上所述', '总而言之',
        '值得注意的是', '具体而言', '由此可见',
        # AI高频表达
        '极大地', '具有重要意义', '发挥着重要作用',
        '得到了广泛应用', '取得了显著成效',
        '不难发现', '显而易见', '毋庸置疑',
        # 过渡词
        '因此', '所以', '故而', '从而', '进而',
        '一方面', '另一方面', '在此基础上',
        # AI特有句式
        '是...的', '通过...可以', '对于...来说',
    ]

    marker_count = 0
    for marker in ai_markers:
        marker_count += text.count(marker)

    # AI标记词密度（每100字符的标记词数量）
    marker_density = (marker_count / len(text)) * 100

    # 5. 检测"完美"句式比例
    # AI倾向于写出结构完整、过于规范的句子
    perfect_pattern_count = 0
    perfect_patterns = [
        r'首先.*其次.*最后',
        r'一方面.*另一方面',
        r'不仅.*而且',
        r'虽然.*但是',
    ]
    for pattern in perfect_patterns:
        if re.search(pattern, text, re.DOTALL):
            perfect_pattern_count += 1

    # 6. 检测段落结构一致性
    paragraphs = text.split('\n')
    paragraphs = [p.strip() for p in paragraphs if p.strip()]
    para_lengths = [len(p) for p in paragraphs] if len(paragraphs) > 1 else [len(text)]
    para_variance = statistics.variance(para_lengths) if len(para_lengths) > 1 else 0

    # 7. 计算综合AI分数
    # 各指标权重
    score_components = {
        'burstiness': 0,  # 突发性分数 (0-40)
        'markers': 0,     # AI标记词分数 (0-30)
        'patterns': 0,    # 完美句式分数 (0-20)
        'uniformity': 0   # 均匀性分数 (0-10)
    }

    # Burstiness分数计算（权重40%）
    # burstiness_ratio < 0.25 -> 高AI可能性
    # burstiness_ratio > 0.6 -> 低AI可能性
    if burstiness_ratio < 0.2:
        score_components['burstiness'] = 40
    elif burstiness_ratio < 0.3:
        score_components['burstiness'] = 32
    elif burstiness_ratio < 0.4:
        score_components['burstiness'] = 24
    elif burstiness_ratio < 0.5:
        score_components['burstiness'] = 16
    elif burstiness_ratio < 0.6:
        score_components['burstiness'] = 8
    else:
        score_components['burstiness'] = 0

    # AI标记词分数（权重30%）
    # marker_density > 1.5 -> 高AI可能性
    if marker_density > 2.0:
        score_components['markers'] = 30
    elif marker_density > 1.5:
        score_components['markers'] = 24
    elif marker_density > 1.0:
        score_components['markers'] = 18
    elif marker_density > 0.5:
        score_components['markers'] = 12
    elif marker_density > 0.2:
        score_components['markers'] = 6
    else:
        score_components['markers'] = 0

    # 完美句式分数（权重20%）
    score_components['patterns'] = min(perfect_pattern_count * 7, 20)

    # 均匀性分数（权重10%）
    # 句子长度方差太小 -> 高AI可能性
    normalized_variance = length_variance / (avg_length ** 2) if avg_length > 0 else 0
    if normalized_variance < 0.1:
        score_components['uniformity'] = 10
    elif normalized_variance < 0.2:
        score_components['uniformity'] = 7
    elif normalized_variance < 0.3:
        score_components['uniformity'] = 4
    else:
        score_components['uniformity'] = 0

    # 总分
    ai_score = sum(score_components.values())

    # 确定置信度
    if len(text) > 1000 and len(sentences) > 10:
        confidence = 'high'
    elif len(text) > 500 and len(sentences) > 5:
        confidence = 'medium'
    else:
        confidence = 'low'

    # 生成详细分析
    details = {
        'text_length': len(text),
        'sentence_count': len(sentences),
        'avg_sentence_length': round(avg_length, 1),
        'sentence_length_std': round(std_length, 1),
        'burstiness_ratio': round(burstiness_ratio, 3),
        'ai_marker_count': marker_count,
        'marker_density': round(marker_density, 3),
        'perfect_patterns': perfect_pattern_count,
        'score_breakdown': score_components,
        'analysis': generate_analysis_text(ai_score, burstiness_ratio, marker_density, perfect_pattern_count)
    }

    return {
        'ai_score': min(ai_score, 100),
        'confidence': confidence,
        'details': details
    }


def generate_analysis_text(ai_score, burstiness_ratio, marker_density, pattern_count):
    """生成分析文本"""
    analysis = []

    # AI分数判断
    if ai_score >= 70:
        analysis.append("该文本具有较高的AI生成特征，建议进行深度优化。")
    elif ai_score >= 50:
        analysis.append("该文本存在一定的AI生成特征，建议进行适度优化。")
    elif ai_score >= 30:
        analysis.append("该文本AI特征较少，但仍有优化空间。")
    else:
        analysis.append("该文本AI特征不明显，接近人类写作风格。")

    # 具体问题
    if burstiness_ratio < 0.3:
        analysis.append("• 句式突发性低：句子长度过于均匀，建议增加长短句交替。")
    if marker_density > 1.0:
        analysis.append("• AI标记词过多：存在较多'首先、其次、综上所述'等AI高频词，建议替换或删除。")
    if pattern_count > 0:
        analysis.append("• 结构过于规整：存在'首先...其次...最后'等AI典型结构，建议打散。")

    return '\n'.join(analysis)


@app.route('/api/estimate-ai-score', methods=['POST'])
def api_estimate_ai_score():
    """AI检测率预估API - 免费功能，无需登录（用于文本优化器）"""
    try:
        data = request.json
        text = data.get('text', '').strip()

        if not text:
            return jsonify({'error': '请提供要检测的文本'}), 400

        if len(text) > 100000:
            return jsonify({'error': '文本长度超过限制（最大10万字符）'}), 400

        # 计算AI检测分数
        result = calculate_ai_detection_score(text)

        return jsonify({
            'success': True,
            'ai_score': result['ai_score'],
            'confidence': result['confidence'],
            'details': result['details'],
            'recommendation': get_recommendation(result['ai_score'])
        })

    except Exception as e:
        print(f"AI检测API错误: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'检测失败: {str(e)}'}), 500


def get_recommendation(ai_score):
    """根据AI分数给出建议"""
    if ai_score >= 70:
        return {
            'level': 'high',
            'text': '高风险',
            'suggestion': '建议使用"深度优化"模式，强度设为4-5级',
            'color': '#ef4444'
        }
    elif ai_score >= 50:
        return {
            'level': 'medium',
            'text': '中风险',
            'suggestion': '建议使用"适度优化"模式，强度设为3-4级',
            'color': '#f59e0b'
        }
    elif ai_score >= 30:
        return {
            'level': 'low',
            'text': '低风险',
            'suggestion': '可选择"轻度优化"或手动调整部分内容',
            'color': '#3b82f6'
        }
    else:
        return {
            'level': 'safe',
            'text': '安全',
            'suggestion': '文本AI特征不明显，可直接使用',
            'color': '#10b981'
        }


@app.route('/api/text-optimize-cost', methods=['POST'])
def api_text_optimize_cost():
    """获取文本优化费用预估API"""
    try:
        data = request.json or {}
        char_count = data.get('char_count', 0)

        # 获取费用配置
        base_cost = float(user_manager.get_system_config('text_optimize_base_cost', 1.00))
        per_1000_chars_cost = float(user_manager.get_system_config('text_optimize_per_1000_cost', 0.50))

        # 计算总费用
        total_cost = base_cost + (char_count / 1000) * per_1000_chars_cost
        total_cost = round(total_cost, 2)

        # 获取用户余额
        balance = 0
        logged_in = 'user_id' in session
        if logged_in:
            user_info = user_manager.get_user_info(session['user_id'])
            if user_info:
                balance = float(user_info['balance'])

        return jsonify({
            'success': True,
            'base_cost': base_cost,
            'per_1000_cost': per_1000_chars_cost,
            'total_cost': total_cost,
            'char_count': char_count,
            'balance': balance,
            'logged_in': logged_in,
            'sufficient': balance >= total_cost if logged_in else False
        })

    except Exception as e:
        app.logger.error(f"获取文本优化费用失败: {e}")
        return jsonify({'success': False, 'message': str(e)}), 500


@app.route('/api/optimize-text', methods=['POST'])
@login_required
def api_optimize_text():
    """文本优化API接口 - 支持多种优化模式和强度（收费功能）"""
    try:
        data = request.json
        text_content = data.get('text', '').strip()

        # 获取优化选项
        options = data.get('options', {})
        intensity = options.get('intensity', 3)  # 1-5，默认3
        mode = options.get('mode', 'balanced')  # balanced/aggressive/conservative
        preserve_terms = options.get('preserve_terms', True)
        diversify_sentence = options.get('diversify_sentence', True)
        natural_transition = options.get('natural_transition', True)
        add_human_touch = options.get('add_human_touch', True)

        if not text_content:
            return jsonify({'error': '请输入要优化的文本'}), 400

        if len(text_content) > 50000:  # 限制文本长度
            return jsonify({'error': '文本长度超限，请控制在5万字符以内'}), 400

        # === 收费逻辑 ===
        user_id = session.get('user_id')

        # 根据文本长度计算费用：基础费用 + 每1000字额外费用
        base_cost = float(user_manager.get_system_config('text_optimize_base_cost', 1.00))
        per_1000_chars_cost = float(user_manager.get_system_config('text_optimize_per_1000_cost', 0.50))

        # 计算总费用：基础费 + (字数/1000) * 每千字费用
        char_count = len(text_content)
        total_cost = base_cost + (char_count / 1000) * per_1000_chars_cost
        total_cost = round(total_cost, 2)  # 保留两位小数

        # 检查用户余额
        user_info = user_manager.get_user_info(user_id)
        if not user_info:
            return jsonify({'error': '获取用户信息失败，请重新登录'}), 401

        if user_info['balance'] < total_cost:
            return jsonify({
                'error': f'余额不足，文本优化需要 {total_cost:.2f} 元，当前余额 {user_info["balance"]:.2f} 元',
                'need_recharge': True,
                'cost': total_cost,
                'balance': float(user_info['balance'])
            }), 400

        app.logger.info(f"开始优化文本，原文长度: {len(text_content)}, 模式: {mode}, 强度: {intensity}, 费用: {total_cost}元")

        # 分段处理
        segments = split_text_intelligently(text_content)
        optimized_segments = []

        app.logger.info(f"文本分为 {len(segments)} 个片段进行处理")

        # 构建优化配置
        optimization_config = {
            'intensity': intensity,
            'mode': mode,
            'preserve_terms': preserve_terms,
            'diversify_sentence': diversify_sentence,
            'natural_transition': natural_transition,
            'add_human_touch': add_human_touch
        }

        for i, segment in enumerate(segments):
            try:
                # 提供上下文（前一段的结尾和后一段的开头）
                context = ""
                if i > 0:
                    context += f"前文：...{segments[i-1][-200:]}\n"
                if i < len(segments) - 1:
                    context += f"后文：{segments[i+1][:200]}..."

                app.logger.info(f"正在优化第 {i+1}/{len(segments)} 个片段")

                optimized = optimize_text_with_deepseek(segment, context, optimization_config)
                if optimized:
                    optimized_segments.append(optimized)
                    app.logger.info(f"第 {i+1} 个片段优化成功")
                else:
                    optimized_segments.append(segment)  # 失败时保留原文
                    app.logger.warning(f"第 {i+1} 个片段优化失败，保留原文")

                # 短暂延迟，避免API调用过于频繁
                time.sleep(0.5)

            except Exception as segment_error:
                app.logger.error(f"优化第 {i+1} 个片段时出错: {segment_error}")
                optimized_segments.append(segment)  # 出错时保留原文

        # 组装最终结果
        final_result = '\n\n'.join(optimized_segments)

        # 检查是否有实际优化（至少有一个片段成功优化）
        optimization_success = any(opt != orig for opt, orig in zip(optimized_segments, segments))

        # 模式和强度名称映射
        mode_names = {'balanced': '均衡模式', 'aggressive': '强力模式', 'conservative': '保守模式', 'moderate': '适度模式', 'rewrite': '完全重写'}
        intensity_names = {1: '轻微', 2: '适度', 3: '中等', 4: '深度', 5: '极限'}

        # === 扣费逻辑：只有优化成功才扣费 ===
        charged = False
        if optimization_success:
            consume_result = user_manager.consume_balance(
                user_id,
                total_cost,
                'text_optimize',
                f'AI文本优化 - {char_count}字 | {mode_names.get(mode, "均衡模式")}'
            )
            charged = consume_result['success']
            if charged:
                app.logger.info(f"文本优化成功，已扣费 {total_cost} 元")
            else:
                app.logger.warning(f"文本优化扣费失败: {consume_result.get('message', '未知错误')}")
        else:
            app.logger.warning("文本优化失败，未扣费")

        # 生成优化报告
        processing_info = f'共处理 {len(segments)} 个文本片段 | {mode_names.get(mode, "均衡模式")} | {intensity_names.get(intensity, "中等")}优化'

        app.logger.info(f"文本优化完成，优化后长度: {len(final_result)}")

        # 获取扣费后的最新余额
        updated_user_info = user_manager.get_user_info(user_id)
        new_balance = float(updated_user_info['balance']) if updated_user_info else 0

        return jsonify({
            'success': True,
            'original': text_content,
            'optimized': final_result,
            'segments_count': len(segments),
            'processing_info': processing_info,
            'original_length': len(text_content),
            'optimized_length': len(final_result),
            'optimization_mode': mode,
            'optimization_intensity': intensity,
            'charged': charged,
            'cost': total_cost,
            'new_balance': new_balance
        })

    except Exception as e:
        app.logger.error(f"文本优化API错误: {e}")
        return jsonify({'error': f'优化失败: {str(e)}'}), 500


def generate_paper_with_custom_outline(task_id, title, field, paper_type, abstract, keywords, requirements, custom_outline):
    """根据用户自定义目录生成论文内容并更新进度"""
    try:
        app.logger.info(f"开始根据自定义目录生成论文: {title}")

        # 使用用户提供的自定义目录结构
        sections = custom_outline

        # 更新任务状态
        paper_generation_tasks[task_id].update({
            'status': 'generating',
            'progress': 5,
            'total_sections': len(sections),
            'message': f'开始生成论文，共{len(sections)}个章节'
        })

        # 初始化记忆系统用于文献引用收集  
        memory = {
            'collected_references': [],
            'reference_counter': 0,
            'system_context': {
                'tech_stack': 'Spring Boot + Vue.js + MySQL',
                'database_info': 'MySQL关系型数据库',
                'key_features': '信息管理、数据统计、权限控制',
                'research_objectives': '提升管理效率，实现数字化转型'
            }
        }

        complete_content = ""

        # 逐个生成章节
        for i, section in enumerate(sections):
            try:
                # 更新进度 - 确保进度从5%开始，到95%结束
                progress = int(5 + (i / len(sections)) * 90)
                paper_generation_tasks[task_id].update({
                    'progress': progress,
                    'current_section': section['name'],
                    'message': f'正在生成: {section["name"]} ({i+1}/{len(sections)})'
                })

                app.logger.info(f"生成章节 {i+1}/{len(sections)}: {section['name']}")

                # 生成章节内容
                if section['name'] == "参考文献":
                    # 使用收集的文献引用生成参考文献
                    section_content = generate_collected_references(memory)
                else:
                    section_content = generate_simple_section_content(
                        title, field, paper_type, section, abstract, keywords, requirements, i+1, memory
                    )

                if section_content and len(section_content.strip()) > 50:
                    complete_content += section_content + "\n\n"
                    app.logger.info(f"章节 {section['name']} 生成成功，长度: {len(section_content)}")
                else:
                    if section['name'] == "参考文献":
                        # 参考文献生成失败，给出明确提示
                        error_message = """<h2>参考文献</h2>
<div style="color: #d32f2f; background: #ffebee; border: 1px solid #f8bbd9; padding: 15px; border-radius: 5px; margin: 10px 0;">
<h3>⚠️ 参考文献生成失败</h3>
<p><strong>原因：</strong>AI生成参考文献时遇到问题，无法获取真实的学术文献数据。</p>
<p><strong>建议：</strong>请手动添加真实的参考文献。您可以到以下数据库搜索相关文献：</p>
<ul>
<li>中国知网(CNKI) - <a href="https://cnki.net" target="_blank">cnki.net</a></li>
<li>IEEE Xplore - <a href="https://ieeexplore.ieee.org" target="_blank">ieeexplore.ieee.org</a></li>
<li>ACM Digital Library - <a href="https://dl.acm.org" target="_blank">dl.acm.org</a></li>
<li>万方数据库 - <a href="https://wanfangdata.com.cn" target="_blank">wanfangdata.com.cn</a></li>
</ul>
<p><strong>格式要求：</strong>[序号] 作者. 标题[J/M/C/D]. 期刊/出版社/会议/学校, 年份, 卷(期): 页码.</p>
</div>"""
                        complete_content += error_message + "\n\n"
                        app.logger.error(f"参考文献生成失败，已添加错误提示")
                    else:
                        # 其他章节生成失败，跳过该章节
                        app.logger.error(f"章节 {section['name']} 生成失败，跳过该章节")
                        continue

                # 更新已完成章节数
                paper_generation_tasks[task_id].update({
                    'sections_completed': i + 1,
                    'message': f'{section["name"]} 章节生成完成'
                })

                # 短暂延迟，避免API调用过于频繁
                time.sleep(1)

            except Exception as section_error:
                app.logger.error(f"生成章节 {section['name']} 时出错: {section_error}")
                # 跳过失败的章节，不使用备用内容
                continue

        # 生成完成
        paper_generation_tasks[task_id].update({
            'status': 'completed',
            'progress': 100,
            'content': complete_content,
            'message': '论文生成完成！'
        })

        app.logger.info(f"论文生成完成: {title}")

    except Exception as e:
        app.logger.error(f"论文生成过程中发生错误: {e}")
        paper_generation_tasks[task_id].update({
            'status': 'error',
            'error': str(e),
            'message': f'生成失败: {str(e)}'
        })


def generate_paper_with_progress(task_id, title, field, paper_type, target_words, abstract, keywords, requirements, use_three_level=False):
    """重写的论文生成流程 - 简化且可靠，支持二级和三级标题"""
    try:
        # 根据用户选择生成不同的章节结构
        if use_three_level:
            sections = generate_three_level_sections(target_words)
        else:
            sections = generate_two_level_sections(target_words)
        
        # 更新任务状态
        paper_generation_tasks[task_id].update({
            'total_sections': len(sections),
            'progress': 5,
            'message': f'开始生成{"三级标题" if use_three_level else "二级标题"}标准本科毕业论文...',
            'start_time': time.time()
        })
        
        # 初始化记忆系统用于文献引用收集
        memory = {
            'collected_references': [],
            'reference_counter': 0,
            'system_context': {
                'tech_stack': 'Spring Boot + Vue.js + MySQL',
                'database_info': 'MySQL关系型数据库',
                'key_features': '信息管理、数据统计、权限控制',
                'research_objectives': '提升管理效率，实现数字化转型'
            }
        }
        
        # 生成完整论文内容
        complete_content = f'<h1 style="text-align: center; margin-bottom: 30px;">{title}</h1>\n\n'
        
        # 逐章节生成内容
        for i, section in enumerate(sections):
            try:
                # 更新进度
                progress = 10 + (i * 80 // len(sections))
                paper_generation_tasks[task_id].update({
                    'current_section': section['name'],
                    'sections_completed': i,
                    'progress': progress,
                    'message': f'正在生成 {section["name"]}...'
                })
                
                app.logger.info(f"开始生成章节: {section['name']}")
                
                # 生成章节内容
                if section['name'] == "参考文献":
                    # 使用收集的文献引用生成参考文献
                    section_content = generate_collected_references(memory)
                else:
                    section_content = generate_simple_section_content(
                        title, field, paper_type, section, abstract, keywords, requirements, i+1, memory
                    )
                
                if section_content and len(section_content.strip()) > 50:
                    complete_content += section_content + "\n\n"
                    app.logger.info(f"章节 {section['name']} 生成成功，长度: {len(section_content)}")
                else:
                    if section['name'] == "参考文献":
                        # 参考文献生成失败，给出明确提示
                        error_message = """<h2>参考文献</h2>
<div style="color: #d32f2f; background: #ffebee; border: 1px solid #f8bbd9; padding: 15px; border-radius: 5px; margin: 10px 0;">
<h3>⚠️ 参考文献生成失败</h3>
<p><strong>原因：</strong>AI生成参考文献时遇到问题，无法获取真实的学术文献数据。</p>
<p><strong>建议：</strong>请手动添加真实的参考文献。您可以到以下数据库搜索相关文献：</p>
<ul>
<li>中国知网(CNKI) - <a href="https://cnki.net" target="_blank">cnki.net</a></li>
<li>IEEE Xplore - <a href="https://ieeexplore.ieee.org" target="_blank">ieeexplore.ieee.org</a></li>
<li>ACM Digital Library - <a href="https://dl.acm.org" target="_blank">dl.acm.org</a></li>
<li>万方数据库 - <a href="https://wanfangdata.com.cn" target="_blank">wanfangdata.com.cn</a></li>
</ul>
<p><strong>格式要求：</strong>[序号] 作者. 标题[J/M/C/D]. 期刊/出版社/会议/学校, 年份, 卷(期): 页码.</p>
</div>"""
                        complete_content += error_message + "\n\n"
                        app.logger.error(f"参考文献生成失败，已添加错误提示")
                    else:
                        # 其他章节生成失败，跳过该章节
                        app.logger.error(f"章节 {section['name']} 生成失败，跳过该章节")
                        continue
                
                # 更新完成状态
                paper_generation_tasks[task_id].update({
                    'sections_completed': i + 1,
                    'message': f'{section["name"]} 章节生成完成'
                })
                
                # 短暂延迟避免API频率限制
                time.sleep(1)
                
            except Exception as section_error:
                app.logger.error(f"生成章节 {section['name']} 时出错: {section_error}")
                # 跳过失败的章节，不使用备用内容
                continue
        
        # 任务完成
        total_length = len(complete_content)
        paper_generation_tasks[task_id].update({
            'progress': 100,
            'message': f'论文生成完成！总字数约 {total_length} 字符',
            'status': 'completed',
            'content': complete_content,
            'sections_completed': len(sections),
            'current_section': '已完成',
            'total_chars': total_length
        })
        
        app.logger.info(f"论文生成完成 - 任务ID: {task_id}, 总长度: {total_length}")
        
    except Exception as e:
        app.logger.error(f"论文生成过程发生错误: {e}")
        paper_generation_tasks[task_id].update({
            'status': 'error',
            'error': str(e),
            'message': f'论文生成失败: {str(e)}'
        })


def generate_intelligent_outline(description, total_words, field, paper_type):
    """根据用户描述智能生成论文目录结构"""
    try:
        # 构建AI提示词
        prompt = f"""作为学术论文专家，请根据以下信息智能生成一个完整的{paper_type}目录结构：

【用户描述】：{description}

【基本要求】：
- 研究领域：{field}
- 论文类型：{paper_type}
- 总字数：{total_words}字
- 目录层级：最多三级标题

【分析任务】：
1. 分析用户描述，理解研究内容和技术特点
2. 确定合适的论文结构（理论研究/系统开发/实验分析等）
3. 生成符合学术规范的章节标题
4. 合理分配各章节字数（确保总和等于{total_words}字）

【标准论文结构参考】：
- 摘要（300字）
- Abstract（300字）
- 第1章 绪论（15-20%）
- 第2章 相关技术/理论基础（15-20%）
- 第3章 需求分析/问题分析（15-20%）
- 第4章 系统设计/方法设计（20-25%）
- 第5章 实现/实验（15-20%）
- 第6章 测试/结果分析（10-15%）
- 第7章 总结与展望（5-10%）
- 参考文献
- 致谢（200字）

【输出格式】：请严格按照以下JSON格式输出，不要包含任何其他文字：
{{
    "title": "根据描述生成的论文标题",
    "analysis": "对用户描述的分析和理解",
    "paper_structure": "确定的论文结构类型",
    "sections": [
        {{"name": "摘要", "words": 300, "description": "中文摘要和关键词", "level": 1}},
        {{"name": "Abstract", "words": 300, "description": "英文摘要和关键词", "level": 1}},
        {{"name": "第1章 绪论", "words": 2500, "description": "具体描述该章节内容", "level": 2}},
        {{"name": "1.1 研究背景", "words": 800, "description": "具体描述", "level": 3}},
        {{"name": "1.2 研究意义", "words": 600, "description": "具体描述", "level": 3}}
    ],
    "total_calculated_words": {total_words}
}}

请确保：
1. 章节标题准确反映研究内容
2. 字数分配合理，总和等于{total_words}
3. 包含必要的三级标题
4. 描述具体且有针对性"""

        # 调用AI生成目录
        headers = {
            'Authorization': f'Bearer {DEEPSEEK_API_KEY}',
            'Content-Type': 'application/json'
        }

        payload = {
            'model': 'deepseek-chat',
            'messages': [{'role': 'user', 'content': prompt}],
            'temperature': 0.7,
            'max_tokens': 3000
        }

        try:
            response = requests.post(DEEPSEEK_API_URL, headers=headers, json=payload, timeout=120)

            if response.status_code == 200:
                result = response.json()
                ai_response = result['choices'][0]['message']['content']

                # 尝试解析JSON响应
                try:
                    import json
                    # 提取JSON部分
                    start_idx = ai_response.find('{')
                    end_idx = ai_response.rfind('}') + 1
                    if start_idx != -1 and end_idx != 0:
                        json_str = ai_response[start_idx:end_idx]
                        outline_data = json.loads(json_str)
                        return outline_data
                    else:
                        raise json.JSONDecodeError("未找到JSON格式", ai_response, 0)
                except json.JSONDecodeError:
                    app.logger.warning("AI返回的不是有效JSON，使用默认结构")
                    return generate_default_outline_from_description(description, total_words, field, paper_type)
            else:
                app.logger.error(f"AI API调用失败: {response.status_code}")
                return generate_default_outline_from_description(description, total_words, field, paper_type)

        except Exception as e:
            app.logger.error(f"AI API调用异常: {e}")
            return generate_default_outline_from_description(description, total_words, field, paper_type)

    except Exception as e:
        app.logger.error(f"智能生成目录失败: {e}")
        # 出错时返回默认结构
        return generate_default_outline_from_description(description, total_words, field, paper_type)


def generate_intelligent_outline_enhanced(title, field, paper_type, total_words, abstract, keywords, special_requirements, outline_level='two'):
    """增强版智能目录生成 - 基于完整的用户输入信息"""
    try:
        # 根据级别设置标题结构
        level_description = ""
        if outline_level == 'three':
            level_description = """
【标题结构要求】
- 使用三级标题结构（如：第1章 → 1.1 → 1.1.1）
- 每个一级章节下包含2-4个二级标题
- 重要的二级标题下可包含2-3个三级标题
- 三级标题要具体到实现细节或具体方法
- 适合硕士及以上学位论文的详细结构"""
        else:
            level_description = """
【标题结构要求】
- 使用二级标题结构（如：第1章 → 1.1）
- 每个一级章节下包含2-4个二级标题
- 结构简洁明了，适合本科毕业论文
- 二级标题要涵盖该章节的主要内容"""

        # 构建更详细的AI提示词
        prompt = f"""作为资深学术论文专家，请根据以下详细信息智能生成一个高质量的{paper_type}目录结构：

【论文基本信息】
论文题目：{title}
研究领域：{field}
论文类型：{paper_type}
目标字数：{total_words}字

【研究内容】
论文摘要：{abstract}
关键词：{keywords}
详细描述：{special_requirements}

{level_description}

【生成要求】
1. 根据论文内容智能分析，生成最适合的章节结构
2. 字数分配要合理，总和必须等于{total_words}字
3. 章节标题要准确反映研究内容，具有学术性
4. 每个章节要有具体的描述说明
5. 结构要符合{paper_type}的学术规范

【输出格式】
请严格按照以下JSON格式输出，不要包含任何其他文字：
{{
    "sections": [
        {{"name": "摘要", "words": 300, "description": "中文摘要和关键词"}},
        {{"name": "Abstract", "words": 250, "description": "英文摘要和关键词"}},
        {{"name": "第1章 绪论", "words": 1200, "description": "研究背景、意义、现状、内容与方法"}},
        {{"name": "第2章 相关技术介绍", "words": 1500, "description": "相关技术和理论基础"}},
        {{"name": "第3章 需求分析与系统设计", "words": 2000, "description": "需求分析和系统设计"}},
        {{"name": "第4章 系统详细设计与实现", "words": 2500, "description": "详细设计和实现"}},
        {{"name": "第5章 系统测试", "words": 1500, "description": "测试方案和结果分析"}},
        {{"name": "第6章 总结与展望", "words": 800, "description": "工作总结和未来展望"}},
        {{"name": "参考文献", "words": 0, "description": "参考文献列表"}}
    ]
}}

请确保：
1. 章节标题准确反映研究内容
2. 字数分配合理，总和等于{total_words}
3. 描述具体且有针对性
4. 严格按照JSON格式输出"""

        # 调用AI生成目录
        headers = {
            'Authorization': f'Bearer {DEEPSEEK_API_KEY}',
            'Content-Type': 'application/json'
        }

        payload = {
            'model': 'deepseek-chat',
            'messages': [{'role': 'user', 'content': prompt}],
            'temperature': 0.7,
            'max_tokens': 3000
        }

        try:
            app.logger.info("开始调用AI生成智能目录...")
            response = requests.post(DEEPSEEK_API_URL, headers=headers, json=payload, timeout=120)

            if response.status_code == 200:
                result = response.json()
                ai_response = result['choices'][0]['message']['content']
                app.logger.info(f"AI响应长度: {len(ai_response)}")

                # 尝试解析JSON响应
                try:
                    import json
                    # 提取JSON部分
                    start_idx = ai_response.find('{')
                    end_idx = ai_response.rfind('}') + 1
                    if start_idx != -1 and end_idx != 0:
                        json_str = ai_response[start_idx:end_idx]
                        outline_data = json.loads(json_str)

                        # 验证并返回sections数组
                        if 'sections' in outline_data and isinstance(outline_data['sections'], list) and len(outline_data['sections']) > 0:
                            app.logger.info(f"AI智能目录生成成功，共{len(outline_data['sections'])}个章节")
                            return outline_data['sections']
                        else:
                            app.logger.warning("AI返回的JSON格式不正确或为空")
                            raise json.JSONDecodeError("JSON格式不正确", ai_response, 0)
                    else:
                        app.logger.warning("AI响应中未找到JSON格式")
                        raise json.JSONDecodeError("未找到JSON格式", ai_response, 0)
                except json.JSONDecodeError as je:
                    app.logger.warning(f"AI返回的不是有效JSON: {je}，使用默认结构")
                    return generate_default_outline_from_description(special_requirements, total_words, field, paper_type)
            else:
                app.logger.error(f"AI API调用失败: {response.status_code}, 响应: {response.text}")
                return generate_default_outline_from_description(special_requirements, total_words, field, paper_type)

        except requests.exceptions.Timeout:
            app.logger.error("AI API调用超时")
            return generate_default_outline_from_description(special_requirements, total_words, field, paper_type)
        except requests.exceptions.RequestException as re:
            app.logger.error(f"AI API网络请求异常: {re}")
            return generate_default_outline_from_description(special_requirements, total_words, field, paper_type)
        except Exception as e:
            app.logger.error(f"AI API调用异常: {e}")
            return generate_default_outline_from_description(special_requirements, total_words, field, paper_type)

    except Exception as e:
        app.logger.error(f"智能生成目录失败: {e}")
        # 出错时返回默认结构
        return generate_default_outline_from_description(special_requirements, total_words, field, paper_type)


def generate_default_outline_from_description(description, total_words, field, paper_type):
    """根据描述生成默认目录结构"""
    # 分析描述中的关键词来调整标题
    is_system = any(keyword in description.lower() for keyword in ['系统', 'system', '平台', 'platform', '网站', 'website'])
    is_algorithm = any(keyword in description.lower() for keyword in ['算法', 'algorithm', '模型', 'model'])
    is_analysis = any(keyword in description.lower() for keyword in ['分析', 'analysis', '研究', 'research'])

    # 基础结构
    sections = [
        {"name": "摘要", "words": 300, "description": "中文摘要和关键词", "level": 1},
        {"name": "Abstract", "words": 300, "description": "英文摘要和关键词", "level": 1},
    ]

    # 根据内容类型调整章节
    remaining_words = total_words - 600  # 减去摘要字数

    if is_system:
        # 系统开发类论文
        sections.extend([
            {"name": "第1章 绪论", "words": int(remaining_words * 0.18), "description": "研究背景、意义、现状分析", "level": 2},
            {"name": "第2章 相关技术介绍", "words": int(remaining_words * 0.18), "description": "技术基础和开发框架", "level": 2},
            {"name": "第3章 需求分析与系统设计", "words": int(remaining_words * 0.20), "description": "需求分析和总体设计", "level": 2},
            {"name": "第4章 系统详细设计与实现", "words": int(remaining_words * 0.25), "description": "详细设计和功能实现", "level": 2},
            {"name": "第5章 系统测试", "words": int(remaining_words * 0.12), "description": "测试方案和结果分析", "level": 2},
            {"name": "第6章 总结与展望", "words": int(remaining_words * 0.07), "description": "工作总结和未来展望", "level": 2},
        ])
    elif is_algorithm:
        # 算法研究类论文
        sections.extend([
            {"name": "第1章 绪论", "words": int(remaining_words * 0.18), "description": "研究背景和问题提出", "level": 2},
            {"name": "第2章 相关工作", "words": int(remaining_words * 0.20), "description": "相关算法和理论基础", "level": 2},
            {"name": "第3章 问题分析与建模", "words": int(remaining_words * 0.18), "description": "问题建模和分析", "level": 2},
            {"name": "第4章 算法设计与实现", "words": int(remaining_words * 0.25), "description": "算法设计和实现方案", "level": 2},
            {"name": "第5章 实验与分析", "words": int(remaining_words * 0.12), "description": "实验设计和结果分析", "level": 2},
            {"name": "第6章 总结与展望", "words": int(remaining_words * 0.07), "description": "研究总结和未来工作", "level": 2},
        ])
    else:
        # 通用研究类论文
        sections.extend([
            {"name": "第1章 绪论", "words": int(remaining_words * 0.18), "description": "研究背景、意义和内容", "level": 2},
            {"name": "第2章 理论基础", "words": int(remaining_words * 0.20), "description": "相关理论和技术基础", "level": 2},
            {"name": "第3章 研究方法", "words": int(remaining_words * 0.18), "description": "研究方法和技术路线", "level": 2},
            {"name": "第4章 研究内容", "words": int(remaining_words * 0.25), "description": "主要研究内容和实现", "level": 2},
            {"name": "第5章 结果分析", "words": int(remaining_words * 0.12), "description": "结果分析和讨论", "level": 2},
            {"name": "第6章 总结与展望", "words": int(remaining_words * 0.07), "description": "研究总结和展望", "level": 2},
        ])

    # 添加结尾部分
    sections.extend([
        {"name": "参考文献", "words": 0, "description": "参考文献列表", "level": 1},
        {"name": "致谢", "words": 200, "description": "致谢", "level": 1}
    ])

    return {
        "title": f"基于{field}的{description[:20]}研究",
        "analysis": f"根据描述分析，这是一个{field}领域的{'系统开发' if is_system else '算法研究' if is_algorithm else '理论研究'}类论文",
        "paper_structure": "系统开发类" if is_system else "算法研究类" if is_algorithm else "理论研究类",
        "sections": sections,
        "total_calculated_words": total_words
    }


def generate_two_level_sections(target_words):
    """生成二级标题章节结构"""
    return [
        {"name": "摘要", "words": 300, "description": "中文摘要和关键词"},
        {"name": "Abstract", "words": 250, "description": "英文摘要和关键词"},
        {"name": "第1章 绪论", "words": 1200, "description": "研究背景、意义、现状、内容与方法"},
        {"name": "第2章 相关技术介绍", "words": 1500, "description": "开发框架、数据库、前端技术、系统架构"},
        {"name": "第3章 需求分析与系统设计", "words": 2000, "description": "需求分析、系统总体设计、数据库设计"},
        {"name": "第4章 系统详细设计与实现", "words": 2500, "description": "系统功能模块、关键技术实现、部署配置"},
        {"name": "第5章 系统测试", "words": 1500, "description": "测试环境、功能测试、性能测试、结果分析"},
        {"name": "第6章 总结与展望", "words": 800, "description": "工作总结、不足改进、未来展望"},
        {"name": "参考文献", "words": 0, "description": "参考文献列表"}
    ]


def generate_three_level_sections(target_words):
    """生成三级标题章节结构（15000字版本）"""
    return [
        {"name": "摘要", "words": 300, "description": "中文摘要和关键词"},
        {"name": "Abstract", "words": 250, "description": "英文摘要和关键词"},
        
        # 第1章 绪论 (1800字)
        {"name": "第1章 绪论", "words": 300, "description": "章节引言", "level": 2},
        {"name": "1.1 研究背景与意义", "words": 600, "description": "研究背景与意义详述", "level": 3},
        {"name": "1.2 国内外研究现状", "words": 700, "description": "国内外研究现状分析", "level": 3},
        {"name": "1.3 研究内容与方法", "words": 350, "description": "研究内容与方法", "level": 3},
        {"name": "1.4 主要贡献与创新点", "words": 250, "description": "主要贡献与创新点", "level": 3},
        {"name": "1.5 论文组织结构", "words": 150, "description": "论文组织结构", "level": 3},
        
        # 第2章 相关技术介绍 (2300字)
        {"name": "第2章 相关技术介绍", "words": 200, "description": "章节引言", "level": 2},
        {"name": "2.1 开发框架技术", "words": 300, "description": "开发框架技术概述", "level": 3},
        {"name": "2.1.1 前端框架技术", "words": 300, "description": "前端框架技术详述", "level": 4},
        {"name": "2.1.2 后端框架技术", "words": 300, "description": "后端框架技术详述", "level": 4},
        {"name": "2.2 数据库技术", "words": 250, "description": "数据库技术概述", "level": 3},
        {"name": "2.2.1 关系型数据库", "words": 250, "description": "关系型数据库技术", "level": 4},
        {"name": "2.2.2 NoSQL数据库", "words": 250, "description": "NoSQL数据库技术", "level": 4},
        {"name": "2.3 Web开发技术", "words": 300, "description": "Web开发技术概述", "level": 3},
        {"name": "2.3.1 前端技术栈", "words": 300, "description": "前端技术栈详述", "level": 4},
        {"name": "2.3.2 API设计与开发", "words": 300, "description": "API设计与开发", "level": 4},
        {"name": "2.4 系统架构设计模式", "words": 300, "description": "系统架构设计模式概述", "level": 3},
        {"name": "2.4.1 MVC架构模式", "words": 300, "description": "MVC架构模式详述", "level": 4},
        {"name": "2.4.2 微服务架构", "words": 300, "description": "微服务架构详述", "level": 4},
        
        # 第3章 需求分析与系统设计 (3000字)
        {"name": "第3章 需求分析与系统设计", "words": 200, "description": "章节引言", "level": 2},
        {"name": "3.1 需求分析", "words": 200, "description": "需求分析概述", "level": 3},
        {"name": "3.1.1 业务需求分析", "words": 400, "description": "业务需求分析", "level": 4},
        {"name": "3.1.2 功能需求分析", "words": 400, "description": "功能需求分析", "level": 4},
        {"name": "3.1.3 非功能需求分析", "words": 200, "description": "非功能需求分析", "level": 4},
        {"name": "3.2 系统总体设计", "words": 200, "description": "系统总体设计概述", "level": 3},
        {"name": "3.2.1 系统架构设计", "words": 400, "description": "系统架构设计", "level": 4},
        {"name": "3.2.2 技术架构设计", "words": 300, "description": "技术架构设计", "level": 4},
        {"name": "3.2.3 系统流程设计", "words": 300, "description": "系统流程设计", "level": 4},
        {"name": "3.3 数据库设计", "words": 200, "description": "数据库设计概述", "level": 3},
        {"name": "3.3.1 概念模型设计", "words": 300, "description": "概念模型设计", "level": 4},
        {"name": "3.3.2 逻辑模型设计", "words": 350, "description": "逻辑模型设计", "level": 4},
        {"name": "3.3.3 物理模型设计", "words": 350, "description": "物理模型设计", "level": 4},
        
        # 第4章 系统详细设计与实现 (3800字)
        {"name": "第4章 系统详细设计与实现", "words": 200, "description": "章节引言", "level": 2},
        {"name": "4.1 系统功能模块设计", "words": 200, "description": "系统功能模块设计概述", "level": 3},
        {"name": "4.1.1 用户管理模块", "words": 300, "description": "用户管理模块", "level": 4},
        {"name": "4.1.2 权限控制模块", "words": 300, "description": "权限控制模块", "level": 4},
        {"name": "4.1.3 核心业务模块", "words": 400, "description": "核心业务模块", "level": 4},
        {"name": "4.1.4 数据管理模块", "words": 200, "description": "数据管理模块", "level": 4},
        {"name": "4.2 关键技术实现", "words": 200, "description": "关键技术实现概述", "level": 3},
        {"name": "4.2.1 用户认证与权限控制", "words": 400, "description": "用户认证与权限控制", "level": 4},
        {"name": "4.2.2 数据交互接口设计", "words": 400, "description": "数据交互接口设计", "level": 4},
        {"name": "4.2.3 前端交互实现", "words": 300, "description": "前端交互实现", "level": 4},
        {"name": "4.2.4 缓存机制设计", "words": 300, "description": "缓存机制设计", "level": 4},
        {"name": "4.3 系统安全设计", "words": 200, "description": "系统安全设计概述", "level": 3},
        {"name": "4.3.1 数据安全策略", "words": 300, "description": "数据安全策略", "level": 4},
        {"name": "4.3.2 系统安全防护", "words": 300, "description": "系统安全防护", "level": 4},
        {"name": "4.4 系统部署与配置", "words": 200, "description": "系统部署与配置概述", "level": 3},
        {"name": "4.4.1 开发环境配置", "words": 200, "description": "开发环境配置", "level": 4},
        {"name": "4.4.2 生产环境部署", "words": 200, "description": "生产环境部署", "level": 4},
        {"name": "4.4.3 性能优化配置", "words": 200, "description": "性能优化配置", "level": 4},
        
        # 第5章 系统测试 (2200字)
        {"name": "第5章 系统测试", "words": 200, "description": "章节引言", "level": 2},
        {"name": "5.1 测试环境搭建", "words": 400, "description": "测试环境搭建", "level": 3},
        {"name": "5.2 功能测试", "words": 200, "description": "功能测试概述", "level": 3},
        {"name": "5.2.1 单元测试", "words": 300, "description": "单元测试", "level": 4},
        {"name": "5.2.2 集成测试", "words": 300, "description": "集成测试", "level": 4},
        {"name": "5.2.3 系统测试", "words": 200, "description": "系统测试", "level": 4},
        {"name": "5.3 性能测试", "words": 200, "description": "性能测试概述", "level": 3},
        {"name": "5.3.1 负载测试", "words": 250, "description": "负载测试", "level": 4},
        {"name": "5.3.2 压力测试", "words": 250, "description": "压力测试", "level": 4},
        {"name": "5.4 安全性测试", "words": 300, "description": "安全性测试", "level": 3},
        {"name": "5.5 测试结果分析", "words": 200, "description": "测试结果分析", "level": 3},
        
        # 第6章 总结与展望 (1200字)
        {"name": "第6章 总结与展望", "words": 200, "description": "章节引言", "level": 2},
        {"name": "6.1 工作总结", "words": 500, "description": "工作总结", "level": 3},
        {"name": "6.2 系统特色与创新", "words": 300, "description": "系统特色与创新", "level": 3},
        {"name": "6.3 不足与改进", "words": 200, "description": "不足与改进", "level": 3},
        {"name": "6.4 未来工作展望", "words": 200, "description": "未来工作展望", "level": 3},
        
        {"name": "参考文献", "words": 0, "description": "参考文献列表"}
    ]


def clean_ai_generated_content(content):
    """简化版内容清理函数 - 解决格式损坏问题"""
    import re
    
    if not content or not isinstance(content, str):
        return content
    
    app.logger.info(f"开始内容清理，原始长度: {len(content)}")
    
    # 第一步：移除明显的AI解释性文字
    simple_patterns = [
        r'```[\s\S]*?```',  # 代码块
        r'这个HTML格式的.*?(?=<|$)',
        r'以上是.*?的内容[。！？]*',
        r'您可以根据.*?[。！？]*',
        r'希望这.*?[。！？]*',
        r'以下是.*?：\s*',
        r'注意：.*?(?=<|$)',
        r'说明：.*?(?=<|$)',
    ]
    
    cleaned_content = content
    for pattern in simple_patterns:
        cleaned_content = re.sub(pattern, '', cleaned_content, flags=re.DOTALL | re.IGNORECASE)
    
    # 第二步：简单处理转义字符 - 核心问题修复
    cleaned_content = cleaned_content.replace('\\\\n', '\n')
    cleaned_content = cleaned_content.replace('\\n', '\n')
    cleaned_content = cleaned_content.replace('\\\\', '')
    
    # 第三步：标准化换行符
    cleaned_content = re.sub(r'\n{3,}', '\n\n', cleaned_content)
    
    # 第四步：移除空的HTML标签
    cleaned_content = re.sub(r'<p[^>]*>\s*</p>', '', cleaned_content)
    cleaned_content = re.sub(r'<h[1-6][^>]*>\s*</h[1-6]>', '', cleaned_content)
    
    result = cleaned_content.strip()
    
    app.logger.info(f"内容清理完成，最终长度: {len(result)}")
    
    return result


def generate_section_with_detailed_content(title, field, paper_type, section, abstract, keywords, requirements, section_num, memory=None, target_words=None):
    """详细章节内容生成 - 支持分小节多次调用API"""
    try:
        section_name = section['name']
        section_words = target_words or section.get('words', 1500)
        section_desc = section.get('description', '')

        # 获取标题级别
        level = section.get('level', 2)
        header_tag = f"h{level}"
        
        # 初始化文献引用计数器
        if memory and 'reference_counter' not in memory:
            memory['reference_counter'] = 0
            memory['collected_references'] = []
        
        # 获取已生成内容的上下文
        context_info = ""
        if memory and 'generated_sections' in memory and len(memory['generated_sections']) > 0:
            recent_sections = memory['generated_sections'][-2:]  # 最近2个章节
            context_summary = []
            for sec in recent_sections:
                if sec.get('summary'):
                    context_summary.extend(sec['summary'][:2])
            if context_summary:
                context_info = f"\n【前文要点】：{'; '.join(context_summary[:3])}"

        # 根据不同章节类型使用不同策略
        if "摘要" in section_name:
            return generate_abstract_section(title, field, abstract, keywords, section_words, memory)
        elif "Abstract" in section_name:
            return generate_english_abstract(title, field, abstract, keywords, section_words, memory)
        elif "第1章" in section_name or "绪论" in section_name:
            return generate_introduction_chapter(title, field, section_words, memory, context_info)
        elif "第2章" in section_name or "技术" in section_name:
            return generate_technology_chapter(title, field, section_words, memory, context_info)
        elif "第3章" in section_name or "需求" in section_name or "设计" in section_name:
            return generate_design_chapter(title, field, section_words, memory, context_info)
        elif "第4章" in section_name or "实现" in section_name:
            return generate_implementation_chapter(title, field, section_words, memory, context_info)
        elif "第5章" in section_name or "测试" in section_name:
            return generate_testing_chapter(title, field, section_words, memory, context_info)
        elif "第6章" in section_name or "总结" in section_name:
            return generate_conclusion_chapter(title, field, section_words, memory, context_info)
        else:
            return generate_generic_section(title, field, section, section_words, memory, context_info)
            
    except Exception as e:
        app.logger.error(f"详细章节生成失败 {section_name}: {e}")
        return generate_fallback_section(section, section_num)


def generate_introduction_chapter(title, field, section_words, memory, context_info):
    """生成绪论章节 - 分小节多次调用"""
    try:
        # 分成4个小节，每个小节单独生成
        subsections = [
            {"name": "1.1 研究背景与意义", "words": int(section_words * 0.3)},
            {"name": "1.2 国内外研究现状", "words": int(section_words * 0.3)},
            {"name": "1.3 研究内容与方法", "words": int(section_words * 0.25)},
            {"name": "1.4 论文组织结构", "words": int(section_words * 0.15)}
        ]
        
        complete_content = "<h2>第1章 绪论</h2>\n\n"
        current_ref_start = memory.get('reference_counter', 0) + 1
        
        for i, subsection in enumerate(subsections):
            subsection_content = generate_subsection_content(
                title, field, subsection, memory, context_info, current_ref_start + i*3
            )
            complete_content += subsection_content + "\n\n"
            
        return complete_content
        
    except Exception as e:
        app.logger.error(f"绪论章节生成失败: {e}")
        return f"<h2>第1章 绪论</h2>\n<p>内容生成失败，请重试。</p>"


def generate_subsection_content(title, field, subsection, memory, context_info, ref_start_num):
    """生成单个小节内容 - 高质量长文本"""
    try:
        subsection_name = subsection['name']
        target_words = subsection['words']
        
        # 构建详细的小节生成提示词
        prompt = f"""请为{field}领域的论文《{title}》生成{subsection_name}小节内容。

目标字数：{target_words}字（必须达到）
{context_info}

【关键要求】：
1. 内容必须达到{target_words}字，不得少于目标字数的90%
2. 必须包含3-5个文献引用，引用编号从[{ref_start_num}]开始连续编号
3. 每个段落250-350字，共需要{target_words//300 + 1}个段落
4. 内容要具体、深入、专业，避免空洞表述
5. 必须紧密结合{title}系统的特点

【写作要求】：
- 使用学术论文的正式语言
- 每段必须包含1-2个引用标注
- 段落间要有逻辑递进关系
- 避免重复和冗余表达
- 数据和结论要有说服力

【输出格式】：
<h3>{subsection_name}</h3>
<p>第一段内容，至少250字...引用[{ref_start_num}]</p>
<p>第二段内容，至少250字...引用[{ref_start_num+1}]</p>
...

在内容最后添加本小节的参考文献：
<div class="temp-references">
[{ref_start_num}] 作者1. 相关研究1[J]. 期刊名, 年份, 卷(期): 页码.
[{ref_start_num+1}] 作者2. 相关研究2[C]. 会议名, 年份: 页码.
...
</div>

请确保内容质量高、字数充足、引用规范。"""

        # 调用API生成内容，使用更高的token限制
        content = call_deepseek_api(prompt, min(target_words * 4, 8000))
        
        if content and len(content.strip()) > 100:
            # 处理文献引用
            processed_content = process_references_in_content(content, memory, ref_start_num)
            
            # 验证字数
            text_content = re.sub(r'<[^>]+>', '', processed_content)
            actual_words = len(text_content.replace(' ', '').replace('\n', ''))
            
            if actual_words < target_words * 0.8:
                app.logger.warning(f"{subsection_name} 字数不足: {actual_words}/{target_words}")
                # 尝试补充内容
                processed_content = supplement_content_if_needed(processed_content, target_words, subsection_name)
            
            return processed_content
        else:
            return generate_fallback_subsection(subsection_name, target_words)
            
    except Exception as e:
        app.logger.error(f"小节内容生成失败 {subsection_name}: {e}")
        return generate_fallback_subsection(subsection_name, target_words)


def process_references_in_content(content, memory, ref_start_num):
    """处理内容中的文献引用 - 全新的全局引用管理系统"""
    try:
        import re

        if not memory:
            memory = {'reference_counter': 0, 'collected_references': []}

        # 确保引用收集列表存在
        if 'collected_references' not in memory:
            memory['collected_references'] = []

        # 获取当前全局引用计数器
        current_counter = memory.get('reference_counter', 0)

        # 第一步：重新编号正文中的引用，确保全文连续
        citation_count = 0
        def replace_citation(match):
            nonlocal citation_count
            citation_count += 1
            new_num = current_counter + citation_count
            return f'[{new_num}]'

        # 替换正文中的引用编号
        content = re.sub(r'\[(\d+)\]', replace_citation, content)

        # 第二步：处理临时引用部分
        temp_refs_match = re.search(r'<div class="temp-references">(.*?)</div>', content, re.DOTALL)
        if temp_refs_match:
            refs_text = temp_refs_match.group(1).strip()
            ref_lines = [line.strip() for line in refs_text.split('\n') if line.strip()]

            # 处理每个引用
            for line in ref_lines:
                if not line:
                    continue

                # 提取引用内容（支持多种格式）
                ref_content = None

                # 格式1: [数字] 内容
                if line.startswith('['):
                    ref_content = re.sub(r'^\[\d+\]\s*', '', line)
                # 格式2: 数字. 内容
                elif re.match(r'^\d+\.\s*', line):
                    ref_content = re.sub(r'^\d+\.\s*', '', line)
                # 格式3: 直接内容
                else:
                    ref_content = line

                if ref_content and ref_content.strip():
                    current_counter += 1
                    memory['collected_references'].append({
                        'number': current_counter,
                        'content': ref_content.strip()
                    })

            # 移除临时引用部分
            content = re.sub(r'<div class="temp-references">.*?</div>', '', content, flags=re.DOTALL)
        else:
            # 如果没有临时引用，只更新计数器
            current_counter += citation_count

        # 更新全局计数器
        memory['reference_counter'] = current_counter

        app.logger.info(f"引用处理完成，当前计数器: {current_counter}，已收集引用: {len(memory['collected_references'])}")
        return content

    except Exception as e:
        app.logger.error(f"引用处理失败: {e}")
        return content


def supplement_content_if_needed(content, target_words, section_name):
    """如果内容字数不足，尝试补充"""
    try:
        text_content = re.sub(r'<[^>]+>', '', content)
        actual_words = len(text_content.replace(' ', '').replace('\n', ''))
        
        if actual_words < target_words * 0.9:
            shortage = target_words - actual_words
            supplement = f"""
<div class="content-supplement">
<p><strong>📝 内容扩展提示:</strong> 本节内容当前{actual_words}字，建议补充{shortage}字以达到{target_words}字目标。可考虑：</p>
<ul>
<li>增加具体的技术实现细节和案例分析</li>
<li>补充相关研究的对比分析和发展趋势</li>
<li>添加更多的数据支撑和实证研究结果</li>
<li>扩展对未来发展方向的深入讨论</li>
</ul>
</div>"""
            content += supplement
        
        return content
        
    except Exception as e:
        app.logger.error(f"内容补充失败: {e}")
        return content


def generate_enhanced_section_content(title, field, paper_type, section, abstract, keywords, requirements, section_num, memory, target_words):
    """增强的章节内容生成 - 确保达到目标字数"""
    try:
        section_name = section['name']
        
        # 构建更详细的提示词，强调字数要求
        prompt = f"""请为{field}领域的{paper_type}《{title}》生成{section_name}章节内容。

【基本信息】：
- 章节：{section_name}
- 描述：{section.get('description', '')}
- 目标字数：{target_words}字（必须严格达到）
- 摘要参考：{abstract[:200] if abstract else ''}
- 关键词：{keywords}

【严格要求】：
1. 内容字数必须达到{target_words}字，不得少于{int(target_words * 0.95)}字
2. 包含5-8个文献引用，使用[1][2][3]格式标注
3. 分为{max(3, target_words//400)}个段落，每段300-400字
4. 内容要学术化、专业化，避免空话套话
5. 必须紧密结合{title}系统的具体特点

【内容要求】：
- 使用准确的学术术语和专业表达
- 每段都要有具体的技术细节或实证数据
- 段落间要有清晰的逻辑递进关系
- 引用要分布均匀，每段1-2个引用
- 避免重复表述，每句话都要有实质内容

【输出格式】：
<{section.get('level', 'h2')}>{section_name}</{section.get('level', 'h2')}>
<p>第一段内容（300-400字）...引用[1]</p>
<p>第二段内容（300-400字）...引用[2]</p>
...

在内容最后添加临时引用：
<div class="temp-references">
[1] 作者1. 相关研究1[J]. 期刊名, 年份, 卷(期): 页码.
[2] 作者2. 相关研究2[C]. 会议名, 年份: 页码.
...
</div>

请确保内容质量高、字数充足、逻辑清晰。"""

        # 使用更高的token限制生成长内容
        content = call_deepseek_api(prompt, min(target_words * 5, 10000))
        
        if content and len(content.strip()) > 100:
            # 处理引用并验证字数
            processed_content = process_references_in_content(content, memory, memory.get('reference_counter', 0) + 1)
            
            # 检查字数是否达标
            text_content = re.sub(r'<[^>]+>', '', processed_content)
            actual_words = len(text_content.replace(' ', '').replace('\n', ''))
            
            if actual_words < target_words * 0.9:
                app.logger.warning(f"{section_name} 字数不足: {actual_words}/{target_words}")
                processed_content = supplement_content_if_needed(processed_content, target_words, section_name)
            
            return processed_content
        else:
            return generate_fallback_section(section, section_num)
            
    except Exception as e:
        app.logger.error(f"增强章节生成失败 {section_name}: {e}")
        return generate_fallback_section(section, section_num)


def calculate_section_target_words(section, total_words):
    """计算章节目标字数"""
    try:
        section_words = section.get('words', 0)
        if section_words > 0:
            return section_words
        
        # 如果没有指定字数，按默认比例分配
        section_name = section['name']
        if '摘要' in section_name or 'Abstract' in section_name:
            return int(total_words * 0.06)
        elif '第1章' in section_name or '绪论' in section_name:
            return int(total_words * 0.20)
        elif '第2章' in section_name or '技术' in section_name:
            return int(total_words * 0.14)
        elif '第3章' in section_name or '需求' in section_name or '设计' in section_name:
            return int(total_words * 0.18)
        elif '第4章' in section_name or '实现' in section_name:
            return int(total_words * 0.22)
        elif '第5章' in section_name or '测试' in section_name:
            return int(total_words * 0.10)
        elif '第6章' in section_name or '总结' in section_name:
            return int(total_words * 0.04)
        else:
            return 1500  # 默认字数
            
    except Exception as e:
        app.logger.error(f"字数计算失败: {e}")
        return 1500


def generate_fallback_subsection(subsection_name, target_words):
    """生成备用小节内容"""
    return f"""
<h3>{subsection_name}</h3>
<div class="content-placeholder" style="background: #f8f9fa; border: 1px solid #dee2e6; padding: 20px; border-radius: 5px;">
<p><strong>⚠️ 内容生成异常</strong></p>
<p>目标字数：{target_words}字</p>
<p>建议手动补充以下内容：</p>
<ul>
<li>相关理论基础和技术背景</li>
<li>具体的实现方案和技术细节</li>
<li>数据分析和实证研究结果</li>
<li>与其他研究的比较分析</li>
</ul>
</div>"""


def generate_section_content_by_type(title, field, paper_type, section_name, section_words, header_tag, abstract, keywords, memory):
    """根据章节类型生成内容"""
    try:
        if section_name == "摘要":
            prompt = f"""请为{paper_type}《{title}》生成专业的中文摘要。

研究领域：{field}
目标字数：{section_words}字
预设摘要：{abstract if abstract else '无'}
预设关键词：{keywords if keywords else '无'}

内容要求：
1. 研究背景和意义
2. 主要研究方法和技术路线
3. 核心成果和创新点
4. 结论和应用价值

写作要求：
- 使用连贯的段落形式
- 包含5-7个专业关键词
- 内容详实，字数充足
- 体现学术价值

请直接输出HTML格式内容：
<{header_tag}>摘要</{header_tag}>
<p>摘要内容...</p>
<p><strong>关键词：</strong>关键词1；关键词2；关键词3</p>"""

        elif section_name == "Abstract":
            prompt = f"""Please generate a professional English abstract for the {paper_type} titled "{title}".

Research Field: {field}
Target Words: {section_words} words
Chinese Abstract: {abstract[:200] if abstract else 'Not provided'}

Requirements:
1. Research background and significance
2. Main research methods and technical approaches
3. Key findings and innovations
4. Conclusions and practical applications

Writing Requirements:
- Use coherent paragraph format
- Include 5-7 professional keywords
- Comprehensive content with sufficient word count
- Demonstrate academic value

Please output HTML format directly:
<{header_tag}>Abstract</{header_tag}>
<p>Abstract content...</p>
<p><strong>Keywords:</strong> keyword1; keyword2; keyword3</p>"""

        elif section_name == "第1章 绪论":
            # 获取用户具体系统信息
            system_info = ""
            if memory and 'system_context' in memory:
                ctx = memory['system_context']
                system_info = f"""
【用户系统具体信息】：
- 系统名称：{title}
- 技术栈：{ctx.get('tech_stack', 'Spring Boot + Vue.js + MySQL')}
- 数据库信息：{ctx.get('database_info', 'MySQL关系型数据库')}
- 核心功能：{ctx.get('key_features', '信息管理、数据统计、权限控制')}
- 研究目标：{ctx.get('research_objectives', '提升管理效率，实现数字化转型')}
"""
            
            prompt = f"""请为{paper_type}《{title}》生成第1章绪论。

研究领域：{field}
目标字数：{section_words}字
{system_info}

【重要要求】：
1. 必须紧密结合用户具体的系统信息，不要泛泛而谈
2. 在适当位置添加文献引用标注[1][2][3]等，每段至少1-2个引用
3. 引用应该真实反映该领域的研究现状和技术发展
4. 每个引用后要在内容最后列出对应的参考文献信息

【严格格式要求】：
- 主标题使用<h2>标签，子标题使用<h3>标签
- 绝对禁止生成四级标题（如1.1.1、1.1.2等），最多只能到三级
- 内容必须是连贯的段落形式，严格禁止使用任何形式的分点列表

【统一章节结构】：
1.1 研究背景与意义 - 结合{title}的具体背景和意义
1.2 国内外研究现状 - 分析{field}领域相关技术和方法的发展状况
1.3 论文研究内容 - 说明本{title}系统的具体研究内容和技术路线
1.4 论文组织结构 - 介绍各章节的主要内容安排

【文献引用要求】：
- 每段必须包含1-2个文献引用，使用[1][2]格式
- 引用要与内容相关，体现该领域的研究现状
- 在内容最后添加：
<div class="temp-references">
[1] 作者1. 相关研究标题[J]. 期刊名称, 年份, 卷(期): 页码.
[2] 作者2. 相关技术论文[C]. 会议名称, 年份: 页码.
...
</div>

【严格写作要求】：
- 必须具体提及{title}系统的特点和功能
- 技术发展背景要与{field}领域紧密相关
- 研究现状要体现当前{field}信息化建设的具体挑战
- 每个子节必须包含4-5个自然段落，每段280-320字，确保达到目标字数
- 使用"随着{field}信息化的发展"等具体表述

请严格按照统一结构输出HTML格式内容，确保引用编号连续且在最后提供对应的参考文献。"""

        elif "第2章" in section_name:
            prompt = f"""请为{paper_type}《{title}》生成第2章相关技术介绍。

研究领域：{field}
目标字数：{section_words}字

【严格格式要求】：
- 主标题使用<h2>标签，子标题使用<h3>标签
- 绝对禁止生成四级标题（如2.1.1、2.1.2等），最多只能到三级
- 内容必须是连贯的段落形式，严格禁止使用任何形式的分点列表

【统一章节结构】：
2.1 Spring Boot开发框架 - 详细介绍Spring Boot框架特性和应用
2.2 数据库技术 - 阐述MySQL数据库设计和Redis缓存技术
2.3 前端开发技术 - 说明Vue.js框架和前端技术栈
2.4 系统架构设计 - 介绍整体架构模式和设计思路

【严格写作要求】：
- 使用自然流畅的段落形式，严格避免任何分点列表（1.、2.、3.或•、-、①②③等）
- 技术介绍要用连贯的叙述语言描述技术原理和应用，不要写成特性罗列
- 每个技术点都要融入到段落的自然表达中，避免生硬的技术名词堆砌
- 每个子节必须包含4-5个自然段落，每段250-300字，内容要详细深入
- 段落之间要有逻辑递进关系，从技术原理到实际应用
- 使用"在技术选型方面"、"考虑到项目需求"、"通过对比分析发现"等自然过渡语
- 必须包含具体的技术原理、配置方法、性能优化等深度内容

【严格禁止的错误示例】：
"Spring框架具有以下特点：1）依赖注入 2）面向切面编程 3）声明式事务"
"数据库技术包括：①MySQL关系型数据库 ②Redis缓存数据库 ③连接池技术"
"前端技术栈包含：• Vue.js框架 • Element UI组件库 • Axios网络库"
"架构设计采用：1）分层架构 2）微服务架构 3）前后端分离"

【正确示例（必须参考）】：
"在现代Web应用开发中，Spring Boot框架凭借其成熟的设计理念和强大的功能特性，已经成为Java企业级开发的首选解决方案。该框架通过依赖注入机制实现了组件间的松耦合，使得系统的可维护性和可测试性得到显著提升。同时，其面向切面编程的特性为横切关注点的处理提供了优雅的解决方案，特别是在日志记录、事务管理等方面表现出色。考虑到项目需求中对开发效率的要求，Spring Boot的自动配置特性能够显著减少样板代码的编写，让开发者能够专注于业务逻辑的实现。在具体应用中，框架提供的starter依赖管理机制能够自动解决版本冲突问题，而内嵌的Tomcat服务器则简化了部署流程，这些特性共同构成了一个高效、稳定的开发平台。"

请严格按照统一结构输出HTML格式内容：
<h2>第2章 相关技术介绍</h2>
<h3>2.1 Spring Boot开发框架</h3>
<p>第一段：框架概述和技术背景...</p>
<p>第二段：核心特性和设计理念...</p>
<p>第三段：自动配置和依赖管理...</p>
<p>第四段：在项目中的具体应用和优势...</p>
<h3>2.2 数据库技术</h3>
<p>第一段：MySQL数据库的选型和特性...</p>
<p>第二段：数据库设计和性能优化...</p>
<p>第三段：Redis缓存技术的应用...</p>
<p>第四段：数据持久化和一致性保障...</p>
<h3>2.3 前端开发技术</h3>
<p>第一段：Vue.js框架的特性和优势...</p>
<p>第二段：组件化开发和状态管理...</p>
<p>第三段：前端工程化和构建优化...</p>
<p>第四段：用户体验和界面设计...</p>
<h3>2.4 系统架构设计</h3>
<p>第一段：架构设计理念和模式选择...</p>
<p>第二段：分层架构和模块划分...</p>
<p>第三段：前后端分离和API设计...</p>
<p>第四段：系统扩展性和维护性...</p>"""

        elif "第3章" in section_name:
            prompt = f"""请为{paper_type}《{title}》生成第3章需求分析与系统设计。

研究领域：{field}
目标字数：{section_words}字

【严格格式要求】：
- 主标题使用<h2>标签，子标题使用<h3>标签
- 绝对禁止生成四级标题（如3.1.1、3.1.2、3.2.1等），最多只能到三级
- 内容必须是连贯的段落形式，严格禁止使用任何形式的分点列表

【统一章节结构】：
3.1 系统需求分析 - 详细分析系统的功能需求和非功能需求
3.2 系统总体设计 - 介绍系统的整体架构和设计思路
3.3 数据库设计 - 阐述数据库的设计方案和实现策略

【严格写作要求】：
- 使用自然流畅的段落形式，严格避免任何分点列表（1.、2.、3.或•、-、①②③等）
- 需求分析要用连贯的叙述语言描述用户需求和系统目标，不要写成需求条目
- 设计方案要融入到段落的自然表达中，避免生硬的设计要素罗列
- 每个子节必须包含5-6个自然段落，每段250-300字，内容要详细深入
- 段落之间要有逻辑递进关系，从需求识别到设计方案
- 使用"通过深入调研发现"、"在设计过程中"、"为了满足...需求"等自然过渡语
- 必须包含具体的需求分析方法、设计原则、技术方案等深度内容

【严格禁止的错误示例】：
"系统功能需求包括：1）用户管理 2）数据处理 3）报表生成"
"非功能需求主要有：①性能要求 ②安全要求 ③可用性要求"
"数据库设计包含：• 用户表 • 图书表 • 借阅记录表"
"核心实体包括：1. 用户实体 2. 图书实体 3. 借阅记录实体"
"设计原则采用：1）高内聚低耦合 2）可扩展性 3）可维护性"

【正确示例（必须参考）】：
"通过对目标用户群体的深入调研和需求访谈，发现当前系统面临的核心挑战主要集中在用户体验和数据处理效率两个方面。用户普遍反映现有系统的操作复杂度较高，特别是在数据录入和查询方面缺乏直观的交互界面。同时，随着数据量的不断增长，系统的响应速度和处理能力已经成为制约业务发展的重要瓶颈。基于这些实际需求，本系统在设计理念上强调简洁高效的用户界面和强大的后台数据处理能力。通过采用现代化的技术架构和优化的算法设计，系统能够在保证功能完整性的同时，显著提升用户操作的便捷性和系统运行的稳定性。"

请严格按照统一结构输出HTML格式内容：
<h2>第3章 需求分析与系统设计</h2>
<h3>3.1 系统需求分析</h3>
<p>第一段：需求调研的背景和重要性...</p>
<p>第二段：功能需求的详细分析和用户场景...</p>
<p>第三段：非功能需求的识别和量化指标...</p>
<p>第四段：需求优先级划分和实现策略...</p>
<p>第五段：需求变更管理和风险控制...</p>
<h3>3.2 系统总体设计</h3>
<p>第一段：系统架构设计的理念和原则...</p>
<p>第二段：分层架构的详细设计和职责划分...</p>
<p>第三段：核心业务模块的设计和交互关系...</p>
<p>第四段：技术选型的依据和集成方案...</p>
<p>第五段：系统扩展性和可维护性设计...</p>
<h3>3.3 数据库设计</h3>
<p>第一段：数据库设计的原则和方法论...</p>
<p>第二段：概念模型和逻辑模型设计...</p>
<p>第三段：物理模型和表结构设计...</p>
<p>第四段：索引优化和查询性能设计...</p>
<p>第五段：数据安全和备份恢复设计...</p>"""

        elif "第4章" in section_name:
            prompt = f"""请为{paper_type}《{title}》生成第4章系统详细设计与实现。

研究领域：{field}
目标字数：{section_words}字

【严格格式要求】：
- 主标题使用<h2>标签，子标题使用<h3>标签
- 绝对禁止生成四级标题（如4.1.1、4.1.2等），最多只能到三级
- 内容必须是连贯的段落形式，严格禁止使用任何形式的分点列表

【统一章节结构】：
4.1 系统功能模块设计 - 详细设计各个功能模块的架构和交互机制
4.2 关键技术实现 - 深入阐述核心技术的具体实现方案和技术细节
4.3 系统安全设计 - 全面介绍系统的多层次安全保障体系

【严格写作要求】：
- 使用自然流畅的段落形式，严格避免任何分点列表（1.、2.、3.或•、-、①②③等）
- 每个技术点都要用连贯的叙述语言描述，不要写成列表形式
- 技术实现要融入到段落的自然表达中，避免生硬的技术堆砌
- 每个子节必须包含5-6个自然段落，每段250-300字，内容要详细深入
- 段落之间要有逻辑递进关系，形成完整的技术叙述
- 使用"在技术实现上"、"为了保证...本系统采用了"、"通过深入分析发现"等自然过渡语
- 必须包含具体的技术细节、代码实现思路、性能优化方案等深度内容

【严格禁止的错误示例】：
"认证体系包括：1）用户登录 2）权限验证 3）令牌管理"
"功能模块主要有：①用户管理模块 ②图书管理模块 ③借阅管理模块"
"安全措施包含：• 数据加密 • 访问控制 • 日志审计"
"系统采用以下技术：1）Spring Boot 2）MySQL 3）Redis"

【正确示例（必须参考）】：
"在认证体系的设计过程中，本系统深入分析了传统认证机制的不足，发现单纯的session管理在分布式环境下存在状态同步问题。为了解决这一技术难题，系统采用了基于JWT的无状态认证方案，通过在每个请求中携带包含用户信息的令牌来实现身份验证。这种设计不仅提高了系统的扩展性，还显著降低了服务器的内存负担。考虑到安全性要求，系统还实现了令牌的定期刷新机制，确保在保持用户体验的同时最大化系统安全性。在具体实现上，系统使用HMAC-SHA256算法对JWT进行签名，设置合理的过期时间，并在Redis中维护黑名单机制来处理用户主动登出的场景。"

请严格按照统一结构输出HTML格式内容：
<h2>第4章 系统详细设计与实现</h2>
<h3>4.1 系统功能模块设计</h3>
<p>第一段：模块设计理念和整体架构，详细阐述设计思路...</p>
<p>第二段：用户管理模块的详细设计和实现方案...</p>
<p>第三段：图书管理模块的核心功能和技术实现...</p>
<p>第四段：借阅管理模块的业务流程和数据处理...</p>
<p>第五段：智能推荐模块的算法设计和实现细节...</p>
<p>第六段：模块间交互机制和接口设计...</p>
<h3>4.2 关键技术实现</h3>
<p>第一段：Spring Boot框架的深度应用和配置优化...</p>
<p>第二段：数据持久化技术的实现和性能优化...</p>
<p>第三段：缓存机制的设计和Redis集成方案...</p>
<p>第四段：搜索引擎的集成和全文检索实现...</p>
<p>第五段：前后端分离架构的API设计和实现...</p>
<h3>4.3 系统安全设计</h3>
<p>第一段：安全架构设计理念和多层防护体系...</p>
<p>第二段：身份认证和授权机制的详细实现...</p>
<p>第三段：数据安全和加密技术的应用...</p>
<p>第四段：系统监控和日志审计的实现方案...</p>
<p>第五段：安全测试和漏洞防护的具体措施...</p>"""

        elif "第5章" in section_name:
            prompt = f"""请为{paper_type}《{title}》生成第5章系统测试。

研究领域：{field}
目标字数：{section_words}字

【严格格式要求】：
- 主标题使用<h2>标签，子标题使用<h3>标签
- 绝对禁止生成四级标题（如5.1.1、5.2.1等），最多只能到三级
- 内容必须是连贯的段落形式，严格禁止使用任何形式的分点列表

【统一章节结构】：
5.1 测试环境与方法 - 详细描述测试环境配置和测试方法选择
5.2 单元测试与集成测试 - 进行系统功能的全面测试验证
5.3 系统测试 - 分析系统的功能、性能和安全性测试

【严格写作要求】：
- 使用自然流畅的段落形式，严格避免任何分点列表（1.、2.、3.或•、-、①②③等）
- 测试过程要用连贯的叙述语言描述测试方法和结果，不要写成测试用例清单
- 每个测试环节都要融入到段落的自然表达中，避免生硬的测试数据堆砌
- 每个子节必须包含5-6个自然段落，每段250-300字，内容要详细深入
- 段落之间要有逻辑递进关系，从测试准备到结果分析
- 使用"在测试过程中发现"、"通过全面测试验证"、"测试结果表明"等自然过渡语
- 必须包含具体的测试数据、性能指标、覆盖率统计等详细信息

【严格禁止的错误示例】：
"功能测试包括：1）登录功能测试 2）数据处理测试 3）报表生成测试"
"测试工具主要有：①JUnit ②Selenium ③JMeter"
"性能指标包含：• 响应时间 • 并发用户数 • 吞吐量"
"测试用例分为：1）正常流程测试 2）异常流程测试 3）边界值测试"

【正确示例（必须参考）】：
"为了全面验证系统的功能完整性和稳定性，本研究设计了一套系统性的测试方案。在功能测试阶段，首先对系统的核心业务流程进行了深入的验证，包括用户认证机制、数据处理逻辑以及界面交互响应等关键环节。测试过程中采用了黑盒测试和白盒测试相结合的策略，既验证了系统功能的正确性，又深入检查了代码逻辑的合理性。通过大量真实场景的模拟测试，发现系统在各种复杂条件下都能保持良好的运行状态。测试结果表明，系统的核心功能模块在正常操作流程下的成功率达到99.8%，异常处理机制能够有效捕获和处理各种边界情况。"

请严格按照统一结构输出HTML格式内容：
<h2>第5章 系统测试</h2>
<h3>5.1 测试环境与方法</h3>
<p>第一段：测试环境的搭建和配置详细过程...</p>
<p>第二段：测试方法论的选择和测试策略制定...</p>
<p>第三段：测试工具的选型和集成配置...</p>
<p>第四段：测试数据的准备和测试用例设计...</p>
<p>第五段：测试流程的规范化和质量控制...</p>
<h3>5.2 单元测试与集成测试</h3>
<p>第一段：单元测试的设计理念和实施方案...</p>
<p>第二段：核心业务逻辑的单元测试覆盖...</p>
<p>第三段：集成测试的策略和模块间接口验证...</p>
<p>第四段：数据库集成测试和事务一致性验证...</p>
<p>第五段：API接口测试和前后端集成验证...</p>
<p>第六段：测试覆盖率分析和问题修复过程...</p>
<h3>5.3 系统测试</h3>
<p>第一段：功能测试的全面执行和结果分析...</p>
<p>第二段：性能测试的设计和压力测试结果...</p>
<p>第三段：安全性测试和漏洞扫描验证...</p>
<p>第四段：兼容性测试和用户体验验证...</p>
<p>第五段：系统稳定性测试和长期运行验证...</p>"""

        elif "第6章" in section_name:
            prompt = f"""请为{paper_type}《{title}》生成第6章总结与展望。

研究领域：{field}
目标字数：{section_words}字

【严格格式要求】：
- 主标题使用<h2>标签，子标题使用<h3>标签
- 绝对禁止生成四级标题，最多只能到三级
- 绝对禁止在段落中生成任何小标题（如"系统性能与可靠性需求"等）
- 内容必须是纯段落形式，严格禁止使用任何形式的分点列表

【严格章节要求】：
必须按照6.1、6.2、6.3、6.4的顺序生成，不能出现章节编号错乱

【统一章节结构】：
6.1 工作总结 - 总结本研究的主要工作和成果
6.2 系统特色与创新 - 归纳本研究的创新点和技术特色
6.3 不足与改进 - 分析研究的不足之处和改进方向
6.4 未来展望 - 展望后续研究方向和应用前景

【严格写作要求】：
- 使用自然流畅的段落形式，严格避免任何分点列表（1.、2.、3.或•、-、①②③等）
- 绝对禁止在段落内部生成小标题或子标题
- 工作总结要用连贯的叙述语言回顾研究历程，不要写成成果清单
- 创新点和不足要融入到段落的自然表达中，避免生硬的要点罗列
- 每个子节必须包含4-5个自然段落，每段250-300字，内容要详细深入
- 段落之间要有逻辑递进关系，从现状总结到未来展望
- 使用"回顾整个研究过程"、"通过深入研究发现"、"展望未来发展趋势"等自然过渡语

【严格禁止的错误示例】：
"本研究的创新点包括：1）提出了新算法 2）设计了新架构 3）实现了新功能"
"系统特色主要有：①智能推荐 ②高性能 ③安全可靠"
"不足之处包括：• 性能优化 • 功能扩展 • 用户体验"
"系统性能与可靠性需求"（禁止这种小标题）
"安全性与数据完整性需求"（禁止这种小标题）

【正确示例（必须参考）】：
"回顾整个研究过程，本论文在{field}领域取得了一系列有价值的研究成果。在理论层面，通过深入分析现有技术的局限性，提出了一种全新的解决思路，这种方法不仅在理论上具有创新性，更重要的是为实际应用提供了可行的技术路径。在实践层面，本研究成功开发了完整的系统原型，并通过大量实验验证了方法的有效性和实用性，为相关领域的技术发展奠定了坚实的基础。通过系统化的研究方法和严谨的实验验证，本研究不仅解决了传统系统存在的关键问题，还为后续研究提供了重要的理论基础和实践经验。"

请严格按照顺序输出HTML格式内容：
<h2>第6章 总结与展望</h2>
<h3>6.1 工作总结</h3>
<p>第一段：研究过程回顾和主要成果...</p>
<p>第二段：理论贡献和技术突破...</p>
<p>第三段：系统实现和验证结果...</p>
<p>第四段：研究价值和学术意义...</p>
<h3>6.2 系统特色与创新</h3>
<p>第一段：技术架构创新和设计理念...</p>
<p>第二段：核心算法创新和性能提升...</p>
<p>第三段：功能特色和用户体验创新...</p>
<p>第四段：整体创新价值和竞争优势...</p>
<h3>6.3 不足与改进</h3>
<p>第一段：当前系统存在的技术局限...</p>
<p>第二段：性能和功能方面的不足...</p>
<p>第三段：用户体验和扩展性问题...</p>
<p>第四段：改进方案和优化策略...</p>
<h3>6.4 未来展望</h3>
<p>第一段：技术发展趋势和应用前景...</p>
<p>第二段：功能扩展和性能优化方向...</p>
<p>第三段：应用场景拓展和生态建设...</p>
<p>第四段：长远发展规划和研究价值...</p>"""

        else:
            # 通用章节提示词 - 严格控制标题层级，生成自然流畅内容
            if level == 3:
                # 三级标题章节 - 严格禁止生成四级标题
                prompt = f"""请为{paper_type}《{title}》生成{section_name}章节。

研究领域：{field}
目标字数：{section_words}字
章节描述：{section_desc}
标题级别：{level}级标题

【严格格式要求】：
- 这是一个三级标题章节，只能使用<h3>标签作为主标题
- 绝对禁止生成四级标题（如2.1.1、2.1.2等），最多只能到三级
- 内容必须是连贯的段落形式，不能使用分点列表
- 如果内容较多，通过自然段落分隔，而不是子标题

【重要写作要求】：
- 使用自然流畅的段落形式，严格避免分点列表（1.、2.、3.或•、-等）
- 每个技术点都要融入到段落的自然表达中，避免生硬的要点罗列
- 必须包含4-6个自然段落，每段200-250字
- 段落之间要有逻辑递进关系，形成完整的学术论述
- 使用"在...方面"、"通过深入分析"、"考虑到...需求"等自然过渡语

错误示例（严格避免）：
"技术特点包括：1）高性能 2）易扩展 3）安全可靠"

正确示例（参考）：
"在技术选型过程中，本系统深入分析了多种解决方案的优劣势。通过对比研究发现，当前主流的技术框架在性能表现和扩展能力方面都有显著提升，特别是在处理大规模并发请求时展现出了良好的稳定性。考虑到项目的实际需求和长期发展规划，最终选择了成熟度较高且社区支持活跃的技术栈，这不仅保证了开发效率，也为后续的维护和升级奠定了坚实基础。"

请输出HTML格式内容：
<h3>{section_name}</h3>
<p>第一段详细内容...</p>
<p>第二段详细内容...</p>
<p>第三段详细内容...</p>"""
            else:
                # 二级标题章节 - 可以包含三级子标题
                prompt = f"""请为{paper_type}《{title}》生成{section_name}章节。

研究领域：{field}
目标字数：{section_words}字
章节描述：{section_desc}

【严格格式要求】：
- 主标题使用<h2>标签
- 如果需要子节，最多使用<h3>标签，绝对禁止四级标题
- 内容必须是连贯的段落形式，不能使用分点列表

【重要写作要求】：
- 使用自然流畅的段落形式，严格避免分点列表（1.、2.、3.或•、-等）
- 每个子节包含4-5个自然段落，每段200-250字
- 段落之间要有逻辑递进关系，形成完整的学术论述
- 使用"在...方面"、"通过深入分析"、"考虑到...需求"等自然过渡语

请输出HTML格式内容：
<h2>{section_name}</h2>
<h3>子节标题</h3>
<p>详细段落内容...</p>"""

        # 调用API生成内容
        content = call_deepseek_api(prompt, min(section_words * 3, 6000))
        
        # 清理内容并处理文献引用
        if content:
            cleaned_content = clean_ai_generated_content(content)
            
            # 处理文献引用收集
            if memory and 'collected_references' in memory:
                import re
                # 提取临时引用
                temp_refs_match = re.search(r'<div class="temp-references">(.*?)</div>', cleaned_content, re.DOTALL)
                if temp_refs_match:
                    refs_text = temp_refs_match.group(1).strip()
                    # 解析引用并添加到收集器
                    ref_lines = [line.strip() for line in refs_text.split('\n') if line.strip() and line.strip().startswith('[')]
                    ref_start_num = memory['reference_counter'] + 1
                    
                    for i, ref_line in enumerate(ref_lines):
                        memory['reference_counter'] += 1
                        # 重新编号引用
                        new_ref = re.sub(r'^\[\d+\]', f'[{memory["reference_counter"]}]', ref_line)
                        memory['collected_references'].append(new_ref)
                    
                    # 从内容中移除临时引用部分
                    cleaned_content = re.sub(r'<div class="temp-references">.*?</div>', '', cleaned_content, flags=re.DOTALL).strip()
                    
                    # 更新内容中的引用编号
                    current_refs = re.findall(r'\[(\d+)\]', cleaned_content)
                    if current_refs and ref_lines:
                        # 创建引用映射
                        ref_mapping = {}
                        for i, old_num in enumerate(sorted(set(map(int, current_refs)))):
                            ref_mapping[old_num] = ref_start_num + i
                        
                        # 替换引用编号
                        for old_num, new_num in ref_mapping.items():
                            cleaned_content = re.sub(f'\\[{old_num}\\]', f'[{new_num}]', cleaned_content)
            
            # 验证和调整内容长度
            cleaned_content = validate_and_adjust_content_length(cleaned_content, section_words, section_name)
            
            return cleaned_content
        else:
            return generate_fallback_section(section, section_num)
        
    except Exception as e:
        app.logger.error(f"生成章节 {section_name} 失败: {e}")
        return generate_fallback_section(section, section_num)


def search_academic_literature(field, title, keywords):
    """联网搜索学术文献"""
    try:
        # 构建文献搜索提示词
        search_prompt = f"""作为学术文献搜索专家，请联网搜索{field}领域与《{title}》相关的真实学术文献。

搜索关键词：{keywords}

请搜索以下类型的真实文献：
1. 中文核心期刊论文（知网、万方、维普）
2. 国际期刊论文（IEEE、ACM、Springer、Elsevier）
3. 重要会议论文
4. 权威学位论文
5. 专业书籍和专著

搜索要求：
- 优先搜索2019-2024年的最新文献
- 确保文献真实存在且可验证
- 包含准确的作者、标题、期刊、年份、页码等信息
- 涵盖理论基础、技术方法、应用实践等多个方面

请返回JSON格式的搜索结果：
{{
    "literature_list": [
        {{
            "id": 1,
            "title": "文献标题",
            "authors": ["作者1", "作者2"],
            "journal": "期刊名称",
            "year": 2023,
            "volume": "卷号",
            "issue": "期号",
            "pages": "页码范围",
            "type": "期刊论文/会议论文/学位论文/专著",
            "doi": "DOI号（如有）",
            "relevance": "与论文的相关性描述"
        }}
    ]
}}"""

        app.logger.info(f"开始联网搜索{field}领域的学术文献")

        # 调用联网搜索API
        search_result = call_deepseek_api_with_search(search_prompt, 4000)

        if search_result:
            try:
                # 尝试解析JSON结果
                import json
                start_idx = search_result.find('{')
                end_idx = search_result.rfind('}') + 1
                if start_idx != -1 and end_idx != 0:
                    json_str = search_result[start_idx:end_idx]
                    literature_data = json.loads(json_str)
                    return literature_data.get('literature_list', [])
            except json.JSONDecodeError:
                app.logger.warning("文献搜索结果不是有效JSON格式")
                return []

        return []

    except Exception as e:
        app.logger.error(f"搜索学术文献失败: {e}")
        return []


def generate_references_with_search(field, title, keywords):
    """基于联网搜索生成参考文献"""
    try:
        current_year = datetime.now().year

        # 首先搜索真实文献
        literature_list = search_academic_literature(field, title, keywords)

        if literature_list and len(literature_list) >= 10:
            # 基于搜索结果生成参考文献
            prompt = f"""基于以下搜索到的真实学术文献，为{field}领域的论文《{title}》生成标准格式的参考文献列表。

搜索到的文献信息：
{json.dumps(literature_list, ensure_ascii=False, indent=2)}

请按照GB/T 7714-2015标准格式整理这些文献：
- 期刊论文：[序号] 作者. 论文标题[J]. 期刊名称, 年份, 卷号(期号): 页码.
- 会议论文：[序号] 作者. 论文标题[C]. 会议名称, 年份: 页码.
- 学位论文：[序号] 作者. 论文标题[D]. 学校名称, 年份.
- 专著：[序号] 作者. 书名[M]. 出版社, 年份.

要求：
1. 严格按照搜索结果的真实信息整理
2. 确保格式规范统一
3. 按照相关性和重要性排序
4. 如果文献不足15条，可以适当补充相关的权威文献

请直接输出HTML格式：
<h2>参考文献</h2>
<p>[1] ...</p>
<p>[2] ...</p>
..."""

            # 调用AI生成格式化的参考文献
            formatted_refs = call_deepseek_api(prompt, 3000)
            if formatted_refs and len(formatted_refs.strip()) > 200:
                return formatted_refs, literature_list

        # 如果搜索失败，使用AI生成备用文献
        return generate_ai_only_references(field, title), []

    except Exception as e:
        app.logger.error(f"基于搜索生成参考文献失败: {e}")
        return generate_ai_only_references(field, title), []


def generate_collected_references(memory):
    """基于收集的引用生成参考文献章节 - 改进版本"""
    try:
        if not memory or 'collected_references' not in memory or not memory['collected_references']:
            app.logger.warning("没有收集到文献引用，生成默认参考文献")
            return generate_default_references()
        
        references = memory['collected_references']
        
        # 按类型分组并排序
        journal_refs = []
        conference_refs = []
        book_refs = []
        other_refs = []
        
        for ref in references:
            if '[J]' in ref:
                journal_refs.append(ref)
            elif '[C]' in ref:
                conference_refs.append(ref)
            elif '[M]' in ref:
                book_refs.append(ref)
            else:
                other_refs.append(ref)
        
        # 重新编号，确保连续性
        all_refs = journal_refs + conference_refs + book_refs + other_refs
        
        # 生成HTML
        references_html = '<h2>参考文献</h2>\n'
        references_html += '<div class="references-container">\n'
        
        for i, ref in enumerate(all_refs, 1):
            # 重新编号引用
            formatted_ref = re.sub(r'^\[\d+\]', f'[{i}]', ref)
            references_html += f'<p class="reference-item">{formatted_ref}</p>\n'
        
        references_html += '</div>\n'
        
        # 添加引用统计信息
        references_html += f'''
<div class="reference-stats">
    <div class="stats-item">
        <i class="fas fa-book"></i>
        <span>期刊论文: {len(journal_refs)}</span>
    </div>
    <div class="stats-item">
        <i class="fas fa-users"></i>
        <span>会议论文: {len(conference_refs)}</span>
    </div>
    <div class="stats-item">
        <i class="fas fa-bookmark"></i>
        <span>专著: {len(book_refs)}</span>
    </div>
    <div class="stats-item">
        <i class="fas fa-globe"></i>
        <span>其他: {len(other_refs)}</span>
    </div>
</div>

<style>
.references-container {
    padding: 20px 0;
    line-height: 1.8;
}

.reference-item {
    margin-bottom: 12px;
    padding-left: 20px;
    text-indent: -20px;
    font-size: 14px;
    color: #2d3748;
}

.reference-stats {
    display: flex;
    justify-content: space-around;
    flex-wrap: wrap;
    gap: 15px;
    margin-top: 20px;
    padding: 15px;
    background: linear-gradient(135deg, #f8f9fa 0%, #e9ecef 100%);
    border-radius: 10px;
    border: 1px solid #dee2e6;
}

.stats-item {
    display: flex;
    align-items: center;
    gap: 8px;
    font-size: 13px;
    color: #495057;
    font-weight: 500;
}

.stats-item i {
    color: #667eea;
    font-size: 14px;
}
</style>
'''
        
        app.logger.info(f"基于收集的{len(all_refs)}条引用生成参考文献，包含{len(journal_refs)}篇期刊论文")
        return references_html
        
    except Exception as e:
        app.logger.error(f"生成收集的参考文献失败: {e}")
        return generate_default_references()


def generate_default_references():
    """生成默认参考文献"""
    return """<h2>参考文献</h2>
<div style="color: #d32f2f; background: #ffebee; border: 1px solid #f8bbd9; padding: 15px; border-radius: 5px; margin: 10px 0;">
<h3>⚠️ 参考文献待完善</h3>
<p><strong>提示：</strong>系统未能收集到足够的文献引用。请手动添加真实的参考文献。</p>
<p><strong>建议：</strong>可到以下数据库搜索相关文献：</p>
<ul>
<li>中国知网(CNKI) - <a href="https://cnki.net" target="_blank">cnki.net</a></li>
<li>IEEE Xplore - <a href="https://ieeexplore.ieee.org" target="_blank">ieeexplore.ieee.org</a></li>
<li>ACM Digital Library - <a href="https://dl.acm.org" target="_blank">dl.acm.org</a></li>
</ul>
</div>"""
def calculate_word_distribution(total_words, section_count):
    """智能计算各章节字数分配 - 改进版本"""
    if section_count <= 0:
        return {}
    
    # 基础字数分配比例 (更精确的分配)
    base_ratios = {
        '摘要': 0.06,          # 6% - 约720字
        'Abstract': 0.06,      # 6% - 约720字
        '第1章 绪论': 0.20,    # 20% - 约2400字 (增加权重)
        '第2章 相关技术介绍': 0.14,  # 14% - 约1680字
        '第3章 需求分析与系统设计': 0.18,  # 18% - 约2160字
        '第4章 系统详细设计与实现': 0.22,  # 22% - 约2640字 (最重要章节)
        '第5章 系统测试': 0.10,  # 10% - 约1200字
        '第6章 总结与展望': 0.04,  # 4% - 约480字
        '参考文献': 0.0        # 固定500字，不占比例
    }
    
    # 计算分配字数
    word_distribution = {}
    reference_words = 500  # 参考文献固定字数
    available_words = total_words - reference_words
    
    for section_name, ratio in base_ratios.items():
        if section_name == '参考文献':
            word_distribution[section_name] = reference_words
        else:
            # 按比例分配剩余字数
            allocated_words = int(available_words * ratio)
            # 确保每个章节至少有最小字数
            min_words = 300 if '摘要' in section_name else 800
            word_distribution[section_name] = max(allocated_words, min_words)
    
    # 验证总字数不超标
    total_allocated = sum(word_distribution.values())
    if total_allocated > total_words:
        # 按比例缩减
        scale_factor = total_words / total_allocated
        for section_name in word_distribution:
            if section_name != '参考文献':
                word_distribution[section_name] = int(word_distribution[section_name] * scale_factor)
    
    return word_distribution


def validate_and_adjust_content_length(content, target_words, section_name):
    """验证和调整内容长度"""
    import re
    
    if not content:
        return content
    
    # 统计实际字数（去除HTML标签）
    text_content = re.sub(r'<[^>]+>', '', content)
    actual_words = len(text_content.replace(' ', '').replace('\n', ''))
    
    # 计算达成率
    completion_rate = actual_words / target_words if target_words > 0 else 0
    
    app.logger.info(f"章节 {section_name}: 目标{target_words}字，实际{actual_words}字，达成率{completion_rate:.1%}")
    
    # 如果字数明显不足（小于70%），添加补充提示
    if completion_rate < 0.7 and actual_words < target_words - 200:
        supplement = f"""
<div style="background: #fff3cd; border: 1px solid #ffeaa7; padding: 10px; margin: 10px 0; border-radius: 4px;">
<p><strong>📝 内容扩展提示:</strong> 本章节当前字数为{actual_words}字，建议扩展到{target_words}字。可以考虑：</p>
<ul>
<li>增加技术细节和实现原理的深入分析</li>
<li>补充相关案例研究和对比分析</li>
<li>扩展对未来发展趋势的讨论</li>
<li>添加更多的技术背景和理论基础</li>
</ul>
</div>"""
        content += supplement
    
    return content

def generate_ai_only_references(field, title):
    """纯AI生成参考文献 - 备用方案"""
    try:
        current_year = datetime.now().year

        # 构建专门的AI参考文献生成提示词
        prompt = f"""作为学术文献专家，请为{field}领域的论文《{title}》生成15条高质量的真实参考文献。

要求：
1. 必须生成真实存在的学术文献，包含准确的作者、期刊、年份信息
2. 优先选择近5年内的高影响因子期刊论文
3. 包含中文核心期刊（如计算机学报、软件学报等）和国际顶级期刊（如IEEE、ACM等）
4. 包含期刊论文、会议论文、学位论文、专著等多种类型
5. 严格按照学术引用格式
6. 作者姓名要多样化，避免重复使用相同的人名
7. 每次生成的文献都应该不同，体现真实的学术多样性

文献类型分布建议：
- 期刊论文：8-10篇
- 会议论文：2-3篇
- 学位论文：1-2篇
- 专著教材：1-2篇

格式要求：
- 期刊论文：[序号] 作者. 论文标题[J]. 期刊名称, 年份, 卷号(期号): 页码.
- 会议论文：[序号] 作者. 论文标题[C]. 会议名称, 年份: 页码.
- 学位论文：[序号] 作者. 论文标题[D]. 学校名称, 年份.
- 专著：[序号] 作者. 书名[M]. 出版社, 年份.

请直接输出HTML格式：
<h2>参考文献</h2>
<p>[1] ...</p>
<p>[2] ...</p>
...

注意：请确保每次生成的参考文献都有所不同，作者姓名要多样化，避免使用模板化的内容。"""

        app.logger.info(f"开始AI生成{field}领域的参考文献")

        # 调用DeepSeek API生成参考文献
        ai_references = call_deepseek_api(prompt, 3000)

        if ai_references and len(ai_references.strip()) > 200:
            # 清理AI生成的内容
            cleaned_refs = clean_ai_generated_content(ai_references)
            if cleaned_refs and ("<h2>参考文献</h2>" in cleaned_refs or "[1]" in cleaned_refs):
                app.logger.info("成功通过AI生成参考文献")
                # 确保格式正确
                if not cleaned_refs.startswith("<h2>参考文献</h2>"):
                    cleaned_refs = "<h2>参考文献</h2>\n" + cleaned_refs
                return cleaned_refs
            else:
                app.logger.error("AI生成的参考文献格式不正确")
                return None
        else:
            app.logger.error("AI生成参考文献失败或内容过短")
            return None

    except Exception as e:
        app.logger.error(f"AI生成参考文献时出错: {e}")
        return None


def generate_advanced_references_with_search(field, title, use_search=False):
    """高级参考文献生成 - 可选择启用网络搜索获取真实文献"""
    try:
        current_year = datetime.now().year
        
        if use_search:
            # 使用AI + 网络搜索生成真实参考文献
            search_prompt = f"""作为学术文献专家，请为{field}领域的研究生成15条高质量的真实参考文献。

研究主题：{title}
研究领域：{field}

要求：
1. 生成的文献必须是真实存在的，包含准确的作者、期刊、年份信息
2. 优先选择近5年内的高影响因子期刊论文
3. 包含中文核心期刊（如计算机学报、软件学报等）和国际顶级期刊（如IEEE、ACM等）
4. 包含经典教材和学位论文
5. 严格按照学术引用格式

请搜索并提供真实的参考文献，格式如下：
[序号] 作者. 论文标题[J]. 期刊名称, 年份, 卷号(期号): 页码.

如果无法搜索到足够的真实文献，请基于该领域的知名学者和权威期刊生成高质量的参考文献。"""
            
            try:
                # 尝试使用AI生成更真实的参考文献
                real_refs = call_deepseek_api(search_prompt, 2000)
                if real_refs and len(real_refs.strip()) > 200:
                    # 清理AI生成的参考文献
                    cleaned_refs = clean_ai_generated_content(real_refs)
                    if cleaned_refs:
                        app.logger.info("成功生成AI增强的真实参考文献")
                        return cleaned_refs
            except Exception as e:
                app.logger.warning(f"AI搜索参考文献失败: {e}")
        
        # 不使用任何备用方案，直接返回None
        return None

    except Exception as e:
        app.logger.error(f"高级参考文献生成失败: {e}")
        return None

def generate_real_references_with_search(field, title):
    """使用DeepSeek联网搜索生成真实参考文献"""
    try:
        current_year = datetime.now().year
        
        # 构建专门的联网搜索提示词
        search_prompt = f"""我需要为{field}领域的学术论文生成真实的参考文献。请帮我联网搜索相关的学术资料。

论文信息：
- 研究领域：{field}
- 论文标题：{title}
- 目标：生成15条真实可靠的参考文献

请执行以下搜索任务：
1. 搜索知网(CNKI)数据库中与"{field}"相关的近期论文
2. 搜索IEEE Xplore中的相关技术文献
3. 查找ACM Digital Library中的会议论文
4. 搜索谷歌学术中的权威期刊文章

搜索关键词建议：
- {field}
- {title.split('基于')[1] if '基于' in title else title}
- 相关的技术术语

请返回真实的参考文献，包含：
- 真实的作者姓名
- 真实的论文标题  
- 真实的期刊/会议名称
- 准确的发表年份和页码

输出格式：
<h2>参考文献</h2>
<p>[1] 作者. 论文标题[J]. 期刊名称, 年份, 卷(期): 页码.</p>
<p>[2] 作者. 书名[M]. 出版社, 年份.</p>
...

注意：请确保所有文献都是通过网络搜索获得的真实数据，不要编造任何信息。"""

        app.logger.info(f"开始联网搜索{field}领域的真实参考文献")
        
        # 调用DeepSeek API进行联网搜索
        real_references = call_deepseek_api_with_search(search_prompt, 3000)
        
        if real_references and len(real_references.strip()) > 200:
            # 清理搜索结果
            cleaned_refs = clean_ai_generated_content(real_references)
            if cleaned_refs and ("<h2>参考文献</h2>" in cleaned_refs or "[1]" in cleaned_refs):
                app.logger.info("成功通过联网搜索获取真实参考文献")
                # 确保格式正确
                if not cleaned_refs.startswith("<h2>参考文献</h2>"):
                    cleaned_refs = "<h2>参考文献</h2>\n" + cleaned_refs
                return cleaned_refs
            else:
                app.logger.warning("搜索结果格式不正确，尝试重新请求")
                # 尝试更简单的搜索提示词
                simple_prompt = f"""请搜索{field}领域的真实学术文献，生成15条参考文献。要求：
1. 搜索真实的学术数据库
2. 返回真实的作者、期刊、年份信息
3. 格式：[1] 作者. 标题[J]. 期刊, 年份, 卷(期): 页码.

领域：{field}
题目：{title}

请开始搜索并返回结果："""
                
                retry_refs = call_deepseek_api_with_search(simple_prompt, 2000)
                if retry_refs and len(retry_refs.strip()) > 100:
                    cleaned_retry = clean_ai_generated_content(retry_refs)
                    if cleaned_retry:
                        app.logger.info("重试搜索成功")
                        return f"<h2>参考文献</h2>\n{cleaned_retry}"
        else:
            app.logger.warning("联网搜索返回结果为空")
            
    except Exception as e:
        app.logger.error(f"联网搜索参考文献失败: {e}")
    
    # 不使用任何备用方案
    app.logger.error("所有搜索尝试都失败，无法生成参考文献")
    return None


# 删除了generate_reference_template函数 - 不再使用静态模板


def call_deepseek_api_with_search(prompt, max_tokens=3000):
    """调用DeepSeek API进行联网搜索"""
    import time
    
    try:
        headers = {
            'Authorization': f'Bearer {DEEPSEEK_API_KEY}',
            'Content-Type': 'application/json'
        }

        # 简化配置，依靠提示词指导模型进行搜索
        payload = {
            'model': 'deepseek-chat',
            'messages': [
                {
                    'role': 'system',
                    'content': '你是一个专业的学术研究助手，具备联网搜索能力。请根据用户要求搜索真实的学术文献和资料。'
                },
                {
                    'role': 'user', 
                    'content': prompt
                }
            ],
            'temperature': 0.3,  # 降低温度以获得更准确的搜索结果
            'max_tokens': max_tokens,
            'stream': False
        }

        app.logger.info(f"发起DeepSeek联网搜索请求，提示词长度: {len(prompt)}")
        
        # 联网搜索通常需要更长时间
        response = requests.post(DEEPSEEK_API_URL, headers=headers, json=payload, timeout=180)
        
        if response.status_code == 200:
            result = response.json()
            content = result['choices'][0]['message']['content']
            
            app.logger.info(f"联网搜索API调用成功，返回内容长度: {len(content)}")
            app.logger.debug(f"搜索返回内容预览: {content[:200]}...")
            return content
            
        elif response.status_code == 429:
            app.logger.warning("API频率限制，等待后重试")
            time.sleep(5)
            # 重试一次
            response = requests.post(DEEPSEEK_API_URL, headers=headers, json=payload, timeout=180)
            if response.status_code == 200:
                result = response.json()
                return result['choices'][0]['message']['content']
        
        app.logger.error(f"联网搜索API调用失败: {response.status_code} - {response.text}")
        return None
        
    except requests.Timeout:
        app.logger.error("联网搜索请求超时")
        return None
    except Exception as e:
        app.logger.error(f"联网搜索API调用异常: {e}")
        return None


# 删除了generate_enhanced_static_references函数 - 不再使用静态参考文献


def generate_fallback_section(section, section_num):
    """生成备用章节内容"""
    section_name = section['name']
    section_desc = section['description']
    
    if section_name == "摘要":
        return f"""<h2>摘要</h2>
<p>本研究针对当前{section_desc}中存在的关键问题进行深入分析和研究。通过采用先进的技术方法和理论框架，本文提出了创新的解决方案。研究结果表明，所提出的方法在理论和实践层面都具有重要意义。</p>
<p><strong>关键词：</strong>系统设计；技术实现；性能优化；创新方法</p>"""
        
    elif section_name == "Abstract":
        return f"""<h2>Abstract</h2>
<p>This research addresses key issues in {section_desc}. Through advanced technical methods and theoretical frameworks, this paper proposes innovative solutions. The results demonstrate significant theoretical and practical importance.</p>
<p><strong>Keywords:</strong> system design; technical implementation; performance optimization; innovative methods</p>"""
        
    elif "第" in section_name and "章" in section_name:
        chapter_num = section_name.split("第")[1].split("章")[0]
        return f"""<h2>第{chapter_num}章 {section_name.split(' ', 1)[1] if ' ' in section_name else section_desc}</h2>
<p>本章主要讨论{section_desc}相关内容。通过详细的分析和研究，为后续章节奠定理论基础。</p>
<p>相关技术和方法的应用将在本章中得到充分的阐述和说明。</p>"""
        
    else:
        return f"""<h2>{section_name}</h2>
<p>本节介绍{section_desc}的相关内容。通过系统性的分析，为研究提供必要的支撑。</p>"""
def call_deepseek_api(prompt, max_tokens=3000):
    """调用DeepSeek API - 高质量版本，只返回真实AI内容"""
    import time
    
    try:
        headers = {
            'Authorization': f'Bearer {DEEPSEEK_API_KEY}',
            'Content-Type': 'application/json'
        }

        payload = {
            'model': 'deepseek-chat',
            'messages': [{'role': 'user', 'content': prompt}],
            'temperature': 0.7,
            'max_tokens': min(max_tokens, 8000),  # 确保在API限制内
            'stream': False
        }

        # 重试机制：确保获得真实内容
        for attempt in range(3):
            try:
                app.logger.info(f"API调用尝试 {attempt + 1}/3，max_tokens: {max_tokens}")
                
                response = requests.post(DEEPSEEK_API_URL, headers=headers, json=payload, timeout=120)
                
                if response.status_code == 200:
                    result = response.json()
                    content = result['choices'][0]['message']['content']
                    
                    # 验证内容质量 - 必须是真实的AI生成内容
                    if content and len(content.strip()) > 200:  # 确保内容充实
                        app.logger.info(f"API调用成功，返回内容长度: {len(content)}")
                        return content
                    else:
                        app.logger.warning(f"API返回内容过短，尝试重新生成")
                        if attempt < 2:
                            time.sleep(2)
                            continue
                        
                elif response.status_code == 429:  # 频率限制
                    wait_time = 2 ** attempt
                    app.logger.warning(f"API频率限制，等待 {wait_time} 秒后重试")
                    time.sleep(wait_time)
                    continue
                    
                else:
                    app.logger.warning(f"API调用失败: {response.status_code} - {response.text}")
                    if attempt < 2:
                        time.sleep(3)
                        continue
                    
            except requests.Timeout:
                app.logger.warning(f"API调用超时 - 尝试 {attempt + 1}/3")
                if attempt < 2:
                    time.sleep(2)
                    continue
                    
            except Exception as e:
                app.logger.error(f"API调用异常: {e} - 尝试 {attempt + 1}/3")
                if attempt < 2:
                    time.sleep(3)
                    continue
        
        # 所有重试都失败 - 抛出异常，不使用备用内容
        raise Exception("API调用失败，无法生成内容。请检查网络连接或稍后重试。")
        
    except Exception as e:
        app.logger.error(f"DeepSeek API调用严重错误: {e}")
        raise e  # 向上抛出异常，不返回备用内容


def calculate_optimal_tokens(section_name, target_words, context_length=0):
    """
    动态计算最优token分配，支持长论文生成
    
    参数:
    - section_name: 章节名称
    - target_words: 目标字数
    - context_length: 上下文长度
    
    返回:
    - 最优的max_tokens值
    """
    # DeepSeek API的实际限制（实测安全值）
    MAX_CONTEXT_TOKENS = 65000  # 留出安全边界
    MAX_OUTPUT_TOKENS = 8000    # 单次输出安全上限，极大幅提高以满足字数要求
    
    # 估算中文token比例（中文字符通常1字符约等于1.5-2个token）
    estimated_output_tokens = int(target_words * 2)
    
    # 考虑上下文占用
    available_tokens = MAX_CONTEXT_TOKENS - context_length
    
    # 根据章节类型动态调整 - 极大幅增加token分配以满足用户字数要求
    section_multipliers = {
        "摘要": 2.0,           # 摘要需要更详细内容
        "第1章 绪论": 3.5,           # 需要非常详细的背景介绍
        "第2章 相关技术介绍": 4.0,       # 文献综述需要大量内容
        "第3章 需求分析与系统设计": 4.5,       # 技术细节极多
        "第4章 系统详细设计与实现": 4.2,       # 数据分析极其详细
        "第5章 系统测试": 4.0,       # 深度分析需要大量内容
        "第6章 总结与展望": 2.8,           # 总结性内容也要极其详细
    }
    
    multiplier = section_multipliers.get(section_name, 1.2)
    optimal_tokens = int(estimated_output_tokens * multiplier)
    
    # 确保不超过各种限制
    final_tokens = min(
        optimal_tokens,
        MAX_OUTPUT_TOKENS,
        available_tokens - 1000,  # 预留安全边界
        target_words * 4          # 最大不超过字数的4倍token（增加倍数）
    )
    
    app.logger.info(f"Token计算 - 章节:{section_name}, 目标字数:{target_words}, 分配tokens:{final_tokens}")
    return max(final_tokens, 3000)  # 最少保证3000个token（大幅提高最小值以满足字数要求）


def split_text_intelligently(text, max_length=2000):
    """
    智能分段处理长文本
    
    参数:
    - text: 要分段的文本
    - max_length: 每段的最大长度
    
    返回:
    - 分段后的文本列表
    """
    if len(text) <= max_length:
        return [text]
    
    paragraphs = text.split('\n\n')
    segments = []
    current_segment = ""
    
    for paragraph in paragraphs:
        if len(current_segment + paragraph) <= max_length:
            current_segment += paragraph + '\n\n'
        else:
            if current_segment:
                segments.append(current_segment.strip())
                current_segment = paragraph + '\n\n'
            else:
                # 单个段落太长，按句子分割
                sentences = paragraph.split('。')
                temp_segment = ""
                for sentence in sentences:
                    if len(temp_segment + sentence + '。') <= max_length:
                        temp_segment += sentence + '。'
                    else:
                        if temp_segment:
                            segments.append(temp_segment)
                        temp_segment = sentence + '。'
                if temp_segment:
                    current_segment = temp_segment + '\n\n'
    
    if current_segment:
        segments.append(current_segment.strip())
    
    return segments


def optimize_text_with_deepseek(text_segment, context="", config=None):
    """
    使用DeepSeek优化单个文本片段 - 基于AI检测原理的深度优化

    核心原理：
    1. Perplexity（困惑度）：AI文本词汇选择高度可预测，需要增加词汇随机性
    2. Burstiness（突发性）：AI文本句式均匀，人类写作长短句交替、节奏变化
    3. 模式识别：AI有固定的连接词、句式结构，需要打破这些模式

    参数:
    - text_segment: 要优化的文本片段
    - context: 上下文信息
    - config: 优化配置

    返回:
    - 优化后的文本
    """
    if config is None:
        config = {
            'intensity': 3,
            'mode': 'balanced',
            'preserve_terms': True,
            'diversify_sentence': True,
            'natural_transition': True,
            'add_human_touch': True
        }

    intensity = config.get('intensity', 3)
    mode = config.get('mode', 'moderate')
    text_type = config.get('text_type', 'academic')
    preserve_terms = config.get('preserve_terms', True)
    diversify_sentence = config.get('diversify_sentence', True)
    natural_transition = config.get('natural_transition', False)
    add_human_touch = config.get('add_human_touch', False)
    keep_formal_style = config.get('keep_formal_style', True)

    # 新增的高级优化选项
    increase_perplexity = config.get('increase_perplexity', True)  # 增加困惑度
    add_burstiness = config.get('add_burstiness', True)  # 增加突发性
    break_patterns = config.get('break_patterns', True)  # 打破AI模式
    add_specificity = config.get('add_specificity', False)  # 增加具体细节
    vary_paragraph = config.get('vary_paragraph', True)  # 段落结构变化

    # 根据强度级别调整温度和改写程度描述
    intensity_settings = {
        1: {'temp': 0.6, 'degree': '轻微调整，仅修改最明显的AI痕迹，保持90%原文'},
        2: {'temp': 0.7, 'degree': '适度改写，保持较高的原文相似度，改写约40%内容'},
        3: {'temp': 0.8, 'degree': '中等强度改写，平衡原意保持和降AI率，改写约60%内容'},
        4: {'temp': 0.9, 'degree': '深度改写，大幅调整表达方式，改写约80%内容'},
        5: {'temp': 1.0, 'degree': '极限改写，最大限度重构表达，几乎完全重写'}
    }

    # 根据模式设置不同的优化策略 - 基于AI检测原理优化
    mode_strategies = {
        'moderate': """## 适度优化模式 - 学术论文专用（推荐）

### 基于AI检测原理的优化策略：

**1. 降低可预测性(Perplexity优化)**
- 替换30%的常用词为同义词或近义词
- 避免使用AI偏好的"标准化"表达
- 适当使用不那么"完美"的词语组合

**2. 增加句式突发性(Burstiness优化)**
- 在一个段落内混合使用：短句(10字以内)、中句(20字左右)、长句(40字以上)
- 不要让每个段落都是4-5个相似长度的句子
- 偶尔使用破折号、省略号制造停顿

**3. 打破AI模式**
- 消除"首先、其次、最后"的三段式
- 替换所有AI高频连接词
- 段落开头不总是主题句

**4. 保持学术规范**
- 专业术语、数据、引用保持原样
- 保持学术论文的正式风格""",

        'aggressive': """## 深度优化模式 - 最大化降AI率

### 基于AI检测原理的深度优化：

**1. 大幅降低可预测性**
- 替换50%以上的词汇
- 使用更多样化的同义表达
- 引入一些学术领域的专业变体表达

**2. 强化句式突发性**
- 刻意制造句子长度的剧烈变化
- 一个长句后跟2-3个短句，再来一个中等句
- 打破每段句数的规律性（有的段3句，有的段7句）

**3. 彻底打破AI模式**
- 删除60%的连接词，用句号直接分隔
- 调整论述顺序：可以先给结论再给论据
- 段落结构完全打散重组

**4. 增加人类写作痕迹**
- 适当加入"本研究认为"、"笔者发现"等主观表达
- 使用"相对而言"、"在一定程度上"等程度限定
- 偶尔使用括号补充说明""",

        'rewrite': """## 完全重写模式（激进）

### 完全基于人类写作特征重写：

**1. 最大化困惑度**
- 几乎完全更换词汇表达
- 使用更口语化但仍正式的表达方式
- 引入该领域人类作者常用的特色表达

**2. 模拟真实写作节奏**
- 像人类一样"思考着写"：有时详细，有时简略
- 制造"写作中的犹豫"：使用"或者说"、"换个角度看"
- 段落长度差异明显

**3. 完全消除AI痕迹**
- 不使用任何AI高频词汇
- 论述结构完全重组
- 加入适当的个人化学术表达"""
    }

    # 构建动态特征替换规则 - 基于AI检测研究更新
    ai_signature_rules = """
## AI特征词替换规则（必须严格执行）

### 知网/Turnitin等平台重点检测的AI特征：

**第一类：结构性标记词（检测权重最高，必须替换）**
| AI特征表达 | 替换方案（随机选一个或直接删除） |
|-----------|------------------------------|
| 首先/其次/再次/最后 | 完全删除，用句号分开；或只保留一个 |
| 第一/第二/第三 | 删除序号，直接陈述；或用"一是...二是" |
| 一方面...另一方面 | 从X角度看...就Y而言；或删除对仗 |
| 综上所述/总而言之 | 从以上分析来看/基于上述内容/回顾前文 |
| 因此/所以/故而 | 这使得/由此/这样一来（或直接删除） |

**第二类：AI偏好的过渡词（检测权重高）**
| AI特征表达 | 替换方案 |
|-----------|---------|
| 此外/另外/除此之外 | 还有/同时/而且（或删除，用句号） |
| 值得注意的是 | 要说明的是/这里需要指出/特别是 |
| 具体而言/具体来说 | 详细来看/说得细一点/展开来讲 |
| 与此同时 | 同时/在此期间/这一时期 |
| 在此基础上 | 基于此/由此/在这个基础上 |
| 从某种程度上说 | 可以说/在一定程度上/相对来说 |

**第三类：AI特有的"完美"表达（需要"弱化"）**
| AI特征表达 | 替换为更"人类"的表达 |
|-----------|-------------------|
| 极大地促进了 | 促进了/推动了/有利于 |
| 具有重要意义 | 有一定意义/有其价值/值得关注 |
| 发挥着重要作用 | 有一定作用/起到作用/产生影响 |
| 得到了广泛应用 | 应用较多/有所应用/被采用 |
| 取得了显著成效 | 取得成效/有所成效/产生效果 |
| 研究表明/研究发现 | 有研究指出/从研究来看/相关研究显示 |
| 不难发现/可以发现 | 能够看出/分析显示/观察到 |
| 显而易见 | 明显/显然/可以看出 |

**第四类：句式结构（必须打破）**
- "是...的"句式过多 → 减少使用，换成直接陈述
- 每句话都很完整 → 偶尔使用省略、破折号补充
- 句子长度均匀 → 刻意制造长短句交替
- 段落都是总-分-总 → 有时先举例再总结，有时只分析不总结
"""

    # 构建可选的优化规则 - 基于Perplexity和Burstiness原理
    optional_rules = ""

    if diversify_sentence or add_burstiness:
        optional_rules += """
## 句式突发性优化（Burstiness Enhancement）

AI检测器会分析句子长度的方差。人类写作的burstiness（突发性）更高，表现为：
- 句子长度变化大：有很短的句子（5-10字），也有很长的句子（50字以上）
- 节奏不规律：不是均匀地一句接一句
- 情绪起伏：有时详细展开，有时简略带过

**执行要求：**
1. 每个段落必须包含至少一个短句（15字以内）和一个长句（35字以上）
2. 相邻两句的长度差异要明显（不要都是20-25字的中等句）
3. 使用破折号（——）、省略号（……）制造停顿和节奏变化
4. 偶尔使用括号（）补充说明，模拟人类边写边想的状态
"""

    if increase_perplexity:
        optional_rules += """
## 降低可预测性（Perplexity Enhancement）

AI检测器通过预测下一个词来判断文本。AI文本的perplexity（困惑度）低，因为词汇选择高度可预测。

**执行要求：**
1. 避免使用"最常见"的词语搭配，选择同样正确但稍微少见的表达
2. 示例替换：
   - "进行研究" → "开展研究工作" / "着手研究" / "投入研究"
   - "取得进展" → "有所推进" / "获得进展" / "实现突破"
   - "产生影响" → "带来影响" / "形成影响" / "造成影响"
3. 不要每个概念都用最"标准"的表达，适当使用变体
"""

    if break_patterns:
        optional_rules += """
## 打破AI模式特征（Pattern Breaking）

AI有明显的写作模式，检测器专门识别这些模式：

**必须消除的模式：**
1. 三段式结构："首先...其次...最后" → 删除或只保留一个
2. 对仗结构："一方面...另一方面" → 改为不对称表达
3. 因果链条过于清晰：不是每个论点都要"因为...所以..."
4. 段落结构雷同：不要每段都是"主题句+论证+小结"
5. 连接词过多：删除30-50%的连接词，直接用句号

**替代方案：**
- 用句号直接分隔，让读者自己理解逻辑关系
- 调换句子顺序：有时先给结论，再解释原因
- 段落长度要有变化：有的段落3句话，有的段落6句话
"""

    if natural_transition:
        optional_rules += """
## 过渡词自然化处理

**删除或替换的目标：**
- 删除30%的过渡词，用句号直接分隔
- "此外" → 直接删除，新起一句
- "值得注意的是" → "这里要提一下" 或直接删除
- "具体而言" → "具体来看" 或删除
"""

    if add_human_touch and text_type != 'academic':
        optional_rules += """
## 人性化表达增强（仅非学术文本）
- 适当加入语气词："其实"、"说实话"
- 使用反问句增加表达力度
- 偶尔使用"我认为"、"在我看来"
"""

    if preserve_terms:
        optional_rules += """
## 专业术语与数据保护（绝对不能改）
- 所有专业术语必须100%保持原样
- 数字、百分比、日期不得更改
- 引用内容（引号内）保持原样
- 人名、地名、机构名保持原样
- 公式、代码保持原样
"""

    if vary_paragraph:
        optional_rules += """
## 段落结构变化
- 段落长度要有明显差异：短段落2-3句，长段落5-7句
- 不要每段开头都是总起句，可以先给例子再总结
- 段落之间的过渡可以不那么"顺滑"
"""

    # 根据文本类型添加特定规则 - 基于知网AIGC检测算法优化
    text_type_rules = ""
    if text_type == 'academic' or keep_formal_style:
        text_type_rules = """
## 【核心任务】学术论文降AI率专用规则（基于知网AIGC检测原理）

### 知网AIGC检测算法原理：
知网采用"知识增强AIGC检测技术"，从两条链路检测：
1. **语言模式链路**：检测句式规律性、词汇可预测性、连接词使用模式
2. **语义逻辑链路**：检测论述结构、逻辑链条、段落组织方式

### 为什么AI写作容易被检测？（必须理解）
- **Perplexity过低**：AI选词太"标准"，人类会用更多样的表达
- **Burstiness过低**：AI句子长度均匀（都是20-30字），人类长短句交替
- **模式化严重**：AI总是"首先、其次、最后"，人类不会这么规整
- **表达过于完美**：AI不会有冗余、犹豫、补充，人类写作有

### 核心降AI策略（必须严格执行）

**策略1：增加句式突发性（最重要，权重40%）**
- 刻意制造长短句交替：短句（10字以内）→ 长句（40字以上）→ 中句
- 示例改写：
  - 原文："该系统采用了先进的技术架构，能够有效提升工作效率。"
  - 改为："系统架构先进。它对工作效率的提升效果明显——实测数据显示，处理速度提高了约35%。"
- 每段必须有长度差异明显的句子

**策略2：打破AI连接模式（权重30%）**
- 删除"首先、其次、最后"，直接用句号分隔
- 删除"此外、另外、与此同时"，或换成"还有一点"
- "综上所述"换成"从上面的分析来看"或"回顾本节内容"
- 因果关系不要每次都用"因此、所以"，有时直接陈述结果

**策略3：降低词汇可预测性（权重20%）**
- 不用最"标准"的搭配，用同样正确的变体
- "进行研究"→"开展研究"/"着手研究"
- "取得成效"→"收到效果"/"产生成效"
- "具有重要意义"→"有其研究价值"/"值得关注"

**策略4：模拟人类写作痕迹（权重10%）**
- 适当使用括号补充：（具体数据见表3-1）
- 使用破折号连接：这一发现——虽然需要更多验证——为后续研究指明了方向
- 偶尔使用"本研究认为"、"笔者发现"代替无主语句
- 段落结构要变化：不是每段都"主题句+论证+总结"

**必须保持的学术规范（底线）**
- 章节编号、标题保持原样
- 专业术语、数据、引用100%不改
- 保持第三人称或"本研究"表述
- 禁止口语词汇：我觉得、挺好、蛮不错、啥
- 保持学术论文的严谨性和专业性
"""
    elif text_type == 'article':
        text_type_rules = """
## 一般文章优化规则
- 可以适度口语化，但保持文章的专业性
- 句式变化可以更大胆
- 可以加入更多个人观点表达
- 保持文章的可读性和流畅性
"""
    else:
        text_type_rules = """
## 非正式内容优化规则
- 可以使用口语化表达
- 句式可以不那么完整
- 可以加入语气词和感叹
- 重点是自然、像真人在说话
"""

    # 组装最终的system prompt
    system_prompt = f"""你是一个专业的文本人性化改写助手，专门针对知网、Turnitin、GPTZero等AI检测系统进行优化。

你的核心任务是将AI生成的文本改写成更像人类书写的自然文本，从而降低AI检测率。

{text_type_rules}

## 核心改写原则（基于AI检测算法原理）

### 1. 增加Perplexity（困惑度）
AI检测器通过预测下一个词来判断。AI文本词汇选择高度可预测，需要：
- 使用同义词变体，不总是用"最标准"的表达
- 避免AI偏好的固定搭配

### 2. 增加Burstiness（突发性）
AI文本句式均匀，人类写作有节奏变化：
- 制造长短句交替
- 段落长度要有变化
- 不要每段都是4-5个相似长度的句子

### 3. 打破模式特征
消除AI的规律性模式：
- 删除或替换"首先、其次、最后"
- 不要每段都是"总-分-总"结构
- 减少过渡词的使用

{mode_strategies.get(mode, mode_strategies['moderate'])}

## 当前优化强度：{intensity}/5
{intensity_settings.get(intensity, intensity_settings[3])['degree']}

{ai_signature_rules}

{optional_rules}

## 输出要求（极其重要）
- 直接输出改写后的文本，不要任何解释、标记或说明
- 不要输出"改写说明"、"优化要点"等额外信息
- 不要使用分隔符、标题或格式标记
- 保持原文的段落结构
- 只输出纯净的优化后文本"""

    user_prompt = f"""请对以下文本进行人性化改写优化：

{text_segment}

{f"上下文参考：{context}" if context else ""}

请直接输出改写后的文本："""

    try:
        headers = {
            'Authorization': f'Bearer {DEEPSEEK_API_KEY}',
            'Content-Type': 'application/json'
        }

        # 根据强度调整temperature
        temperature = intensity_settings.get(intensity, intensity_settings[3])['temp']

        payload = {
            'model': 'deepseek-chat',
            'messages': [
                {'role': 'system', 'content': system_prompt},
                {'role': 'user', 'content': user_prompt}
            ],
            'temperature': temperature,
            'max_tokens': min(len(text_segment) * 3, 4000),
            'stream': False
        }

        app.logger.info(f"开始优化文本片段，长度: {len(text_segment)}, 模式: {mode}, 强度: {intensity}, temperature: {temperature}")

        response = requests.post(DEEPSEEK_API_URL, headers=headers, json=payload, timeout=120)

        if response.status_code == 200:
            result = response.json()
            content = result['choices'][0]['message']['content']

            # 后处理：清理可能残留的格式标记
            content = content.strip()
            # 移除可能的markdown格式
            if content.startswith('```'):
                content = content.split('```')[1] if '```' in content[3:] else content
            # 移除可能的标题行
            lines = content.split('\n')
            cleaned_lines = [line for line in lines if not line.startswith('#') and not line.startswith('**改写') and not line.startswith('---')]
            content = '\n'.join(cleaned_lines).strip()

            app.logger.info(f"文本优化成功，优化后长度: {len(content)}")
            return content
        else:
            app.logger.error(f"文本优化API调用失败: {response.status_code}")
            return None

    except Exception as e:
        app.logger.error(f"文本优化失败: {e}")
        return None


def fix_broken_json(json_str):
    """修复损坏的JSON字符串"""
    try:
        import re
        if not json_str:
            return None

        # 移除多余的换行和空格
        cleaned = re.sub(r'\s+', ' ', json_str.strip())

        # 修复常见的JSON问题
        cleaned = re.sub(r',\s*}', '}', cleaned)  # 移除对象末尾多余逗号
        cleaned = re.sub(r',\s*]', ']', cleaned)  # 移除数组末尾多余逗号
        cleaned = re.sub(r'}\s*{', '},{', cleaned)  # 修复缺少逗号的对象
        cleaned = re.sub(r']\s*\[', '],[', cleaned)  # 修复缺少逗号的数组

        # 尝试提取JSON部分
        start_idx = cleaned.find('{')
        end_idx = cleaned.rfind('}') + 1

        if start_idx != -1 and end_idx != 0:
            return cleaned[start_idx:end_idx]

        return None
    except Exception as e:
        app.logger.error(f"JSON修复失败: {e}")
        return None


def extract_user_requirements_context(requirements, abstract, keywords):
    """从用户输入中提取关键上下文信息"""
    # 先记录原始输入
    app.logger.info(f"用户输入分析 - 摘要长度: {len(abstract) if abstract else 0}, 关键词: {keywords}, 要求长度: {len(requirements) if requirements else 0}")

    # 如果有摘要内容，直接从中提取信息
    if abstract and len(abstract.strip()) > 10:
        # 构建更详细的分析提示词，特别关注数据库表结构
        analysis_prompt = f"""
        请详细分析以下论文摘要和系统需求，特别关注数据库表结构信息：

        论文摘要：{abstract}
        关键词：{keywords}
        详细要求：{requirements}

        请仔细分析用户输入中的所有信息，特别是：
        1. 数据库表结构（如CREATE TABLE语句）
        2. 系统功能模块
        3. 技术栈信息
        4. 具体业务需求

        请返回详细的JSON格式分析结果：
        {{
            "tech_stack": "技术栈（如Spring Boot + Vue.js + MySQL）",
            "database_info": {{
                "type": "数据库类型",
                "tables": ["发现的所有数据表名"],
                "table_count": "表的总数量",
                "key_business_tables": ["核心业务表名"]
            }},
            "key_features": ["功能1", "功能2", "功能3"],
            "research_objectives": "研究目标",
            "system_modules": ["模块1", "模块2"],
            "business_scope": "业务范围描述"
        }}

        重要：请仔细检查用户输入中是否包含数据库表结构信息（CREATE TABLE等），如果有请详细提取。
        """

        try:
            result = call_deepseek_api(analysis_prompt, 2000)  # 增加token限制
            if result:
                import json
                import re  # 确保re模块在此函数中可用
                # 尝试解析JSON - 增强版本
                try:
                    # 首先尝试清理JSON
                    cleaned_result = result.strip()

                    # 移除markdown代码块标记
                    if cleaned_result.startswith('```json'):
                        cleaned_result = cleaned_result[7:]
                    if cleaned_result.endswith('```'):
                        cleaned_result = cleaned_result[:-3]

                    # 提取JSON部分
                    start_idx = cleaned_result.find('{')
                    end_idx = cleaned_result.rfind('}') + 1

                    if start_idx != -1 and end_idx != 0:
                        json_str = cleaned_result[start_idx:end_idx]

                        # 修复常见的JSON格式问题
                        json_str = json_str.replace('\n', ' ')  # 移除换行符
                        json_str = re.sub(r',\s*}', '}', json_str)  # 移除对象末尾多余逗号
                        json_str = re.sub(r',\s*]', ']', json_str)  # 移除数组末尾多余逗号
                        json_str = re.sub(r'}\s*{', '},{', json_str)  # 修复缺少逗号的对象
                        
                        # 尝试解析
                        context_data = json.loads(json_str)
                        app.logger.info(f"AI分析结果: {context_data}")
                        return context_data
                        
                except json.JSONDecodeError as json_error:
                    app.logger.warning(f"JSON解析失败: {json_error}, 尝试逐步修复...")
                    
                    # 尝试修复JSON结构
                    try:
                        # 更激进的修复策略
                        fixed_json = fix_broken_json(result)
                        if fixed_json:
                            context_data = json.loads(fixed_json)
                            app.logger.info(f"修复后的AI分析结果: {context_data}")
                            return context_data
                    except Exception as fix_error:
                        app.logger.warning(f"JSON修复失败: {fix_error}")
                        
                except Exception as parse_error:
                    app.logger.warning(f"JSON解析失败: {parse_error}, 原始结果长度: {len(result)}")
        except Exception as e:
            app.logger.warning(f"AI分析失败: {e}")

    # 如果AI分析失败，进行更详细的关键词匹配和信息提取
    if abstract or requirements:
        combined_text = f"{abstract} {requirements}".lower()
        
        # 提取数据库表信息
        tables = []
        table_patterns = [
            r'CREATE TABLE\s+(\w+)',
            r'表名[:：]\s*(\w+)',
            r'(\w+)表',
            r'table\s+(\w+)',
        ]
        
        for pattern in table_patterns:
            matches = re.findall(pattern, combined_text, re.IGNORECASE)
            tables.extend(matches)
        
        # 去重并过滤
        tables = list(set([t for t in tables if len(t) > 2 and t.lower() not in ['table', 'create', '数据', '信息']]))
        
        # 技术栈检测 - 更全面
        tech_components = []
        tech_mappings = {
            'spring boot': 'Spring Boot',
            'springboot': 'Spring Boot', 
            'vue': 'Vue.js',
            'vue.js': 'Vue.js',
            'mysql': 'MySQL',
            'java': 'Java',
            'javascript': 'JavaScript',
            'thymeleaf': 'Thymeleaf',
            'mybatis': 'MyBatis',
            'redis': 'Redis',
            'nginx': 'Nginx'
        }
        
        for key, value in tech_mappings.items():
            if key in combined_text:
                tech_components.append(value)
        
        # 功能特性检测 - 通用动态提取
        features = []
        # 从用户输入中动态提取功能关键词
        feature_keywords = re.findall(r'(\w*管理\w*|\w*系统\w*|\w*平台\w*|\w*服务\w*)', combined_text)
        features.extend([f for f in feature_keywords if len(f) > 2])

        # 通用功能模式检测
        common_patterns = {
            '数据管理': ['数据', '信息', 'data', 'information'],
            '用户系统': ['用户', '账户', 'user', 'account'],
            '统计分析': ['统计', '分析', '报表', 'statistics', 'analysis'],
            '权限控制': ['权限', '角色', 'permission', 'role', '认证', 'auth'],
        }

        for feature, keywords in common_patterns.items():
            if any(keyword in combined_text for keyword in keywords):
                features.append(feature)
        
        # 系统模块检测 - 通用动态提取，不硬编码具体业务
        modules = []
        if tables:
            # 从表名动态推断模块
            for table in tables[:8]:  # 取前8个主要表
                table_lower = table.lower()
                if 'admin' in table_lower or 'manage' in table_lower:
                    modules.append('管理模块')
                elif 'user' in table_lower:
                    modules.append('用户模块')
                elif 'info' in table_lower or 'data' in table_lower:
                    modules.append('信息管理模块')
                else:
                    # 根据表名生成模块名
                    module_name = table.replace('_', '').replace('table', '').replace('info', '') + '模块'
                    if len(module_name) > 2:
                        modules.append(module_name)
        
        return {
            'tech_stack': ' + '.join(tech_components) if tech_components else 'Spring Boot + Vue.js + MySQL',
            'database_info': {
                'type': 'MySQL数据库' if 'mysql' in combined_text else '关系型数据库',
                'tables': tables[:10],  # 最多返回10个表
                'table_count': len(tables),
                'key_business_tables': tables[:5]  # 核心业务表
            },
            'key_features': features if features else ['信息管理', '数据统计', '权限控制'],
            'research_objectives': '设计并实现高效的信息管理系统',
            'system_modules': list(set(modules)) if modules else ['用户管理模块', '信息管理模块'],
            'business_scope': '信息管理系统'  # 通用描述，不硬编码具体业务
        }

    # 最后的默认值
    return {
        'tech_stack': '未明确指定',
        'database_info': '未明确指定',
        'key_features': '未明确指定',
        'research_objectives': '未明确指定'
    }


def extract_section_context(content, section_name):
    """从章节内容中提取关键上下文信息"""
    if not content or len(content.strip()) < 100:
        return {
            'tech_decisions': [],
            'key_terms': {},
            'main_points': [],
            'pending_points': [],
            'continuation_needs': []
        }

    # 清理HTML标签以便分析
    import re
    clean_content = re.sub(r'<[^>]+>', '', content)

    context_prompt = f"""
    分析以下{section_name}章节内容，提取关键信息：

    {clean_content[:1500]}

    请提取并返回JSON格式：
    {{
        "tech_decisions": ["技术决策1", "技术决策2"],
        "key_terms": {{"术语1": "定义1", "术语2": "定义2"}},
        "main_points": ["主要观点1", "主要观点2"],
        "pending_points": ["待完成论证点1"],
        "continuation_needs": ["需在后续章节延续的内容1"]
    }}

    重点关注：
    1. 技术选择和实现方案
    2. 重要的专业术语及其含义
    3. 核心观点和结论
    4. 未完成的论证
    5. 需要在后续章节中继续讨论的内容
    """

    try:
        result = call_deepseek_api(context_prompt, 1200)
        if result:
            import json
            try:
                start_idx = result.find('{')
                end_idx = result.rfind('}') + 1
                if start_idx != -1 and end_idx != 0:
                    json_str = result[start_idx:end_idx]
                    context_data = json.loads(json_str)
                    return context_data
            except:
                pass
    except Exception as e:
        app.logger.warning(f"提取章节上下文失败: {e}")

    # 返回默认值
    return {
        'tech_decisions': [],
        'key_terms': {},
        'main_points': [],
        'pending_points': [],
        'continuation_needs': []
    }


def generate_context_summary(prev_content, max_length=500):
    """生成上下文摘要，保持长论文的连贯性"""
    if not prev_content or len(prev_content) < 200:
        return ""

    # 提取最后几段重要内容
    last_content = prev_content[-2000:] if len(prev_content) > 2000 else prev_content

    summary_prompt = f"""请用{max_length}字以内简要概括以下内容的核心观点和关键信息，用于后续章节的上下文连接：

{last_content}

要求：
1. 突出核心技术点和关键结论
2. 保留重要的术语和概念
3. 语言简洁准确
4. 为后续章节提供必要的背景信息"""

    try:
        summary = call_deepseek_api(summary_prompt, 800)
        return clean_ai_generated_content(summary)
    except Exception as e:
        app.logger.warning(f"生成上下文摘要失败: {e}")
        return ""


def build_contextual_prompt(section, memory, section_index):
    """构建包含完整上下文的智能提示词"""

    global_context = memory['global_context']
    previous_sections = memory['generated_sections']

    # 基础信息
    base_info = f"""
【论文基本信息】
题目：{global_context['title']}
研究领域：{global_context['field']}
论文类型：{global_context['paper_type']}

【用户详细要求】
论文摘要：{global_context['abstract']}
关键词：{global_context['keywords']}
详细要求：{global_context['requirements']}

【确定的技术方案】
技术栈：{global_context.get('tech_stack', '未明确')}
数据库：{global_context.get('database_info', '未明确')}
核心功能：{global_context.get('key_features', '未明确')}
研究目标：{global_context.get('research_objectives', '未明确')}
"""

    # 前文上下文（关键创新）
    context_info = ""
    if previous_sections:
        latest_sections = previous_sections[-2:]  # 最近2个章节
        context_info = f"""
【前文关键内容摘要】
"""
        for prev_section in latest_sections:
            context_info += f"""
{prev_section['name']}：
- 主要内容：{'; '.join(prev_section.get('main_points', [])[:3])}
- 技术决策：{'; '.join(prev_section.get('tech_decisions', [])[:2])}
- 关键术语：{', '.join(list(prev_section.get('key_terms', {}).keys())[:5])}
"""

        # 收集所有待完成的论证点
        all_pending_points = []
        for prev_section in previous_sections:
            all_pending_points.extend(prev_section.get('pending_points', []))

        # 收集所有需要延续的内容
        all_continuation_needs = []
        for prev_section in previous_sections:
            all_continuation_needs.extend(prev_section.get('continuation_needs', []))

        # 逻辑承接要求
        context_info += f"""
【逻辑承接要求】
- 必须基于前文确定的技术方案继续展开
- 保持术语使用的一致性：{', '.join(list(memory.get('key_terms', {}).keys())[:8])}
- 技术实现必须与前文描述的架构匹配
- 避免与前文内容产生矛盾
"""

        if all_pending_points:
            context_info += f"- 需要解决的遗留问题：{'; '.join(all_pending_points[:3])}\n"

        if all_continuation_needs:
            context_info += f"- 需要延续的内容：{'; '.join(all_continuation_needs[:3])}\n"

    # 当前章节要求
    current_section = f"""
【当前章节：{section['name']}】
目标字数：{section['words']}字
章节描述：{section['description']}

【具体写作要求】
1. 如果涉及技术实现，必须严格按照用户要求的技术栈：{global_context.get('tech_stack', '用户指定的技术')}
2. 数据库设计必须与用户描述保持一致：{global_context.get('database_info', '用户指定的数据库')}
3. 系统功能必须与用户需求完全匹配：{global_context.get('key_features', '用户指定的功能')}
4. 保持与前文的逻辑连贯性，避免重复或矛盾
5. 使用专业的学术写作风格，段落形式，避免分点列表
6. 如果是第一章，需要自然地引出后续章节的内容
7. 如果是中间章节，需要承接前文并为后续章节做铺垫
8. 确保内容充实，达到目标字数要求

【格式要求】
- 使用HTML格式输出
- 主标题使用<h2>标签
- 如需子标题，使用<h3>标签
- 内容使用<p>标签，每段200-300字
- 严格避免使用分点列表（1.、2.、3.或•、-等）
"""

    return base_info + context_info + current_section


def generate_references_advanced(field, title, accumulated_content):
    """高级参考文献生成 - 基于论文内容智能生成"""
    try:
        # 构建参考文献生成提示词
        prompt = f"""作为学术论文专家，请为{field}领域的论文《{title}》生成符合学术规范的参考文献。

论文内容概要：
{accumulated_content[:1000] if accumulated_content else '无'}

要求：
1. 生成15-20条高质量的参考文献
2. 包含期刊论文、会议论文、学位论文、专著等多种类型
3. 参考文献要与{field}领域高度相关
4. 遵循国际标准的引用格式
5. 包含近5年的最新研究成果
6. 作者姓名、期刊名称、年份等信息要真实可信

格式要求：
- 期刊论文：[序号] 作者. 论文标题[J]. 期刊名称, 年份, 卷号(期号): 页码.
- 会议论文：[序号] 作者. 论文标题[C]. 会议名称, 年份: 页码.
- 学位论文：[序号] 作者. 论文标题[D]. 学校名称, 年份.

请直接输出HTML格式的参考文献：
<h2>参考文献</h2>
<p>[1] ...</p>"""

        # 调用API生成参考文献
        references_content = call_deepseek_api(prompt, 2000)
        
        if references_content:
            # 清理内容
            cleaned_refs = clean_ai_generated_content(references_content)
            if cleaned_refs and len(cleaned_refs.strip()) > 100:
                return cleaned_refs
        
        # 不使用任何备用方案
        return None

    except Exception as e:
        app.logger.error(f"生成参考文献时出错: {e}")
        return None


# 删除了generate_fallback_references函数 - 不再使用静态备用参考文献


# 删除了无法访问的备用论文内容生成代码


def generate_paper_word_document(title, quill_content, references):
    """生成符合中国学术规范的本科毕业论文Word文档"""
    try:
        doc = Document()
        
        # 设置文档基本格式
        setup_document_format(doc)
        
        # 1. 创建封面页
        create_academic_cover_page(doc, title)
        
        # 2. 插入分页符
        doc.add_page_break()
        
        # 3. 处理正文内容，并收集标题信息用于生成真实目录
        headings = []
        process_academic_content_with_headings(doc, quill_content, title, headings)
        
        # 4. 添加参考文献
        if references and len(references) > 0:
            add_academic_references(doc, references)
        
        # 5. 在封面后插入真实目录
        insert_real_toc_after_cover(doc, headings)
        
        # 6. 设置页眉页脚
        setup_academic_header_footer(doc, title)
        
        # 保存到内存
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        return doc_buffer
        
    except Exception as e:
        app.logger.error(f"生成Word文档时出错: {e}")
        return create_error_document(title, str(e))


def setup_document_format(doc):
    """设置文档基本格式"""
    # 设置页面布局
    section = doc.sections[0]
    section.page_height = Inches(11.69)  # A4纸高度
    section.page_width = Inches(8.27)    # A4纸宽度
    section.left_margin = Inches(1.25)   # 左边距3.17cm
    section.right_margin = Inches(1.0)   # 右边距2.54cm
    section.top_margin = Inches(1.0)     # 上边距2.54cm
    section.bottom_margin = Inches(1.0)  # 下边距2.54cm
    
    # 设置基础正文样式
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Times New Roman'  # 英文字体
    normal_font.size = Pt(12)
    normal_format = normal_style.paragraph_format
    normal_format.line_spacing = 1.5      # 1.5倍行距
    normal_format.space_after = Pt(0)
    normal_format.first_line_indent = Inches(0.5)  # 首行缩进2字符
    
    # 设置中文字体
    normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 创建标题样式
    create_heading_styles(doc)


def create_heading_styles(doc):
    """创建标准的学术标题样式"""
    try:
        # 一级标题样式
        h1_style = doc.styles.add_style('Academic Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        h1_font = h1_style.font
        h1_font.name = 'Times New Roman'
        h1_font.size = Pt(16)
        h1_font.bold = True
        h1_style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        h1_format = h1_style.paragraph_format
        h1_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h1_format.space_before = Pt(18)
        h1_format.space_after = Pt(12)
        h1_format.first_line_indent = 0
        
        # 二级标题样式
        h2_style = doc.styles.add_style('Academic Heading 2', WD_STYLE_TYPE.PARAGRAPH)
        h2_font = h2_style.font
        h2_font.name = 'Times New Roman'
        h2_font.size = Pt(14)
        h2_font.bold = True
        h2_style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        h2_format = h2_style.paragraph_format
        h2_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h2_format.space_before = Pt(15)
        h2_format.space_after = Pt(10)
        h2_format.first_line_indent = 0
        
        # 三级标题样式
        h3_style = doc.styles.add_style('Academic Heading 3', WD_STYLE_TYPE.PARAGRAPH)
        h3_font = h3_style.font
        h3_font.name = 'Times New Roman'
        h3_font.size = Pt(12)
        h3_font.bold = True
        h3_style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        h3_format = h3_style.paragraph_format
        h3_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h3_format.space_before = Pt(12)
        h3_format.space_after = Pt(8)
        h3_format.first_line_indent = 0
        
    except ValueError:
        # 样式已存在，跳过
        pass


def create_academic_cover_page(doc, title):
    """创建符合学术规范的封面页"""
    # 顶部空白
    for _ in range(3):
        doc.add_paragraph()
    
    # 学校名称
    university_para = doc.add_paragraph()
    university_run = university_para.add_run('××大学')
    university_run.font.name = '华文中宋'
    university_run.font.size = Pt(26)
    university_run.font.bold = True
    university_run._element.rPr.rFonts.set(qn('w:eastAsia'), '华文中宋')
    university_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    university_para.paragraph_format.space_after = Pt(18)
    
    # 论文类型
    type_para = doc.add_paragraph()
    type_run = type_para.add_run('本科毕业论文')
    type_run.font.name = '华文中宋'
    type_run.font.size = Pt(22)
    type_run.font.bold = True
    type_run._element.rPr.rFonts.set(qn('w:eastAsia'), '华文中宋')
    type_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    type_para.paragraph_format.space_after = Pt(36)
    
    # 中间空白
    for _ in range(3):
        doc.add_paragraph()
    
    # 论文标题
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.font.name = '黑体'
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(48)
    title_para.paragraph_format.space_before = Pt(24)
    
    # 下方空白
    for _ in range(4):
        doc.add_paragraph()
    
    # 学生信息部分
    info_items = [
        ('学生姓名', ''),
        ('学    号', ''),
        ('专    业', ''),
        ('班    级', ''),
        ('指导教师', ''),
        ('学    院', ''),
        ('完成日期', datetime.now().strftime('%Y年%m月%d日'))
    ]
    
    for label, value in info_items:
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_para.paragraph_format.space_after = Pt(12)
        
        # 标签
        label_run = info_para.add_run(f'{label}：')
        label_run.font.name = '宋体'
        label_run.font.size = Pt(16)
        label_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 下划线空白或值
        if value:
            value_run = info_para.add_run(value)
            value_run.font.name = '宋体'
            value_run.font.size = Pt(16)
            value_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        else:
            # 添加下划线
            underline_run = info_para.add_run('_' * 20)
            underline_run.font.name = '宋体'
            underline_run.font.size = Pt(16)


def insert_real_toc_after_cover(doc, headings):
    """在封面后插入真实目录"""
    try:
        # 添加目录标题
        doc.add_page_break()  # 在封面后分页
        
        # 目录标题
        toc_title = doc.add_paragraph('目  录')
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        toc_title.runs[0].font.name = 'Times New Roman'
        toc_title.runs[0].font.size = Pt(18)
        toc_title.runs[0].font.bold = True
        toc_title.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        toc_title.paragraph_format.space_after = Pt(24)
        toc_title.paragraph_format.space_before = Pt(24)
        
        # 添加目录项
        for heading in headings:
            toc_para = doc.add_paragraph()
            
            # 根据标题级别设置缩进
            indent = (heading['level'] - 1) * 0.5
            toc_para.paragraph_format.left_indent = Inches(indent)
            toc_para.paragraph_format.space_after = Pt(6)
            
            # 标题文本
            title_run = toc_para.add_run(heading['text'])
            title_run.font.name = 'Times New Roman'
            title_run.font.size = Pt(14) if heading['level'] == 1 else Pt(12)
            title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            # 添加制表符和页码
            toc_para.add_run('\t')
            page_run = toc_para.add_run(str(heading['page']))
            page_run.font.name = 'Times New Roman'
            page_run.font.size = Pt(12)
            
            # 设置制表位（右对齐页码）
            tab_stops = toc_para.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(5.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        
        # 在目录后添加分页符
        doc.add_page_break()
        
    except Exception as e:
        app.logger.error(f"插入目录时出错: {e}")
        # 如果目录生成失败，添加简单提示
        doc.add_page_break()
        simple_toc = doc.add_paragraph("目录")
        simple_toc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        simple_toc.runs[0].font.bold = True
        simple_toc.runs[0].font.size = Pt(18)
        doc.add_paragraph("（请在Word中通过引用-目录功能生成目录）")
        doc.add_page_break()


def process_academic_content_with_headings(doc, quill_content, title, headings):
    """处理学术论文正文内容，并收集标题信息"""
    if isinstance(quill_content, dict) and 'ops' in quill_content:
        try:
            current_paragraph = None
            page_counter = 1  # 简单的页码计数器
            
            for op in quill_content['ops']:
                if 'insert' in op:
                    text = op['insert']
                    attributes = op.get('attributes', {})
                    
                    # 处理换行符
                    if text == '\n':
                        current_paragraph = None
                        continue
                    
                    # 跳过空字符串
                    if not text.strip():
                        continue
                    
                    # 创建段落
                    if current_paragraph is None:
                        if attributes.get('header'):
                            level = min(int(attributes['header']), 3)
                            
                            # 收集标题信息
                            headings.append({
                                'text': text.strip(),
                                'level': level,
                                'page': page_counter
                            })
                            
                            # 创建标题段落
                            current_paragraph = doc.add_paragraph(text.strip())
                            
                            # 应用对应的标题样式
                            if level == 1:
                                current_paragraph.style = doc.styles['Academic Heading 1']
                            elif level == 2:
                                current_paragraph.style = doc.styles['Academic Heading 2']
                            else:
                                current_paragraph.style = doc.styles['Academic Heading 3']
                                
                            # 估算页码增长
                            page_counter += 1
                            
                        else:
                            # 正文段落
                            current_paragraph = doc.add_paragraph()
                            current_paragraph.style = doc.styles['Normal']
                            
                            # 添加文本
                            run = current_paragraph.add_run(text)
                            
                            # 应用正文格式
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(12)
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                            
                            # 应用文本属性
                            if attributes.get('bold'):
                                run.font.bold = True
                            if attributes.get('italic'):
                                run.font.italic = True
                            if attributes.get('underline'):
                                run.font.underline = True
                    else:
                        # 继续添加到当前段落
                        run = current_paragraph.add_run(text)
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        
                        # 应用文本属性
                        if attributes.get('bold'):
                            run.font.bold = True
                        if attributes.get('italic'):
                            run.font.italic = True
                        if attributes.get('underline'):
                            run.font.underline = True
                            
        except Exception as e:
            app.logger.error(f"处理正文内容时出错: {e}")
            # 如果出错，添加一个简单的错误提示段落
            error_para = doc.add_paragraph("内容处理出现错误，请检查论文格式。")
            error_para.style = doc.styles['Normal']
    else:
        # 处理其他格式内容
        if isinstance(quill_content, str):
            # 增强的图片和无关内容移除逻辑
            import re
            clean_content = quill_content
            
            # 移除所有图片相关标签和内容
            clean_content = re.sub(r'<img[^>]*>', '', clean_content)  # 完全移除img标签
            clean_content = re.sub(r'<figure[^>]*>.*?</figure>', '', clean_content, flags=re.DOTALL)  # 移除figure
            clean_content = re.sub(r'<picture[^>]*>.*?</picture>', '', clean_content, flags=re.DOTALL)  # 移除picture
            clean_content = re.sub(r'<svg[^>]*>.*?</svg>', '', clean_content, flags=re.DOTALL)  # 移除svg
            
            # 移除图片相关的文字描述
            clean_content = re.sub(r'如图.*?所示[：，。]', '', clean_content)
            clean_content = re.sub(r'见图\s*\d+.*?[：，。]', '', clean_content)
            clean_content = re.sub(r'图\s*\d+.*?显示.*?[：，。]', '', clean_content)
            clean_content = re.sub(r'上图.*?[：，。]', '', clean_content)
            clean_content = re.sub(r'下图.*?[：，。]', '', clean_content)
            clean_content = re.sub(r'图表.*?说明.*?[：，。]', '', clean_content)
            
            # 移除其他HTML标签，但保留文本结构
            clean_content = re.sub(r'<h([1-6])[^>]*>(.*?)</h[1-6]>', r'\\n\\n**\\2**\\n', clean_content)  # 标题转换
            clean_content = re.sub(r'<p[^>]*>(.*?)</p>', r'\\1\\n\\n', clean_content)  # 段落
            clean_content = re.sub(r'<[^>]+>', '', clean_content)  # 移除剩余HTML标签
            
            # 应用AI内容清理函数
            clean_content = clean_ai_generated_content(clean_content)
            
            # 清理多余的空行和空白字符
            clean_content = re.sub(r'\n\s*\n\s*\n+', '\n\n', clean_content)
            clean_content = clean_content.strip()
            
            if clean_content:
                # 按段落分割并添加到文档
                paragraphs = clean_content.split('\n\n')
                for para_text in paragraphs:
                    para_text = para_text.strip()
                    if para_text:
                        # 检查是否是标题格式
                        if para_text.startswith('**') and para_text.endswith('**'):
                            title_text = para_text.strip('*').strip()
                            if title_text:
                                doc.add_heading(title_text, level=2)
                        else:
                            doc.add_paragraph(para_text)
            else:
                doc.add_paragraph("生成的内容为空，请检查输入参数。")
        else:
            doc.add_paragraph("无有效内容可导出。")


def add_academic_references(doc, references):
    """添加学术规范的参考文献部分"""
    # 添加分页符
    doc.add_page_break()
    
    # 参考文献标题
    ref_heading = doc.add_heading('参考文献', level=1)
    ref_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 设置参考文献标题样式
    for run in ref_heading.runs:
        run.font.name = '黑体'
        run.font.size = Pt(16)
        run.font.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    ref_heading.paragraph_format.space_before = Pt(18)
    ref_heading.paragraph_format.space_after = Pt(18)
    
    # 添加参考文献内容
    for ref in references:
        p = doc.add_paragraph()
        ref_text = f"[{ref.get('number', '1')}] {ref.get('formatted', '参考文献格式错误')}"
        run = p.add_run(ref_text)
        
        # 设置参考文献格式
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10.5)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 设置段落格式
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.first_line_indent = 0
        p.paragraph_format.left_indent = Inches(0.25)  # 悬挂缩进


def setup_academic_header_footer(doc, title):
    """设置学术规范的页眉页脚"""
    section = doc.sections[0]
    
    # 设置页眉
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = title
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 设置页眉样式
    for run in header_para.runs:
        run.font.name = '宋体'
        run.font.size = Pt(9)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 设置页脚（页码）
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.text = "- "
    
    # 添加页码字段
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    footer_run = footer_para.runs[0]
    footer_run._r.append(fldChar1)
    footer_run._r.append(instrText)
    footer_run._r.append(fldChar2)
    footer_para.add_run(" -")
    
    # 设置页脚样式
    for run in footer_para.runs:
        run.font.name = '宋体'
        run.font.size = Pt(9)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def create_error_document(title, error_message):
    """创建错误文档"""
    try:
        doc = Document()
        doc.add_heading('文档生成错误', level=1)
        doc.add_paragraph(f'论文标题: {title}')
        doc.add_paragraph(f'错误时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph(f'错误信息: {error_message}')
        doc.add_paragraph('请检查内容格式或联系技术支持。')
        
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        return doc_buffer
    except:
        return None


@app.route('/api/get_defense_cost')
def api_get_defense_cost():
    """获取答辩问题生成费用"""
    try:
        cost = float(user_manager.get_system_config('thesis_defense_cost', 5.00))
        balance = 0
        logged_in = 'user_id' in session
        
        if logged_in:
            user_info = user_manager.get_user_info(session['user_id'])
            if user_info:
                balance = float(user_info['balance'])
        
        return jsonify({
            'success': True,
            'cost': cost,
            'balance': balance,
            'logged_in': logged_in
        })
    except Exception as e:
        app.logger.error(f"获取答辩生成费用失败: {e}")
        return jsonify({'success': False, 'message': str(e)}), 500


@app.route('/api/generate-defense-questions', methods=['POST'])
def api_generate_defense_questions():
    """生成论文答辩问题API"""
    try:
        data = request.get_json()
        mode = data.get('mode', 'normal')  # 新增模式参数

        # 生成唯一的会话ID
        session_id = str(uuid.uuid4())[:8] + str(int(time.time()))[-6:]
        generation_start_time = time.time()

        # 论文片段分析模式
        if mode == 'fragment':
            return handle_fragment_analysis(data, session_id, generation_start_time)

        # 原有的全论文分析模式
        thesis_title = data.get('thesisTitle', '')
        research_field = data.get('researchField', '')
        thesis_abstract = data.get('thesisAbstract', '')
        system_name = data.get('systemName', '')
        tech_stack = data.get('techStack', '')
        system_description = data.get('systemDescription', '')
        question_count = data.get('questionCount', 10)
        difficulty_level = data.get('difficultyLevel', '中等')
        category = data.get('category', None)  # 新增类别参数

        if not all([thesis_title, research_field, thesis_abstract]):
            return jsonify({'success': False, 'message': '请提供完整的论文基本信息'}), 400

        # 检查用户登录状态
        if 'user_id' not in session:
            return jsonify({
                'success': False,
                'message': 'AI功能需要登录后使用',
                'need_login': True
            }), 401
        
        user_id = session['user_id']
        
        # 获取费用并检查余额
        cost = float(user_manager.get_system_config('thesis_defense_cost', 5.00))
        user_info = user_manager.get_user_info(user_id)
        
        if not user_info or user_info['balance'] < cost:
            return jsonify({
                'success': False,
                'message': f'余额不足，AI生成需要 {cost:.2f} 元，当前余额 {user_info["balance"]:.2f} 元',
                'need_recharge': True,
                'cost': cost,
                'balance': float(user_info['balance']) if user_info else 0
            }), 400

        # 调用AI生成答辩问题
        questions, is_ai_success = generate_defense_questions_with_ai(
            thesis_title, research_field, thesis_abstract,
            system_name, tech_stack, system_description,
            question_count, difficulty_level, category  # 传递类别参数
        )

        # 只有AI真正成功才扣费
        charged = False
        if is_ai_success:
            consume_result = user_manager.consume_balance(
                user_id,
                cost,
                'thesis_defense',
                f'AI答辩问题生成 - {thesis_title[:30]}'
            )
            charged = consume_result['success']
            if charged:
                app.logger.info(f"AI答辩问题生成成功，已扣费 {cost} 元")
            else:
                app.logger.warning(f"AI答辩问题生成扣费失败: {consume_result['message']}")
        else:
            app.logger.warning("AI调用失败，使用备用方案，不扣费")

        # 计算生成时间
        generation_time = int(time.time() - generation_start_time)

        # 保存到数据库历史记录
        history_id = None
        try:
            history_id = user_manager.save_defense_question_history(
                user_id=user_id,
                session_id=session_id,
                generation_mode='single' if not category else 'category',
                thesis_data=data,
                questions_data=questions,
                generation_time=generation_time
            )
            app.logger.info(f"历史记录已保存，ID: {history_id}")
        except Exception as e:
            app.logger.error(f"保存历史记录失败: {e}")
            # 不影响主要功能，继续返回结果

        return jsonify({
            'success': True,
            'questions': questions,
            'message': f'成功生成 {len(questions)} 个答辩问题' + ('' if is_ai_success else '（AI服务暂时不可用，已使用备用方案，未扣费）'),
            'session_id': session_id,
            'history_id': history_id,
            'charged': charged,
            'cost': cost,
            'ai_success': is_ai_success
        })

    except Exception as e:
        app.logger.error(f"Error generating defense questions: {e}")
        return jsonify({'success': False, 'message': f'生成答辩问题失败: {str(e)}'}), 500

def handle_fragment_analysis(data, session_id, generation_start_time):
    """处理论文片段分析"""
    try:
        thesis_title = data.get('thesisTitle', '论文')
        research_field = data.get('researchField', '计算机科学')
        fragment = data.get('thesisFragment', '')
        context = data.get('fragmentContext', 'auto')
        question_count = data.get('questionCount', 5)
        difficulty_level = data.get('difficultyLevel', '中等')
        
        if not fragment.strip():
            return jsonify({'success': False, 'message': '论文片段内容不能为空'}), 400
            
        if len(fragment) < 50:
            return jsonify({'success': False, 'message': '论文片段内容过短，请提供更详细的内容'}), 400
            
        if len(fragment) > 3000:
            return jsonify({'success': False, 'message': '论文片段内容过长，请控制在3000字符以内'}), 400
        
        # 调用AI分析论文片段
        questions = generate_fragment_questions_with_ai(
            thesis_title, research_field, fragment, context, question_count, difficulty_level
        )
        
        # 计算生成时间
        generation_time = int(time.time() - generation_start_time)
        
        # 保存到数据库历史记录
        try:
            user_id = session['user_id']
            history_id = user_manager.save_defense_question_history(
                user_id=user_id,
                session_id=session_id,
                generation_mode='fragment',
                thesis_data=data,
                questions_data=questions,
                generation_time=generation_time
            )
            app.logger.info(f"片段分析历史记录已保存，ID: {history_id}")
        except Exception as e:
            app.logger.error(f"保存片段分析历史记录失败: {e}")
        
        return jsonify({
            'success': True,
            'questions': questions,
            'message': f'基于论文片段成功生成 {len(questions)} 个答辩问题',
            'session_id': session_id,
            'history_id': history_id if 'history_id' in locals() else None
        })
        
    except Exception as e:
        app.logger.error(f"Error in fragment analysis: {e}")
        return jsonify({'success': False, 'message': f'片段分析失败: {str(e)}'}), 500


@app.route('/api/export-defense-questions-word', methods=['POST'])
def api_export_defense_questions_word():
    """导出答辩问题到Word文档"""
    try:
        data = request.get_json()
        questions = data.get('questions', [])
        thesis_title = data.get('thesisTitle', '论文')
        research_field = data.get('researchField', '计算机科学')

        if not questions:
            return jsonify({'error': '没有答辩问题数据'}), 400

        # 生成Word文档
        doc_buffer = generate_defense_questions_word(questions, thesis_title, research_field)

        # 返回文件
        return send_file(
            doc_buffer,
            as_attachment=True,
            download_name=f'论文答辩问题_{thesis_title}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        app.logger.error(f"Error exporting defense questions to Word: {e}")
        return jsonify({'error': f'导出Word文档失败: {str(e)}'}), 500


def generate_fragment_questions_with_ai(thesis_title, research_field, fragment, context, question_count, difficulty_level):
    """基于论文片段生成答辩问题"""
    try:
        # 构建内容类型提示
        context_prompts = {
            'algorithm': '这是算法原理相关的内容',
            'experiment': '这是实验设计或结果分析相关的内容',
            'system': '这是系统实现或架构设计相关的内容',
            'innovation': '这是创新点或贡献阐述相关的内容',
            'technical': '这是技术难点或解决方案相关的内容',
            'theory': '这是理论分析或推导相关的内容',
            'evaluation': '这是性能评估或对比分析相关的内容',
            'conclusion': '这是结论总结相关的内容',
            'auto': '请自动识别内容类型'
        }
        
        context_hint = context_prompts.get(context, '请自动识别内容类型')
        
        # 构建专门针对片段分析的提示词
        prompt = f"""你是资深答辩委员会教授，擅长从论文片段中发现深层次问题。请针对以下论文片段，生成{question_count}个高度精准的答辩问题。

论文基本信息：
- 题目：{thesis_title}
- 研究领域：{research_field}
- 学位级别：{difficulty_level}

重点分析片段：
{fragment[:1000]}{"..." if len(fragment) > 1000 else ""}

内容类型：{context_hint}

精准分析要求：
1. 问题必须完全基于这段具体内容，避免泛泛而谈
2. 深度挖掘片段中的技术方法、实现细节、理论依据
3. 模拟真实答辩场景中教授会针对这段内容提出的尖锐问题
4. 关注技术选择的合理性、实现的完整性、结果的可信度
5. 问题要能测试学生对这段内容的深度理解和掌握

答案标准：
- 每个答案300-500字，逻辑清晰，层次分明
- 必须紧扣片段内容，体现专业深度
- 展现对技术细节的准确理解
- 能够说服答辩委员会的专业水准

输出格式：
严格按照JSON格式返回，不要任何额外文字：
[{{"category": "问题分类", "question": "针对片段的精准问题", "answer": "专业详细答案"}}]"""

        headers = {
            'Authorization': f'Bearer {DEEPSEEK_API_KEY}',
            'Content-Type': 'application/json'
        }

        payload = {
            'model': 'deepseek-chat',
            'messages': [{'role': 'user', 'content': prompt}],
            'temperature': 0.7,
            'max_tokens': 3000  # 修改为DeepSeek API安全限制内
        }

        # 重试机制
        for attempt in range(3):
            try:
                # 记录请求信息以便调试
                app.logger.info(f"DeepSeek API 片段分析请求第{attempt+1}次，prompt长度: {len(prompt)}")
                
                response = requests.post(DEEPSEEK_API_URL, headers=headers, json=payload, timeout=90)
                
                if response.status_code == 200:
                    result = response.json()
                    content = result['choices'][0]['message']['content']
                    
                    # 提取JSON部分
                    start_idx = content.find('[')
                    end_idx = content.rfind(']') + 1
                    if start_idx != -1 and end_idx != 0:
                        json_str = content[start_idx:end_idx]
                        questions = json.loads(json_str)
                        
                        # 验证和完善数据
                        for i, question in enumerate(questions):
                            if 'category' not in question:
                                question['category'] = f"片段分析问题{i+1}"
                            if 'question' not in question:
                                question['question'] = "请详细解释这段内容的核心要点"
                            if 'answer' not in question or len(question['answer']) < 100:
                                question['answer'] = generate_fragment_fallback_answer(fragment, question.get('question', ''))
                        
                        app.logger.info(f"AI成功生成{len(questions)}个片段分析问题")
                        return questions
                        
                elif response.status_code == 429:  # 速率限制
                    if attempt < 2:
                        time.sleep(2 ** attempt)
                        continue
                    
                # 记录详细的API错误信息
                app.logger.error(f"DeepSeek API 片段分析第{attempt+1}次尝试失败: {response.status_code}")
                app.logger.error(f"响应内容: {response.text}")
                
            except requests.Timeout:
                app.logger.warning(f"DeepSeek API 片段分析第{attempt+1}次尝试超时")
                if attempt < 2:
                    continue
            except Exception as e:
                app.logger.error(f"DeepSeek API 片段分析第{attempt+1}次尝试出错: {e}")
                if attempt < 2:
                    continue
        
        # 所有尝试都失败，返回智能备用方案
        app.logger.warning("DeepSeek API片段分析调用失败，使用智能备用方案")
        return generate_fragment_fallback_questions(thesis_title, research_field, fragment, context, question_count, difficulty_level)

    except Exception as e:
        app.logger.error(f"片段分析生成问题时出错: {e}")
        return generate_fragment_fallback_questions(thesis_title, research_field, fragment, context, question_count, difficulty_level)

def generate_fragment_fallback_answer(fragment, question):
    """为片段问题生成备用答案"""
    return f"针对论文片段中的内容，{question}可以从以下几个方面回答：1）分析片段中的核心技术点和关键信息；2）解释实现原理和方法选择的依据；3）讨论可能存在的技术挑战和解决方案；4）评估方法的有效性和适用范围。建议结合片段的具体内容进行详细阐述。"

def generate_fragment_fallback_questions(thesis_title, research_field, fragment, context, question_count, difficulty_level):
    """片段分析备用问题生成"""
    # 基于片段内容特征生成问题
    fallback_questions = []
    
    # 分析片段中的关键词
    keywords = []
    if '算法' in fragment or 'algorithm' in fragment.lower():
        keywords.append('algorithm')
    if '实验' in fragment or 'experiment' in fragment.lower():
        keywords.append('experiment')
    if '系统' in fragment or 'system' in fragment.lower():
        keywords.append('system')
    if '结果' in fragment or 'result' in fragment.lower():
        keywords.append('result')
    if '方法' in fragment or 'method' in fragment.lower():
        keywords.append('method')
    
    # 生成通用问题模板
    base_questions = [
        {
            "category": "内容理解",
            "question": f"请详细解释这段内容中提到的核心观点和关键技术？",
            "answer": f"这段内容的核心观点体现在：1）主要技术方法的选择和应用；2）关键实现细节和技术特点；3）与{research_field}领域相关理论的结合；4）实际效果和预期目标的实现。通过深入分析可以看出，该方法在解决相关问题时具有一定的创新性和实用性。"
        },
        {
            "category": "技术分析",
            "question": f"您在这段内容中采用的技术方案有什么优势？为什么选择这种方法？",
            "answer": f"技术方案的优势主要体现在：1）方法的科学性和合理性；2）实现复杂度和性能的平衡；3）与研究目标的匹配度；4）相比其他方案的创新点。选择这种方法的原因包括技术可行性、效果预期、资源限制等多方面考虑。"
        },
        {
            "category": "深入探讨",
            "question": f"这段内容中是否存在技术难点？您是如何解决的？",
            "answer": f"技术难点主要包括：1）理论方法向实际应用的转化；2）性能优化和效率提升；3）特殊情况和异常处理；4）系统集成和兼容性问题。解决方案采用了多种技术手段和策略，通过不断优化和改进达到了预期效果。"
        },
        {
            "category": "效果评估",
            "question": f"如何验证这段内容中提到的方法或结果的有效性？",
            "answer": f"验证方法的有效性采用了多重策略：1）理论分析和数学证明；2）实验设计和数据验证；3）对比分析和基准测试；4）实际应用场景的验证。通过全面的评估体系，确保了方法的科学性和实用性。"
        },
        {
            "category": "扩展思考",
            "question": f"基于这段内容，您认为还有哪些可以改进或扩展的地方？",
            "answer": f"改进和扩展的方向包括：1）算法效率和性能的进一步优化；2）适用范围和应用场景的扩展；3）与其他技术方法的融合；4）理论深度和实践价值的提升。这些改进将有助于推动{research_field}领域的进一步发展。"
        }
    ]
    
    # 根据关键词调整问题
    if 'algorithm' in keywords:
        base_questions.append({
            "category": "算法分析",
            "question": "请详细说明算法的时间复杂度和空间复杂度，以及优化策略？",
            "answer": "算法复杂度分析需要考虑：1）时间复杂度的理论分析和实际表现；2）空间复杂度的内存使用情况；3）不同数据规模下的性能表现；4）针对性的优化策略和改进方案。"
        })
    
    if 'experiment' in keywords:
        base_questions.append({
            "category": "实验设计",
            "question": "实验设计的合理性如何？实验结果说明了什么问题？",
            "answer": "实验设计的合理性体现在：1）实验目标和假设的明确性；2）实验条件和参数的控制；3）数据收集和分析方法的科学性；4）结果解释和结论的可靠性。"
        })
    
    # 返回指定数量的问题
    return base_questions[:question_count]

def generate_defense_questions_with_ai(thesis_title, research_field, thesis_abstract,
                                     system_name, tech_stack, system_description,
                                     question_count, difficulty_level, category=None):
    """使用DeepSeek AI生成论文答辩问题 - 优化版本
    返回: (questions, is_ai_success) - 问题列表和是否AI成功的标志
    """
    try:
        # 构建类别特定的提示词
        category_prompts = {
            'basic': '重点关注研究背景、基础概念、文献综述等基础理论问题',
            'technical': '重点关注技术方案、算法原理、系统设计、实现方法等技术问题',
            'experiment': '重点关注实验设计、数据分析、结果验证、性能评估等实验相关问题',
            'advanced': '重点关注创新点、理论深度、复杂度分析、优缺点等深入分析问题',
            'application': '重点关注实际应用、商业价值、推广前景、社会影响等应用前景问题',
            'system': '重点关注系统架构、模块设计、技术选型、部署方案等系统实现问题',
            'background': '重点关注研究背景、问题意义、现状分析等背景相关问题',
            'innovation': '重点关注创新点、贡献价值、技术突破等创新相关问题',
            'theory': '重点关注理论基础、数学模型、算法证明等理论分析问题',
            'future': '重点关注未来工作、发展方向、改进计划等未来展望问题'
        }
        
        # 构建基础提示词
        base_prompt = f"""你是有20年经验的计算机专业答辩委员会资深教授。请基于真实答辩场景，为这篇论文生成{question_count}个精准的答辩问题。

论文信息：
题目：{thesis_title}
研究领域：{research_field}
论文摘要：{thesis_abstract[:500]}{"..." if len(thesis_abstract) > 500 else ""}

系统信息：
系统名称：{system_name or "无"}
技术栈：{tech_stack or "无"}
系统功能：{system_description[:300] if system_description else "无"}

答辩要求：
- 学位级别：{difficulty_level}
- 问题数量：{question_count}个
- 模拟真实答辩委员会的提问风格和深度
- 问题要能体现学生的专业能力和研究深度

重要指导原则：
1. 问题必须紧密结合论文的具体内容和技术细节
2. 体现答辩委员会教授的专业水平和关注重点
3. 问题层次递进：基础理解→技术深度→创新评价→应用前景
4. 每个问题都应该是答辩现场真实会被问到的
5. 问题要能测试学生对自己研究工作的掌握程度"""

        # 添加类别特定要求
        if category and category in category_prompts:
            base_prompt += f"""
- 问题类型：{category_prompts[category]}
- 每个问题都应该围绕该类别的核心内容进行设计"""
        else:
            base_prompt += """
- 问题类型分布：
  * 研究背景与意义 (15%)
  * 文献综述与相关工作 (10%) 
  * 技术方案与实现 (25%)
  * 实验设计与结果分析 (20%)
  * 创新点与贡献 (15%)
  * 系统架构与设计 (10%)
  * 未来工作与改进 (5%)"""

        # 完成提示词
        prompt = base_prompt + f"""

答案要求：
- 每个答案400-600字，结构清晰，逻辑严密
- 答案要体现深度思考和专业理解
- 包含具体的技术细节和实现要点
- 展现对研究领域的深入认知
- 答案应该是优秀答辩表现的标准

输出格式：
请严格按照JSON数组格式返回，不要添加任何其他文字说明：
[{{"category": "问题分类", "question": "具体问题内容", "answer": "详细专业答案"}}]

注意：确保生成完整的{question_count}个问题，每个问题都要贴合论文实际内容。"""

        headers = {
            'Authorization': f'Bearer {DEEPSEEK_API_KEY}',
            'Content-Type': 'application/json'
        }

        payload = {
            'model': 'deepseek-chat',
            'messages': [{'role': 'user', 'content': prompt}],
            'temperature': 0.7,
            'max_tokens': 3000  # 修改为DeepSeek API安全限制内
        }

        # 增加超时时间并添加重试机制
        for attempt in range(3):  # 最多重试3次
            try:
                # 记录请求信息以便调试
                app.logger.info(f"DeepSeek API 请求第{attempt+1}次，prompt长度: {len(prompt)}")
                
                response = requests.post(DEEPSEEK_API_URL, headers=headers, json=payload, timeout=150)  # 增加超时时间到150秒
                
                if response.status_code == 200:
                    result = response.json()
                    content = result['choices'][0]['message']['content']
                    
                    # 提取JSON部分
                    start_idx = content.find('[')
                    end_idx = content.rfind(']') + 1
                    if start_idx != -1 and end_idx != 0:
                        json_str = content[start_idx:end_idx]
                        questions = json.loads(json_str)
                        
                        # 验证和完善数据
                        for i, question in enumerate(questions):
                            if 'category' not in question:
                                question['category'] = f"第{i+1}类问题"
                            if 'question' not in question:
                                question['question'] = "请详细阐述相关内容"
                            if 'answer' not in question or len(question['answer']) < 100:
                                question['answer'] = generate_fallback_answer(thesis_title, question.get('question', ''))
                        
                        app.logger.info(f"AI成功生成{len(questions)}个问题")
                        return (questions, True)  # AI成功
                        
                elif response.status_code == 429:  # 速率限制
                    if attempt < 2:
                        time.sleep(2 ** attempt)  # 指数退避
                        continue
                    
                # 记录详细的API错误信息
                app.logger.error(f"DeepSeek API 第{attempt+1}次尝试失败: {response.status_code}")
                app.logger.error(f"响应内容: {response.text}")
                
            except requests.Timeout:
                app.logger.warning(f"DeepSeek API 第{attempt+1}次尝试超时")
                if attempt < 2:
                    continue
            except Exception as e:
                app.logger.error(f"DeepSeek API 第{attempt+1}次尝试出错: {e}")
                if attempt < 2:
                    continue
        
        # 所有尝试都失败，返回智能备用方案
        app.logger.warning("DeepSeek API调用失败，使用智能备用方案")
        return (generate_smart_fallback_questions(thesis_title, research_field, thesis_abstract, 
                                               system_name, question_count, difficulty_level, category), False)

    except Exception as e:
        app.logger.error(f"生成答辩问题时出错: {e}")
        return (generate_smart_fallback_questions(thesis_title, research_field, thesis_abstract, 
                                               system_name, question_count, difficulty_level, category), False)


def generate_category_specific_questions(thesis_title, research_field, thesis_abstract, 
                                       system_name, question_count, difficulty_level, category):
    """生成特定类别的问题"""
    category_templates = {
        'background': [
            {
                "category": "研究背景与意义",
                "question": f"请详细阐述您选择'{thesis_title}'这个研究课题的背景和现实意义？",
                "answer": f"本研究基于{research_field}领域的发展需求，主要解决了以下问题：1）当前技术的局限性和挑战；2）研究问题的重要性和紧迫性；3）预期成果的学术价值和实际应用价值；4）对相关领域发展的推动作用。"
            },
            {
                "category": "研究背景与意义", 
                "question": "当前该研究领域存在哪些主要问题和挑战？您的研究如何解决这些问题？",
                "answer": "当前研究领域面临的主要挑战包括：1）技术方法的局限性；2）应用场景的复杂性；3）性能效率的瓶颈；4）实际部署的困难。本研究通过创新的方法和技术手段，为这些问题提供了有效的解决方案。"
            }
        ],
        'technical': [
            {
                "category": "技术方案与实现",
                "question": "请详细介绍您采用的核心技术方案和算法原理？",
                "answer": f"本研究采用的技术方案具有以下特点：1）核心算法的设计理念和创新点；2）技术架构的合理性和先进性；3）实现方案的可行性和高效性；4）与现有技术的对比优势。整体方案充分考虑了{research_field}领域的特殊需求。"
            },
            {
                "category": "技术方案与实现",
                "question": f"在'{system_name or '系统'}'的技术实现过程中，关键技术难点是什么？如何解决的？",
                "answer": "技术实现的关键难点包括：1）算法复杂度的优化；2）系统性能的提升；3）多模块的协调整合；4）异常情况的处理。通过深入的技术研究和大量的实验验证，成功解决了这些技术挑战。"
            }
        ],
        'innovation': [
            {
                "category": "创新点与贡献",
                "question": "您认为本研究的主要创新点和学术贡献是什么？",
                "answer": f"本研究的主要创新点体现在：1）理论层面的突破和发展；2）技术方法的创新和改进；3）应用模式的拓展和优化；4）问题解决思路的创新。这些创新为{research_field}领域的发展提供了新的思路和方法。"
            },
            {
                "category": "创新点与贡献",
                "question": "与现有相关工作相比，您的方法有哪些显著优势？",
                "answer": "与现有方法相比，本研究的优势包括：1）算法效率的显著提升；2）应用范围的扩展；3）实现复杂度的降低；4）结果准确性的改善。通过对比实验验证了这些优势的客观性和可靠性。"
            }
        ],
        'experiment': [
            {
                "category": "实验设计与结果",
                "question": "请介绍您的实验设计思路和主要实验结果？",
                "answer": f"实验设计基于科学严谨的原则：1）实验环境的构建和数据集的准备；2）评估指标的选择和基准方法的确定；3）实验方案的设计和参数的调优；4）结果分析和统计检验。实验结果充分验证了本方法在{research_field}领域的有效性。"
            },
            {
                "category": "实验设计与结果",
                "question": "您如何验证研究方法的有效性和可靠性？",
                "answer": "方法验证采用多重策略：1）对比实验验证相对优势；2）消融实验分析各组件贡献；3）鲁棒性测试验证稳定性；4）实际应用场景的验证。通过全面的实验评估，确保了研究成果的科学性和实用性。"
            }
        ],
        'system': [
            {
                "category": "系统架构与设计",
                "question": f"请详细介绍'{system_name or '系统'}'的整体架构设计和核心模块功能？",
                "answer": f"'{system_name or '系统'}'采用模块化设计架构：1）前端交互层负责用户界面和操作逻辑；2）业务逻辑层处理核心功能和算法；3）数据访问层管理数据存储和检索；4）系统服务层提供公共服务和接口。各模块间通过标准化接口进行通信。"
            },
            {
                "category": "系统架构与设计",
                "question": "系统的可扩展性和维护性如何保证？",
                "answer": "系统的可扩展性和维护性通过以下方式保证：1）采用松耦合的模块化设计；2）标准化的接口规范和数据格式；3）完善的日志记录和监控机制；4）详细的文档和代码注释。这些设计确保了系统的长期稳定运行。"
            }
        ],
        'theory': [
            {
                "category": "理论分析与证明",
                "question": "请从理论角度分析您方法的数学基础和收敛性？",
                "answer": f"从理论角度分析，本方法具有坚实的数学基础：1）基于{research_field}领域的经典理论；2）提供了完整的数学推导过程；3）分析了算法的收敛性质和复杂度；4）给出了理论性能边界。理论分析为实际应用提供了可靠的指导。"
            }
        ],
        'application': [
            {
                "category": "应用价值与前景",
                "question": "您的研究成果有哪些实际应用价值和商业化前景？",
                "answer": f"研究成果的应用价值体现在：1）解决了{research_field}领域的实际问题；2）提升了相关应用的性能和效率；3）降低了实施成本和技术门槛；4）为相关产业的发展提供了技术支撑。具有良好的商业化应用前景。"
            }
        ],
        'future': [
            {
                "category": "未来工作与展望",
                "question": "基于当前研究成果，您计划开展哪些后续工作？",
                "answer": f"后续工作计划包括：1）进一步优化算法性能和稳定性；2）扩展应用场景和适用范围；3）与其他{research_field}技术的融合研究；4）推动成果的产业化应用。这些工作将进一步推动该领域的技术进步。"
            }
        ]
    }
    
    # 获取指定类别的问题模板
    templates = category_templates.get(category, [])
    if not templates:
        return []
    
    # 根据问题数量返回相应的问题
    questions = []
    for i in range(min(question_count, len(templates))):
        questions.append(templates[i])
    
    # 如果需要更多问题，复制并修改现有模板
    while len(questions) < question_count:
        base_template = templates[len(questions) % len(templates)]
        modified_question = {
            "category": base_template["category"],
            "question": f"[补充问题] {base_template['question']}",
            "answer": base_template["answer"]
        }
        questions.append(modified_question)
    
    return questions[:question_count]

def generate_smart_fallback_questions(thesis_title, research_field, thesis_abstract, 
                                    system_name, question_count, difficulty_level, category=None):
    """智能备用问题生成方案 - 基于输入内容动态生成"""
    questions = []
    
    # 如果指定了类别，优先生成该类别的问题
    if category:
        category_questions = generate_category_specific_questions(
            thesis_title, research_field, thesis_abstract, system_name, 
            question_count, difficulty_level, category
        )
        if category_questions:
            return category_questions
    
    # 基础问题模板
    base_questions = [
        {
            "category": "研究背景与意义",
            "question": f"请详细阐述您选择'{thesis_title}'这个研究课题的背景、现实中存在的问题以及研究意义？",
            "answer": f"本研究针对{research_field}领域的实际需求，解决了传统方法中存在的局限性。主要研究背景包括：1）当前技术发展趋势和行业需求；2）现有方案的不足之处；3）本研究的创新价值和应用前景。通过深入分析相关文献和实际应用场景，确定了研究的必要性和重要性。"
        },
        {
            "category": "技术方案与实现",
            "question": "请详细介绍您采用的主要技术方案、算法原理和具体实现方法？",
            "answer": f"本研究采用了先进的技术架构，主要包括：1）核心算法设计和优化策略；2）系统架构的设计原则和模块划分；3）关键技术的选择依据和实现细节；4）性能优化和可扩展性考虑。技术方案充分考虑了实际应用需求，确保了系统的稳定性和高效性。"
        },
        {
            "category": "创新点与贡献",
            "question": "您认为本研究的主要创新点是什么？与现有相关工作相比有哪些优势？",
            "answer": f"本研究的主要创新点体现在：1）提出了新的理论模型或算法改进；2）在技术实现上采用了创新的方法；3）解决了现有方案中的关键问题；4）在应用层面实现了突破性进展。通过对比实验和性能分析，验证了本方案相比传统方法的显著优势。"
        },
        {
            "category": "实验设计与结果",
            "question": "请介绍您的实验设计思路、评估指标选择和主要实验结果？",
            "answer": f"实验设计遵循科学严谨的原则：1）构建了完整的实验环境和数据集；2）设计了合理的对比实验方案；3）选择了客观有效的评估指标；4）通过多轮实验验证了方法的有效性。实验结果表明，本方法在关键指标上取得了显著提升，验证了理论分析的正确性。"
        }
    ]
    
    # 根据系统名称添加系统相关问题
    if system_name:
        system_questions = [
            {
                "category": "系统设计与架构",
                "question": f"请详细介绍'{system_name}'的整体架构设计、核心模块功能和技术选型依据？",
                "answer": f"'{system_name}'采用模块化设计理念，主要包括：1）前端用户界面设计和交互逻辑；2）后端核心业务逻辑和数据处理；3）数据存储和管理策略；4）系统集成和部署方案。技术选型充分考虑了性能、稳定性和可维护性要求。"
            },
            {
                "category": "系统实现与优化",
                "question": f"在'{system_name}'的开发过程中，您遇到了哪些技术难点？是如何解决的？",
                "answer": f"系统开发过程中主要面临以下挑战：1）性能优化和并发处理问题；2）数据一致性和安全性保障；3）用户体验和界面优化；4）系统稳定性和错误处理。通过采用先进的技术方案和最佳实践，成功解决了这些关键问题。"
            }
        ]
        base_questions.extend(system_questions)
    
    # 根据摘要内容智能调整问题
    if "机器学习" in thesis_abstract or "深度学习" in thesis_abstract or "AI" in thesis_abstract:
        ai_questions = [
            {
                "category": "算法原理与优化",
                "question": "请详细说明您使用的机器学习/深度学习算法的原理、网络结构设计和训练策略？",
                "answer": "本研究采用的算法具有以下特点：1）网络架构设计考虑了任务特性和数据特征；2）损失函数和优化器的选择基于充分的理论分析；3）训练策略包括数据增强、正则化和超参数调优；4）模型评估采用了多种指标和验证方法，确保了结果的可靠性。"
            },
            {
                "category": "模型评估与分析",
                "question": "您如何评估模型的性能？在什么数据集上进行了验证？结果如何解释？",
                "answer": "模型评估采用了科学严谨的方法：1）使用了标准的评估指标和基准数据集；2）进行了充分的对比实验和消融实验；3）分析了模型的泛化能力和鲁棒性；4）对结果进行了深入的理论分析和解释，验证了方法的有效性和可靠性。"
            }
        ]
        base_questions.extend(ai_questions)
    
    # 根据难度级别调整问题复杂度
    if difficulty_level == "advanced":
        advanced_questions = [
            {
                "category": "理论深度分析",
                "question": "请从理论角度深入分析您的方法的数学基础、收敛性证明和复杂度分析？",
                "answer": "从理论角度分析，本方法具有坚实的数学基础：1）提供了完整的理论推导和证明过程；2）分析了算法的收敛性质和收敛速度；3）给出了时间和空间复杂度的详细分析；4）讨论了方法的理论局限性和适用范围，为实际应用提供了理论指导。"
            },
            {
                "category": "未来发展方向",
                "question": "基于当前研究成果，您认为该领域未来的发展趋势是什么？您的工作如何推动领域发展？",
                "answer": "基于本研究成果，该领域的发展趋势包括：1）技术方法的不断创新和优化；2）应用场景的扩展和深化；3）与其他领域的交叉融合；4）理论体系的进一步完善。本工作为后续研究提供了新的思路和方法，推动了领域的技术进步和理论发展。"
            }
        ]
        base_questions.extend(advanced_questions)
    
    # 返回指定数量的问题
    selected_questions = base_questions[:min(question_count, len(base_questions))]
    
    # 如果问题不够，补充通用问题
    if len(selected_questions) < question_count:
        generic_questions = generate_default_defense_questions(thesis_title, research_field, 
                                                             question_count - len(selected_questions))
        selected_questions.extend(generic_questions)
    
    return selected_questions[:question_count]

def generate_fallback_answer(thesis_title, question):
    """为问题生成备用答案"""
    return f"针对'{question}'这个问题，可以从以下几个方面进行回答：1）结合论文'{thesis_title}'的核心内容和创新点；2）分析相关的理论基础和技术方法；3）说明实际应用价值和意义；4）总结研究成果和未来展望。建议结合具体的研究内容进行详细阐述，展现对研究领域的深入理解和专业能力。"

def generate_default_defense_questions(thesis_title, research_field, question_count):
    """生成默认的答辩问题（当AI调用失败时使用）"""
    default_questions = [
        {
            "category": "研究背景",
            "question": f"请简要介绍您选择'{thesis_title}'这个研究课题的背景和意义？",
            "answer": "可以从以下几个方面回答：1）当前该领域存在的问题或挑战；2）研究该问题的重要性和必要性；3）预期的研究成果对学术界或工业界的贡献；4）个人的研究兴趣和专业背景。"
        },
        {
            "category": "技术方案",
            "question": "请详细说明您在研究中采用的主要技术方案和实现方法？",
            "answer": "应该包括：1）总体技术架构设计；2）关键技术选择的依据；3）具体的实现步骤和方法；4）技术方案的创新点和优势；5）与现有方案的对比分析。"
        },
        {
            "category": "创新点",
            "question": "您认为本研究的主要创新点和贡献是什么？",
            "answer": "可以从以下角度阐述：1）理论创新：提出了新的理论模型或算法；2）技术创新：采用了新的技术手段或改进了现有技术；3）应用创新：在新的应用场景中解决了实际问题；4）方法创新：提出了新的研究方法或评估标准。"
        },
        {
            "category": "实验验证",
            "question": "请介绍您的实验设计和主要实验结果？",
            "answer": "应该包括：1）实验环境和数据集的选择；2）评估指标的设定和合理性；3）实验结果的详细分析；4）与基线方法的对比；5）结果的可靠性和统计显著性分析。"
        },
        {
            "category": "技术难点",
            "question": "在研究过程中遇到的主要技术难点是什么？您是如何解决的？",
            "answer": "可以描述：1）具体遇到的技术挑战；2）问题分析和解决思路；3）尝试过的不同方案；4）最终采用的解决方案及其效果；5）从中获得的经验和教训。"
        },
        {
            "category": "相关工作",
            "question": "请比较您的工作与相关研究的异同点？",
            "answer": "应该包括：1）相关工作的梳理和分类；2）现有方法的优缺点分析；3）本研究与现有工作的区别和改进；4）在相关工作基础上的创新和发展；5）未来可能的研究方向。"
        },
        {
            "category": "应用前景",
            "question": "您的研究成果有哪些实际应用价值和推广前景？",
            "answer": "可以从以下方面回答：1）直接的应用场景和目标用户；2）解决的实际问题和带来的效益；3）产业化的可能性和商业价值；4）推广应用的条件和挑战；5）对相关行业的潜在影响。"
        },
        {
            "category": "不足与改进",
            "question": "您认为当前研究还存在哪些不足？未来如何改进？",
            "answer": "应该诚实地分析：1）当前方案的局限性和不足；2）实验验证的不完善之处；3）理论分析的深度有待提高的方面；4）未来的改进方向和计划；5）长期的研究目标和愿景。"
        },
        {
            "category": "系统实现",
            "question": "如果涉及系统开发，请介绍系统的整体架构和关键模块？",
            "answer": "可以介绍：1）系统的总体架构设计；2）各个功能模块的职责和接口；3）关键技术的实现细节；4）系统的性能指标和优化策略；5）系统的可扩展性和维护性考虑。"
        },
        {
            "category": "未来工作",
            "question": "基于当前的研究成果，您计划开展哪些后续工作？",
            "answer": "可以规划：1）短期内可以完成的改进工作；2）中长期的研究目标和计划；3）可能的合作方向和资源需求；4）研究成果的进一步验证和完善；5）向更广泛应用领域的扩展。"
        }
    ]

    # 根据请求的数量返回相应的问题
    return default_questions[:min(question_count, len(default_questions))]


def generate_defense_questions_word(questions, thesis_title, research_field):
    """生成答辩问题Word文档"""
    doc = Document()

    # 设置文档标题
    title = doc.add_heading(f'论文答辩问题及参考答案', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加基本信息
    info_para = doc.add_paragraph()
    info_para.add_run('论文题目：').bold = True
    info_para.add_run(thesis_title)
    info_para.add_run('\n研究领域：').bold = True
    info_para.add_run(research_field)
    info_para.add_run('\n生成时间：').bold = True
    info_para.add_run(datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
    info_para.add_run('\n问题数量：').bold = True
    info_para.add_run(str(len(questions)))

    # 添加空行
    doc.add_paragraph()

    # 添加问题和答案
    for i, question in enumerate(questions, 1):
        # 问题标题
        question_heading = doc.add_heading(f'问题{i}：{question.get("category", "综合问题")}', level=2)

        # 问题内容
        question_para = doc.add_paragraph()
        question_para.add_run('问题：').bold = True
        question_para.add_run(question.get('question', ''))

        # 答案内容
        answer_para = doc.add_paragraph()
        answer_para.add_run('参考答案：').bold = True
        answer_para.add_run(question.get('answer', ''))

        # 添加分隔线（除了最后一个问题）
        if i < len(questions):
            doc.add_paragraph('─' * 50)

    # 保存到内存
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)

    return doc_buffer


# ==================== 答辩问题历史记录API ====================

@app.route('/api/defense-question-history')
@login_required
def api_get_defense_question_history():
    """获取用户的答辩问题历史记录"""
    try:
        page = int(request.args.get('page', 1))
        per_page = int(request.args.get('per_page', 20))
        
        result = user_manager.get_defense_question_history(
            user_id=session['user_id'],
            page=page,
            per_page=per_page
        )
        
        if result:
            return jsonify({
                'success': True,
                'data': result
            })
        else:
            return jsonify({
                'success': True,
                'data': {
                    'records': [],
                    'total': 0,
                    'page': page,
                    'per_page': per_page,
                    'pages': 0
                }
            })
            
    except Exception as e:
        app.logger.error(f"获取历史记录失败: {e}")
        return jsonify({'success': False, 'message': '获取历史记录失败'}), 500

@app.route('/api/defense-question-history/<int:record_id>')
@login_required
def api_get_defense_question_detail(record_id):
    """获取历史记录详情"""
    try:
        record = user_manager.get_defense_question_detail(
            user_id=session['user_id'],
            record_id=record_id
        )
        
        if record:
            return jsonify({
                'success': True,
                'record': record
            })
        else:
            return jsonify({'success': False, 'message': '记录不存在'}), 404
            
    except Exception as e:
        app.logger.error(f"获取历史记录详情失败: {e}")
        return jsonify({'success': False, 'message': '获取记录详情失败'}), 500

@app.route('/api/defense-question-history/<int:record_id>', methods=['DELETE'])
@login_required
def api_delete_defense_question_history(record_id):
    """删除历史记录"""
    try:
        success = user_manager.delete_defense_question_history(
            user_id=session['user_id'],
            record_id=record_id
        )
        
        if success:
            return jsonify({
                'success': True,
                'message': '记录删除成功'
            })
        else:
            return jsonify({'success': False, 'message': '记录不存在或删除失败'}), 404
            
    except Exception as e:
        app.logger.error(f"删除历史记录失败: {e}")
        return jsonify({'success': False, 'message': '删除记录失败'}), 500

@app.route('/api/defense-question-history/clear', methods=['POST'])
@login_required
def api_clear_defense_question_history():
    """清空用户所有历史记录"""
    try:
        count = user_manager.clear_defense_question_history(
            user_id=session['user_id']
        )
        
        return jsonify({
            'success': True,
            'message': f'已清空{count}条历史记录'
        })
            
    except Exception as e:
        app.logger.error(f"清空历史记录失败: {e}")
        return jsonify({'success': False, 'message': '清空记录失败'}), 500












# SEO优化路由
@app.route('/sitemap.xml')
def sitemap():
    """提供sitemap.xml文件"""
    return send_file('static/sitemap.xml', mimetype='application/xml')

@app.route('/robots.txt')
def robots():
    """提供robots.txt文件"""
    return send_file('static/robots.txt', mimetype='text/plain')

# 添加结构化数据API
@app.route('/api/schema')
def schema_data():
    """提供结构化数据"""
    schema = {
        "@context": "https://schema.org",
        "@type": "WebApplication",
        "name": "智能文档处理平台",
        "description": "专业的学术工具集，提供SQL转ER图、测试用例生成等功能",
        "applicationCategory": "EducationalApplication",
        "operatingSystem": "Web Browser",
        "url": request.host_url,
        "offers": {
            "@type": "Offer",
            "price": "0",
            "priceCurrency": "CNY"
        },
        "featureList": [
            "SQL转ER图生成",
            "测试用例自动生成",
            "论文答辩助手",
            "AI文本优化",
            "数据库设计工具"
        ],
        "audience": {
            "@type": "Audience",
            "audienceType": ["学生", "研究人员", "软件开发者", "数据库设计师"]
        },
        "provider": {
            "@type": "Organization",
            "name": "智能文档处理平台"
        }
    }
    return jsonify(schema)


def generate_paper_with_citations_background(task_id, title, field, paper_type, abstract, keywords, requirements, custom_outline):
    """带文献搜索和引用的论文生成后台任务 - 增强版本，支持上下文记忆"""
    try:
        app.logger.info(f"开始带文献引用的论文生成: {title}")

        # 初始化记忆系统
        memory = paper_generation_tasks[task_id]['memory']

        # 提取用户要求的关键信息
        user_context = extract_user_requirements_context(requirements, abstract, keywords)
        memory['global_context'].update(user_context)

        app.logger.info(f"用户要求分析完成: {user_context}")

        # 更新任务状态
        paper_generation_tasks[task_id].update({
            'status': 'searching_literature',
            'progress': 5,
            'message': '正在搜索相关学术文献...',
            'total_sections': len(custom_outline),
            'memory': memory
        })

        # 第一步：搜索学术文献
        search_keywords = f"{title} {field} {keywords}".strip()
        literature_list = search_academic_literature(field, title, search_keywords)

        if literature_list:
            app.logger.info(f"搜索到 {len(literature_list)} 篇相关文献")
            paper_generation_tasks[task_id].update({
                'progress': 15,
                'message': f'文献搜索完成，找到 {len(literature_list)} 篇相关文献',
                'literature_count': len(literature_list)
            })
        else:
            app.logger.warning("未搜索到相关文献，将使用AI生成")
            paper_generation_tasks[task_id].update({
                'progress': 15,
                'message': '未找到相关文献，将使用AI生成参考文献',
                'literature_count': 0
            })

        # 第二步：生成论文内容
        complete_content = f'<h1 style="text-align: center; margin-bottom: 30px;">{title}</h1>\n\n'

        # 逐章节生成内容，使用记忆系统
        for i, section in enumerate(custom_outline):
            try:
                # 更新进度
                progress = 20 + (i * 70 // len(custom_outline))
                paper_generation_tasks[task_id].update({
                    'current_section': section['name'],
                    'sections_completed': i,
                    'progress': progress,
                    'message': f'正在生成 {section["name"]}...',
                    'memory': memory
                })

                app.logger.info(f"开始生成章节 {i+1}/{len(custom_outline)}: {section['name']}")

                # 生成章节内容
                if section['name'] == "参考文献":
                    # 基于搜索结果生成参考文献
                    section_content, _ = generate_references_with_search(field, title, search_keywords)
                else:
                    # 使用上下文感知的章节生成
                    section_content = generate_section_with_memory(
                        section, memory, i, literature_list
                    )

                if section_content and len(section_content.strip()) > 50:
                    complete_content += section_content + "\n\n"

                    # 提取章节上下文并更新记忆
                    if section['name'] != "参考文献":
                        section_context = extract_section_context(section_content, section['name'])

                        # 更新记忆系统
                        memory['generated_sections'].append({
                            'name': section['name'],
                            'content': section_content,
                            'summary': section_context.get('main_points', []),
                            'tech_decisions': section_context.get('tech_decisions', []),
                            'key_terms': section_context.get('key_terms', {}),
                            'pending_points': section_context.get('pending_points', []),
                            'continuation_needs': section_context.get('continuation_needs', [])
                        })

                        # 更新全局术语库
                        memory['key_terms'].update(section_context.get('key_terms', {}))

                        # 更新技术决策记录
                        memory['technical_decisions'].extend(section_context.get('tech_decisions', []))

                        # 更新累积内容（保留最近的内容用于上下文）
                        memory['accumulated_content'] += section_content
                        if len(memory['accumulated_content']) > 15000:
                            memory['accumulated_content'] = memory['accumulated_content'][-12000:]

                        app.logger.info(f"章节 {section['name']} 生成成功，记忆已更新")
                        app.logger.info(f"当前记忆状态 - 章节数: {len(memory['generated_sections'])}, 术语数: {len(memory['key_terms'])}")
                    else:
                        app.logger.info(f"参考文献章节生成成功")
                else:
                    app.logger.error(f"章节 {section['name']} 生成失败，跳过该章节")
                    continue

                # 更新已完成章节数
                paper_generation_tasks[task_id].update({
                    'sections_completed': i + 1,
                    'message': f'{section["name"]} 章节生成完成',
                    'memory': memory
                })

                # 适当延迟，确保生成质量
                time.sleep(3)

            except Exception as section_error:
                app.logger.error(f"生成章节 {section['name']} 时出错: {section_error}")
                continue

        # 生成完成
        paper_generation_tasks[task_id].update({
            'status': 'completed',
            'progress': 100,
            'content': complete_content,
            'message': '带文献引用的论文生成完成！',
            'literature_list': literature_list
        })

        app.logger.info(f"带文献引用的论文生成完成: {title}")

    except Exception as e:
        app.logger.error(f"带文献引用的论文生成过程中发生错误: {e}")
        paper_generation_tasks[task_id].update({
            'status': 'error',
            'error': str(e),
            'message': f'生成失败: {str(e)}'
        })


def generate_paper_background(task_id, title, field, paper_type, abstract, keywords, requirements, custom_outline):
    """普通论文生成后台任务（不带文献搜索）- 增强版本，支持上下文记忆"""
    try:
        app.logger.info(f"开始普通论文生成: {title}")

        # 初始化记忆系统
        memory = paper_generation_tasks[task_id]['memory']

        # 提取用户要求的关键信息
        user_context = extract_user_requirements_context(requirements, abstract, keywords)
        memory['global_context'].update(user_context)

        app.logger.info(f"用户要求分析完成: {user_context}")

        # 更新任务状态
        paper_generation_tasks[task_id].update({
            'status': 'generating',
            'progress': 10,
            'message': '正在生成论文内容...',
            'total_sections': len(custom_outline),
            'memory': memory
        })

        # 生成完整论文内容
        complete_content = f'<h1 style="text-align: center; margin-bottom: 30px;">{title}</h1>\n\n'

        # 逐章节生成内容
        for i, section in enumerate(custom_outline):
            try:
                # 更新进度
                progress = 15 + (i * 80 // len(custom_outline))
                paper_generation_tasks[task_id].update({
                    'current_section': section['name'],
                    'sections_completed': i,
                    'progress': progress,
                    'message': f'正在生成 {section["name"]}...'
                })

                app.logger.info(f"开始生成章节: {section['name']}")

                # 生成章节内容
                if section['name'] == "参考文献":
                    # 使用收集的文献引用生成参考文献
                    section_content = generate_collected_references(memory)
                else:
                    # 使用上下文感知的章节生成
                    section_content = generate_section_with_memory(
                        section, memory, i, None  # 普通生成不带文献引用
                    )

                if section_content and len(section_content.strip()) > 50:
                    complete_content += section_content + "\n\n"

                    # 提取章节上下文并更新记忆
                    if section['name'] != "参考文献":
                        section_context = extract_section_context(section_content, section['name'])

                        # 更新记忆系统
                        memory['generated_sections'].append({
                            'name': section['name'],
                            'content': section_content,
                            'summary': section_context.get('main_points', []),
                            'tech_decisions': section_context.get('tech_decisions', []),
                            'key_terms': section_context.get('key_terms', {}),
                            'pending_points': section_context.get('pending_points', []),
                            'continuation_needs': section_context.get('continuation_needs', [])
                        })

                        # 更新全局术语库
                        memory['key_terms'].update(section_context.get('key_terms', {}))

                        # 更新技术决策记录
                        memory['technical_decisions'].extend(section_context.get('tech_decisions', []))

                        # 更新累积内容
                        memory['accumulated_content'] += section_content
                        if len(memory['accumulated_content']) > 15000:
                            memory['accumulated_content'] = memory['accumulated_content'][-12000:]

                        app.logger.info(f"章节 {section['name']} 生成成功，记忆已更新")
                    else:
                        app.logger.info(f"参考文献章节生成成功")
                else:
                    app.logger.error(f"章节 {section['name']} 生成失败，跳过该章节")
                    continue

                # 更新已完成章节数
                paper_generation_tasks[task_id].update({
                    'sections_completed': i + 1,
                    'message': f'{section["name"]} 章节生成完成',
                    'memory': memory
                })

                # 适当延迟，确保生成质量
                time.sleep(3)

            except Exception as section_error:
                app.logger.error(f"生成章节 {section['name']} 时出错: {section_error}")
                continue

        # 生成完成
        paper_generation_tasks[task_id].update({
            'status': 'completed',
            'progress': 100,
            'content': complete_content,
            'message': '高质量论文生成完成！',
            'memory': memory
        })

        app.logger.info(f"普通论文生成完成: {title}")
        app.logger.info(f"最终记忆状态 - 章节数: {len(memory['generated_sections'])}, 术语数: {len(memory['key_terms'])}, 技术决策数: {len(memory['technical_decisions'])}")

    except Exception as e:
        app.logger.error(f"普通论文生成过程中发生错误: {e}")
        paper_generation_tasks[task_id].update({
            'status': 'error',
            'error': str(e),
            'message': f'生成失败: {str(e)}'
        })


def generate_section_with_memory(section, memory, section_index, literature_list=None):
    """使用记忆系统生成章节内容 - 支持上下文感知"""
    try:
        section_name = section['name']
        section_words = section.get('words', 1000)

        app.logger.info(f"开始生成章节: {section_name}, 目标字数: {section_words}")

        # 构建上下文感知的提示词
        contextual_prompt = build_contextual_prompt(section, memory, section_index)

        # 添加文献引用指导（如果有文献列表）
        if literature_list:
            relevant_refs = [f"[{i}]" for i in range(1, min(16, len(literature_list) + 1))]
            citation_guide = f"""
【文献引用要求】：
- 在适当位置添加文献引用，使用上标格式：<sup>[序号]</sup>
- 可用的引用标记：{', '.join(relevant_refs[:10])}
- 引用位置：理论阐述后、技术方法介绍时、研究现状分析中
- 引用示例：相关研究表明<sup>[1]</sup>，该技术在实际应用中<sup>[2,3]</sup>
- 每段至少包含1-2个文献引用，确保学术性
"""
            contextual_prompt += citation_guide

        # 计算合适的token数量
        max_tokens = min(section_words * 4, 8000)  # 确保有足够的token生成充实内容

        app.logger.info(f"调用AI生成章节，max_tokens: {max_tokens}")

        # 调用AI生成内容
        section_content = call_deepseek_api(contextual_prompt, max_tokens)

        if section_content and len(section_content.strip()) > 100:
            cleaned_content = clean_ai_generated_content(section_content)
            app.logger.info(f"章节 {section_name} 生成成功，内容长度: {len(cleaned_content)}")
            return cleaned_content
        else:
            app.logger.warning(f"章节 {section_name} 生成内容不足，尝试重新生成")
            # 尝试重新生成一次
            retry_content = call_deepseek_api(contextual_prompt, max_tokens)
            if retry_content:
                cleaned_retry = clean_ai_generated_content(retry_content)
                app.logger.info(f"章节 {section_name} 重新生成成功")
                return cleaned_retry
            else:
                app.logger.error(f"章节 {section_name} 重新生成也失败")
                return None

    except Exception as e:
        app.logger.error(f"生成章节 {section_name} 失败: {e}")
        return None


def generate_section_with_citations(section, title, field, paper_type, abstract, keywords, requirements, literature_list):
    """生成带文献引用的章节内容 - 保留原有接口兼容性"""
    try:
        section_name = section['name']
        section_words = section.get('words', 1000)
        section_desc = section.get('description', '')
        level = section.get('level', 2)

        # 为章节选择相关文献引用
        relevant_refs = []
        if literature_list:
            # 根据章节内容选择相关文献（简单实现）
            for i, lit in enumerate(literature_list[:15], 1):
                relevant_refs.append(f"[{i}]")

        # 构建引用指导
        citation_guide = ""
        if relevant_refs:
            citation_guide = f"""
【文献引用要求】：
- 在适当位置添加文献引用，使用上标格式：<sup>[序号]</sup>
- 可用的引用标记：{', '.join(relevant_refs[:10])}
- 引用位置：理论阐述后、技术方法介绍时、研究现状分析中
- 引用示例：相关研究表明<sup>[1]</sup>，该技术在实际应用中<sup>[2,3]</sup>
- 每段至少包含1-2个文献引用，确保学术性
"""

        # 调用带引用的章节生成函数
        return generate_simple_section_content_with_citations(
            title, field, paper_type, section, abstract, keywords, requirements,
            section_name, citation_guide
        )

    except Exception as e:
        app.logger.error(f"生成带引用的章节内容失败: {e}")
        # 回退到普通生成
        return generate_simple_section_content(
            title, field, paper_type, section, abstract, keywords, requirements, 1
        )


def generate_simple_section_content(title, field, paper_type, section, abstract, keywords, requirements, section_num, memory=None):
    """核心章节生成函数 - 支持细粒度生成和引用连续性"""
    try:
        section_name = section['name']
        section_words = section.get('words', 1500)
        section_desc = section.get('description', '')
        level = section.get('level', 2)
        header_tag = f"h{level}"

        # 获取当前引用计数器
        current_ref_counter = memory.get('reference_counter', 0) if memory else 0
        
        app.logger.info(f"生成章节: {section_name}, 目标字数: {section_words}, 当前引用计数: {current_ref_counter}")

        # 获取用户上下文信息
        user_context = memory.get('system_context', {}) if memory else {}
        system_info = ""
        if user_context:
            system_info = f"""
【用户系统具体信息】：
- 系统名称：{title}
- 技术栈：{user_context.get('tech_stack', 'Spring Boot + Vue.js + MySQL')}
- 数据库信息：{user_context.get('database_info', 'MySQL关系型数据库')}
- 核心功能：{user_context.get('key_features', '信息管理、数据统计、权限控制')}
- 研究目标：{user_context.get('research_objectives', '提升管理效率，实现数字化转型')}
"""

        # 检查是否需要分小节生成（用户要求的细粒度生成）
        if should_generate_by_subsections(section_name, section_words):
            return generate_section_by_subsections(
                title, field, paper_type, section, abstract, keywords, 
                requirements, memory, system_info
            )

        # 构建章节生成提示词
        prompt = build_enhanced_section_prompt(
            title, field, paper_type, section_name, section_words, 
            section_desc, header_tag, system_info, abstract, keywords, 
            requirements, current_ref_counter
        )

        # 计算合适的token数量 - 大幅增加以满足字数要求
        # 根据用户反馈，进一步大幅增加token分配
        base_tokens = section_words * 20  # 从12倍提升到20倍
        max_tokens = min(base_tokens, 32000)  # 最大token限制提升到32000
        
        app.logger.info(f"调用API生成章节，max_tokens: {max_tokens}")

        # 调用AI生成内容
        content = call_deepseek_api(prompt, max_tokens)
        
        if content and len(content.strip()) > 100:
            # 处理引用编号连续性
            processed_content = process_references_in_content(content, memory, current_ref_counter + 1)
            
            # 验证字数达标
            cleaned_content = clean_ai_generated_content(processed_content)
            text_content = re.sub(r'<[^>]+>', '', cleaned_content)
            actual_words = len(text_content.replace(' ', '').replace('\n', ''))
            
            app.logger.info(f"章节 {section_name} 生成完成，实际字数: {actual_words}/{section_words}")
            
            # 如果字数不足，尝试扩展内容
            if actual_words < section_words * 0.8:
                cleaned_content = enhance_content_for_word_count(
                    cleaned_content, section_name, section_words, actual_words
                )
            
            return cleaned_content
        else:
            app.logger.warning(f"章节 {section_name} 生成内容不足，使用备用方案")
            return generate_fallback_section(section, section_num)
            
    except Exception as e:
        app.logger.error(f"生成章节 {section_name} 失败: {e}")
        return generate_fallback_section(section, section_num)


def should_generate_by_subsections(section_name, section_words):
    """判断是否需要按小节生成（实现用户要求的细粒度生成）"""
    # 大幅扩大小节生成范围，确保更多API调用
    main_sections = ["第1章", "第2章", "第3章", "第4章", "第5章", "第6章", "第7章",
                    "绪论", "引言", "技术", "需求", "设计", "实现", "测试", "总结", "展望",
                    "相关", "分析", "系统", "功能", "模块", "架构", "数据库", "接口"]
    # 降低字数门槛，让更多章节使用小节生成
    return any(keyword in section_name for keyword in main_sections) and section_words > 500


def generate_section_by_subsections(title, field, paper_type, section, abstract, keywords, requirements, memory, system_info):
    """按小节生成章节内容 - 实现用户要求的4.1、4.2单独调用API"""
    try:
        section_name = section['name']
        section_words = section.get('words', 1500)
        
        # 确定小节结构
        subsections = get_subsection_structure(section_name, section_words)
        
        complete_content = f"<h2>{section_name}</h2>\n\n"
        current_ref_counter = memory.get('reference_counter', 0) if memory else 0
        
        app.logger.info(f"开始按小节生成 {section_name}，共 {len(subsections)} 个小节")
        
        for i, subsection in enumerate(subsections):
            try:
                app.logger.info(f"生成小节 {subsection['name']}，目标字数: {subsection['words']}")
                
                # 为每个小节单独调用API - 确保引用编号连续
                subsection_ref_start = memory.get('reference_counter', 0) + 1 if memory else 1
                subsection_content = generate_single_subsection(
                    title, field, paper_type, subsection, abstract, keywords,
                    requirements, system_info, subsection_ref_start
                )

                if subsection_content:
                    # 处理引用编号连续性
                    processed_content = process_references_in_content(subsection_content, memory, subsection_ref_start)
                    complete_content += processed_content + "\n\n"

                    # 更新引用计数器 - 确保连续性
                    if memory:
                        new_refs = extract_reference_count_from_content(processed_content)
                        memory['reference_counter'] = memory.get('reference_counter', 0) + new_refs
                        
                else:
                    app.logger.warning(f"小节 {subsection['name']} 生成失败")
                    
            except Exception as subsection_error:
                app.logger.error(f"生成小节 {subsection['name']} 时出错: {subsection_error}")
                continue
        
        return complete_content
        
    except Exception as e:
        app.logger.error(f"按小节生成章节失败: {e}")
        return f"<h2>{section['name']}</h2>\n<p>内容生成失败，请重试。</p>"


def get_subsection_structure(section_name, total_words):
    """获取章节的小节结构"""
    if "第1章" in section_name or "绪论" in section_name:
        return [
            {"name": "1.1 研究背景与意义", "words": int(total_words * 0.35)},
            {"name": "1.2 国内外研究现状", "words": int(total_words * 0.35)},
            {"name": "1.3 研究内容与方法", "words": int(total_words * 0.20)},
            {"name": "1.4 论文组织结构", "words": int(total_words * 0.10)}
        ]
    elif "第2章" in section_name or "技术" in section_name:
        return [
            {"name": "2.1 开发框架技术", "words": int(total_words * 0.25)},
            {"name": "2.2 数据库技术", "words": int(total_words * 0.25)},
            {"name": "2.3 前端开发技术", "words": int(total_words * 0.25)},
            {"name": "2.4 系统架构设计", "words": int(total_words * 0.25)}
        ]
    elif "第3章" in section_name or "需求" in section_name or "设计" in section_name:
        return [
            {"name": "3.1 需求分析", "words": int(total_words * 0.35)},
            {"name": "3.2 系统总体设计", "words": int(total_words * 0.35)},
            {"name": "3.3 数据库设计", "words": int(total_words * 0.30)}
        ]
    elif "第4章" in section_name or "实现" in section_name:
        return [
            {"name": "4.1 系统功能模块设计", "words": int(total_words * 0.30)},
            {"name": "4.2 关键技术实现", "words": int(total_words * 0.35)},
            {"name": "4.3 系统安全设计", "words": int(total_words * 0.35)}
        ]
    elif "第5章" in section_name or "测试" in section_name:
        return [
            {"name": "5.1 测试环境与方法", "words": int(total_words * 0.30)},
            {"name": "5.2 功能测试", "words": int(total_words * 0.40)},
            {"name": "5.3 性能测试与分析", "words": int(total_words * 0.30)}
        ]
    else:
        # 默认分为3个小节
        words_per_section = total_words // 3
        return [
            {"name": f"{section_name.split(' ')[0]}.1 概述", "words": words_per_section},
            {"name": f"{section_name.split(' ')[0]}.2 详细内容", "words": words_per_section},
            {"name": f"{section_name.split(' ')[0]}.3 总结", "words": words_per_section}
        ]


def generate_single_subsection(title, field, paper_type, subsection, abstract, keywords, requirements, system_info, ref_start_num):
    """生成单个小节内容 - 单独API调用"""
    try:
        subsection_name = subsection['name']
        target_words = subsection['words']
        
        prompt = f"""请为{field}领域的论文《{title}》生成{subsection_name}小节内容。

目标字数：{target_words}字（必须达到）
{system_info}

【核心要求】：
1. 内容必须达到{target_words}字，不得少于目标字数的90%
2. 必须包含3-5个文献引用，引用编号从[{ref_start_num}]开始连续编号
3. 每个段落300-400字，共需要{target_words//350 + 1}个段落
4. 内容要具体、深入、专业，避免空洞表述
5. 必须紧密结合{title}系统的特点
6. 充分利用用户提供的系统信息和技术栈

【写作要求】：
- 使用学术论文的正式语言
- 每段必须包含1-2个引用标注
- 段落间要有逻辑递进关系
- 避免重复和冗余表达
- 数据和结论要有说服力

【输出格式】：
<h3>{subsection_name}</h3>
<p>第一段内容，至少300字...引用[{ref_start_num}]</p>
<p>第二段内容，至少300字...引用[{ref_start_num+1}]</p>
...

在内容最后添加本小节的参考文献：
<div class="temp-references">
[{ref_start_num}] 作者1. 相关研究1[J]. 期刊名, 年份, 卷(期): 页码.
[{ref_start_num+1}] 作者2. 相关研究2[C]. 会议名, 年份: 页码.
...
</div>

请确保内容质量高、字数充足、引用规范。"""

        # 为小节使用更高的token限制 - 根据用户反馈大幅增加
        base_tokens = target_words * 25  # 从15倍提升到25倍
        max_tokens = min(base_tokens, 40000)  # 最大token限制提升到40000
        
        content = call_deepseek_api(prompt, max_tokens)
        
        if content and len(content.strip()) > 100:
            return clean_ai_generated_content(content)
        else:
            return generate_fallback_subsection_content(subsection_name, target_words)
            
    except Exception as e:
        app.logger.error(f"生成小节 {subsection_name} 失败: {e}")
        return generate_fallback_subsection_content(subsection_name, target_words)


def generate_fallback_subsection_content(subsection_name, target_words):
    """生成备用小节内容"""
    return f"""
<h3>{subsection_name}</h3>
<div class="content-placeholder" style="background: #f8f9fa; border: 1px solid #dee2e6; padding: 20px; border-radius: 5px;">
<p><strong>⚠️ 内容生成异常</strong></p>
<p>目标字数：{target_words}字</p>
<p>建议手动补充以下内容：</p>
<ul>
<li>相关理论基础和技术背景</li>
<li>具体的实现方案和技术细节</li>
<li>数据分析和实证研究结果</li>
<li>与其他研究的比较分析</li>
</ul>
</div>"""


def build_enhanced_section_prompt(title, field, paper_type, section_name, section_words, section_desc, header_tag, system_info, abstract, keywords, requirements, current_ref_counter):
    """构建增强的章节生成提示词"""
    base_prompt = f"""请为{field}领域的{paper_type}《{title}》生成{section_name}章节。

研究领域：{field}
目标字数：{section_words}字（必须严格达到）
章节描述：{section_desc}
{system_info}

【严格字数要求】：
1. 内容字数必须达到{section_words}字，不得少于{int(section_words * 0.95)}字
2. 包含5-8个文献引用，从[{current_ref_counter + 1}]开始连续编号
3. 分为{max(4, section_words//400)}个段落，每段350-450字
4. 内容要学术化、专业化，避免空话套话
5. 必须紧密结合{title}系统的具体特点
6. 充分利用用户提供的技术栈和系统信息

【内容质量要求】：
- 使用准确的学术术语和专业表达
- 每个技术点都要有具体的实现细节或数据支撑
- 段落间要有清晰的逻辑递进关系
- 引用要分布均匀，每段1-2个引用
- 避免重复表述，每句话都要有实质内容

【输出格式】：
<{header_tag}>{section_name}</{header_tag}>
<p>第一段内容（350-450字）...引用[{current_ref_counter + 1}]</p>
<p>第二段内容（350-450字）...引用[{current_ref_counter + 2}]</p>
...

在内容最后添加临时引用：
<div class="temp-references">
[{current_ref_counter + 1}] 作者1. 相关研究1[J]. 期刊名, 年份, 卷(期): 页码.
[{current_ref_counter + 2}] 作者2. 相关研究2[C]. 会议名, 年份: 页码.
...
</div>

请确保内容质量高、字数充足、逻辑清晰。"""

    return base_prompt


def enhance_content_for_word_count(content, section_name, target_words, actual_words):
    """增强内容以达到目标字数"""
    if actual_words >= target_words * 0.9:
        return content
    
    shortage = target_words - actual_words
    enhancement = f"""
<div class="content-enhancement" style="background: #f0f8ff; border-left: 4px solid #1976d2; padding: 15px; margin: 15px 0;">
<h4>📝 内容扩展建议</h4>
<p><strong>当前字数：</strong>{actual_words}字 / 目标字数：{target_words}字</p>
<p><strong>建议补充方向：</strong></p>
<ul>
<li><strong>技术深度：</strong>增加更多的技术实现细节、算法原理分析、性能优化策略</li>
<li><strong>案例分析：</strong>补充具体的应用案例、实验数据、对比分析结果</li>
<li><strong>理论支撑：</strong>扩展相关理论基础、学术背景、前沿发展趋势</li>
<li><strong>实践价值：</strong>详述实际应用价值、社会意义、经济效益分析</li>
</ul>
<p><strong>建议增加：</strong>{shortage}字以达到{target_words}字目标</p>
</div>"""
    
    return content + enhancement


def extract_reference_count_from_content(content):
    """从内容中提取引用数量"""
    import re
    references = re.findall(r'\[(\d+)\]', content)
    return len(set(references)) if references else 0


def generate_simple_section_content_with_citations(title, field, paper_type, section, abstract, keywords, requirements, section_name, citation_guide):
    """生成带文献引用的章节内容"""
    try:
        section_words = section['words']
        section_desc = section['description']
        level = section.get('level', 2)
        header_tag = f"h{level}"

        # 构建带引用的提示词
        if section_name == "摘要":
            prompt = f"""请为{paper_type}《{title}》生成专业的中文摘要。

研究领域：{field}
目标字数：{section_words}字
预设摘要：{abstract if abstract else '无'}
预设关键词：{keywords if keywords else '无'}

内容要求：
1. 研究背景和意义
2. 主要研究方法和技术路线
3. 核心成果和创新点
4. 结论和应用价值

写作要求：
- 使用连贯的段落形式
- 包含5-7个专业关键词
- 内容详实，字数充足
- 体现学术价值

请直接输出HTML格式内容：
<{header_tag}>摘要</{header_tag}>
<p>摘要内容...</p>
<p><strong>关键词：</strong>关键词1；关键词2；关键词3</p>"""

        elif section_name == "Abstract":
            prompt = f"""Please generate a professional English abstract for the {paper_type} titled "{title}".

Research Field: {field}
Target Words: {section_words} words
Chinese Abstract: {abstract[:200] if abstract else 'Not provided'}

Requirements:
1. Research background and significance
2. Main research methods and technical approaches
3. Key findings and innovations
4. Conclusions and practical applications

Writing Requirements:
- Use coherent paragraph format
- Include 5-7 professional keywords
- Comprehensive content with sufficient word count
- Demonstrate academic value

Please output HTML format directly:
<{header_tag}>Abstract</{header_tag}>
<p>Abstract content...</p>
<p><strong>Keywords:</strong> keyword1; keyword2; keyword3</p>"""

        else:
            # 其他章节，添加文献引用指导
            prompt = f"""请为{paper_type}《{title}》生成{section_name}章节。

研究领域：{field}
目标字数：{section_words}字
章节描述：{section_desc}
标题级别：{level}级标题

{citation_guide}

【严格格式要求】：
- 使用<{header_tag}>标签作为主标题
- 内容必须是连贯的段落形式，不能使用分点列表
- 在适当位置添加文献引用，使用上标格式：<sup>[序号]</sup>

【重要写作要求】：
- 使用自然流畅的段落形式，严格避免分点列表（1.、2.、3.或•、-等）
- 每个技术点都要融入到段落的自然表达中，避免生硬的要点罗列
- 必须包含4-6个自然段落，每段200-250字
- 段落之间要有逻辑递进关系，形成完整的学术论述
- 在理论阐述、技术方法介绍、研究现状分析等位置适当添加文献引用
- 使用"在...方面"、"通过深入分析"、"考虑到...需求"等自然过渡语

【引用示例】：
"相关研究表明<sup>[1]</sup>，该技术在实际应用中表现出良好的性能<sup>[2,3]</sup>。"
"根据最新的研究成果<sup>[4]</sup>，这种方法能够有效解决传统系统的局限性。"

请输出HTML格式内容：
<{header_tag}>{section_name}</{header_tag}>
<p>第一段详细内容（包含适当的文献引用）...</p>
<p>第二段详细内容（包含适当的文献引用）...</p>
<p>第三段详细内容（包含适当的文献引用）...</p>"""

        # 调用AI生成内容
        content = call_deepseek_api(prompt, min(section_words * 3, 6000))

        if content:
            cleaned_content = clean_ai_generated_content(content)
            return cleaned_content
        else:
            # 如果生成失败，回退到普通生成
            return generate_simple_section_content(
                title, field, paper_type, section, abstract, keywords, requirements, 1
            )

    except Exception as e:
        app.logger.error(f"生成带引用的章节内容失败: {e}")
        # 回退到普通生成
        return generate_simple_section_content(
            title, field, paper_type, section, abstract, keywords, requirements, 1
        )


# ==================== AI流程图生成功能 ====================

@app.route('/flowchart-generator')
def flowchart_generator_page():
    """AI流程图生成器页面"""
    return render_template('flowchart-generator.html')


@app.route('/api/get_flowchart_cost', methods=['GET'])
def api_get_flowchart_cost():
    """获取AI流程图生成的费用信息"""
    try:
        cost = float(user_manager.get_system_config('flowchart_generation_price', 1.0))

        if 'user_id' not in session:
            return jsonify({
                'is_logged_in': False,
                'cost': cost,
                'balance': 0,
                'sufficient': False
            })

        user_info = user_manager.get_user_info(session['user_id'])
        balance = float(user_info['balance']) if user_info else 0

        return jsonify({
            'is_logged_in': True,
            'cost': cost,
            'balance': balance,
            'sufficient': balance >= cost
        })
    except Exception as e:
        app.logger.error(f"获取流程图费用失败: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/flowchart/generate', methods=['POST'])
def generate_flowchart_api():
    """AI生成流程图API"""
    try:
        # 检查用户是否登录
        if 'user_id' not in session:
            return jsonify({
                'success': False,
                'message': '请先登录后再使用此功能',
                'need_login': True
            })

        # 获取流程图生成费用
        cost = float(user_manager.get_system_config('flowchart_generation_price', 1.0))

        # 获取用户信息并检查余额
        user_info = user_manager.get_user_info(session['user_id'])
        if not user_info:
            return jsonify({
                'success': False,
                'message': '用户信息获取失败，请重新登录',
                'need_login': True
            })

        if user_info['balance'] < cost:
            return jsonify({
                'success': False,
                'message': f'余额不足，流程图生成需要 {cost:.2f} 元，当前余额 {user_info["balance"]:.2f} 元',
                'need_recharge': True
            })

        data = request.get_json()
        description = data.get('description', '').strip()
        direction = data.get('direction', 'TD')

        if not description:
            return jsonify({'success': False, 'message': '请输入流程描述'})

        # 系统流程图提示词
        prompt = f"""你是一个专业的系统流程图设计专家。请将以下描述转换为紧凑的Mermaid流程图。

用户描述：
{description}

【严格要求 - 紧凑布局】：
1. 使用 flowchart {direction} 语法
2. 节点文字要简短（不超过8个汉字），例如："验证格式"而不是"系统验证输入格式是否正确"
3. 合并相似步骤，减少节点数量（整个流程图控制在8-12个节点）
4. 判断分支的两个结果尽量在同一层级并排显示
5. 节点类型：
   - 开始/结束：([开始]) 或 ([结束])
   - 处理步骤：[简短描述]
   - 判断条件：{{条件?}}
   - 输入/输出：[/数据/]

【紧凑布局技巧】：
- 判断节点的"是"和"否"分支指向的节点放在同一行
- 多个结束情况可以合并为一个结束节点
- 使用 & 符号让多个节点指向同一目标

【输出示例】（紧凑版）：
flowchart {direction}
    A([开始]) --> B[/输入信息/]
    B --> C{{格式有效?}}
    C -->|否| B
    C -->|是| D{{数据存在?}}
    D -->|是| E[提示重复]
    D -->|否| F[保存数据]
    E --> G([结束])
    F --> G

只输出Mermaid代码，不要任何解释。"""

        # 调用DeepSeek API
        mermaid_code = call_deepseek_api(prompt, max_tokens=1500)

        if not mermaid_code:
            return jsonify({'success': False, 'message': 'AI生成失败，请重试'})

        # 清理Mermaid代码
        mermaid_code = clean_mermaid_code(mermaid_code)

        # 生成成功后扣费
        consume_result = user_manager.consume_balance(
            session['user_id'],
            cost,
            'flowchart_generation',
            f'AI流程图生成 - {len(description)}字描述'
        )
        if not consume_result:
            app.logger.warning(f"流程图生成扣费失败: user_id={session['user_id']}, cost={cost}")

        # 获取更新后的余额
        updated_user_info = user_manager.get_user_info(session['user_id'])
        new_balance = float(updated_user_info['balance']) if updated_user_info else 0

        return jsonify({
            'success': True,
            'mermaid_code': mermaid_code,
            'message': '生成成功',
            'cost': cost,
            'new_balance': new_balance
        })

    except Exception as e:
        app.logger.error(f"流程图生成失败: {e}")
        return jsonify({'success': False, 'message': f'生成失败: {str(e)}'})


def clean_mermaid_code(code):
    """清理AI返回的Mermaid代码"""
    if not code:
        return ''

    # 移除markdown代码块标记
    code = code.strip()
    if code.startswith('```mermaid'):
        code = code[10:]
    elif code.startswith('```'):
        code = code[3:]
    if code.endswith('```'):
        code = code[:-3]

    # 移除多余的空行
    lines = [line for line in code.strip().split('\n') if line.strip()]
    code = '\n'.join(lines)

    return code.strip()


# ==================== 应用启动 ====================

if __name__ == '__main__':
    # 生产模式：禁用调试模式和自动重载，避免生成过程中断
    app.run(debug=False, host='localhost', port=5000, use_reloader=False)
