"""
Doc Generator Module - Generates database schema documentation
in various formats from parsed SQL metadata.
"""
from typing import Dict, Any, List, Tuple
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re


def _extract_type_and_length(data_type_full: str) -> Tuple[str, str]:
    """
    从完整的数据类型字符串中提取类型和长度
    例如: VARCHAR(50) -> ('VARCHAR', '50')
         INT -> ('INT', '-')
         DECIMAL(10,2) -> ('DECIMAL', '10,2')
    """
    if not data_type_full:
        return 'UNKNOWN', '-'

    # 使用正则表达式匹配类型和长度
    match = re.match(r'^(\w+)(?:\(([^)]+)\))?', data_type_full.upper())
    if match:
        data_type = match.group(1)
        length = match.group(2) if match.group(2) else '-'

        # 标准化数据类型名称
        if data_type in ['VARCHAR', 'CHAR']:
            data_type = 'VARCHAR'
        elif data_type in ['INT', 'INTEGER']:
            data_type = 'INT'
        elif data_type in ['BIGINT']:
            data_type = 'BIGINT'
        elif data_type in ['DECIMAL', 'NUMERIC']:
            data_type = 'DECIMAL'
        elif data_type in ['TEXT', 'LONGTEXT']:
            data_type = 'TEXT'
        elif data_type in ['DATE']:
            data_type = 'DATE'
        elif data_type in ['DATETIME', 'TIMESTAMP']:
            data_type = 'DATETIME'

        return data_type, length
    else:
        return data_type_full.upper(), '-'


def _set_triple_line_style(table):
    """
    应用标准三线表样式 - 只有三条线
    - 表格顶部粗线
    - 标题行下方细线  
    - 表格底部粗线
    - 无任何垂直线或其他横线
    """
    # 清除表格所有默认边框
    tbl = table._tbl
    
    # 移除表格样式，避免样式冲突
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl.insert(0, tbl_pr)
    
    # 清除所有现有边框
    old_borders = tbl_pr.find(qn('w:tblBorders'))
    if old_borders is not None:
        tbl_pr.remove(old_borders)
    
    # 设置表格边框为无
    tbl_borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tbl_borders.append(border)
    tbl_pr.append(tbl_borders)
    
    # 清除所有单元格边框
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            # 移除旧的边框
            old_borders = tc_pr.find(qn('w:tcBorders'))
            if old_borders is not None:
                tc_pr.remove(old_borders)
            
            # 设置无边框
            tc_borders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'nil')
                tc_borders.append(border)
            tc_pr.append(tc_borders)
    
    # 只添加三条线
    # 1. 顶线（第一行的顶部）- 粗线
    for cell in table.rows[0].cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = tc_pr.find(qn('w:tcBorders'))
        if tc_borders is None:
            tc_borders = OxmlElement('w:tcBorders')
            tc_pr.append(tc_borders)
        
        # 添加顶部边框
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), '12')  # 1.5pt 粗线
        top.set(qn('w:space'), '0')
        top.set(qn('w:color'), '000000')
        tc_borders.append(top)
    
    # 2. 栏目线（第一行的底部）- 细线
    for cell in table.rows[0].cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = tc_pr.find(qn('w:tcBorders'))
        if tc_borders is None:
            tc_borders = OxmlElement('w:tcBorders')
            tc_pr.append(tc_borders)
        
        # 添加底部边框
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')  # 0.75pt 细线
        bottom.set(qn('w:space'), '0')
        bottom.set(qn('w:color'), '000000')
        tc_borders.append(bottom)
    
    # 3. 底线（最后一行的底部）- 粗线
    for cell in table.rows[-1].cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = tc_pr.find(qn('w:tcBorders'))
        if tc_borders is None:
            tc_borders = OxmlElement('w:tcBorders')
            tc_pr.append(tc_borders)
        
        # 添加底部边框
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')  # 1.5pt 粗线
        bottom.set(qn('w:space'), '0')
        bottom.set(qn('w:color'), '000000')
        tc_borders.append(bottom)


def generate_html(tables_data: Dict[str, Any]) -> str:
    """
    生成带有三线表样式的HTML格式数据库结构文档
    """
    css = """
    <style>
        /* 完全隔离的预览容器样式 */
        .html-preview-container {
            /* 重置所有可能影响布局的属性 */
            all: initial;
            /* 重新设置需要的样式 */
            font-family: '宋体', 'SimSun', 'Microsoft YaHei', 'Arial', sans-serif;
            line-height: 1.6;
            color: #333;
            background: #f5f7fa;
            padding: 2rem;
            margin: 0;
            width: 100%;
            height: 100%;
            overflow: auto;
            box-sizing: border-box;
            display: block;
        }
        /* 只对预览容器内的元素应用box-sizing */
        .html-preview-container *,
        .html-preview-container *::before,
        .html-preview-container *::after {
            box-sizing: border-box;
        }
        .html-preview-container .document-container {
            max-width: 95%; /* 增加宽度利用率 */
            min-width: 1000px; /* 最小宽度确保表格不会太挤 */
            margin: 0 auto;
            background: white;
            padding: 2rem;
            border-radius: 8px;
            box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        }
        .html-preview-container .document-header {
            text-align: center;
            margin-bottom: 3rem;
            border-bottom: 3px solid #3498db;
            padding-bottom: 1.5rem;
        }
        .html-preview-container .document-title {
            font-size: 28pt;
            font-weight: bold;
            margin-bottom: 1rem;
            color: #2c3e50;
            text-shadow: 1px 1px 2px rgba(0,0,0,0.1);
        }
        .html-preview-container .document-info {
            font-size: 12pt;
            color: #7f8c8d;
            margin-bottom: 1rem;
            background: #ecf0f1;
            padding: 1rem;
            border-radius: 6px;
            display: inline-block;
        }
        .html-preview-container .table-section {
            margin-bottom: 3rem;
            page-break-inside: avoid;
            background: #fff;
            border-radius: 8px;
            padding: 1.5rem;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
            overflow-x: auto; /* 水平滚动支持 */
        }
        .html-preview-container .table-header {
            margin-bottom: 1rem;
        }
        .html-preview-container .table-title {
            font-size: 16pt;
            font-weight: bold;
            margin-bottom: 0.5rem;
            text-align: center;
            color: #2c3e50;
            border-bottom: 2px solid #3498db;
            padding-bottom: 0.5rem;
        }
        .html-preview-container .table-comment {
            font-size: 11pt;
            color: #666;
            font-style: italic;
            margin-bottom: 1rem;
            text-align: center;
        }
        /* 标准三线表样式 - 优化版 */
        .html-preview-container .three-line-table {
            width: 100%;
            border-collapse: collapse;
            font-size: 11pt;
            margin: 0 auto 1rem auto;
            border: none;
            background: #fff;
        }
        .html-preview-container .three-line-table thead tr {
            border-top: 2pt solid #000;
            border-bottom: 1pt solid #000;
            background: #f8f9fa;
        }
        .html-preview-container .three-line-table th {
            padding: 12px 10px;
            text-align: center;
            font-weight: bold;
            background: #f8f9fa;
            border: none;
            color: #333;
            font-size: 12pt;
            white-space: nowrap;
        }
        .html-preview-container .three-line-table td {
            padding: 10px 8px;
            border: none;
            vertical-align: middle;
            line-height: 1.5;
            color: #333;
            word-wrap: break-word;
        }
        .html-preview-container .three-line-table tbody tr:last-child {
            border-bottom: 2pt solid #000;
        }
        .html-preview-container .three-line-table tbody tr:hover {
            background-color: #f5f5f5;
        }
        /* 列宽设置 - 优化7列布局，增加空间 */
        .html-preview-container .three-line-table th:nth-child(1), .html-preview-container .three-line-table td:nth-child(1) {
            width: 16%;
            font-weight: 500;
            text-align: left;
            min-width: 120px;
        }  /* 字段名 */
        .html-preview-container .three-line-table th:nth-child(2), .html-preview-container .three-line-table td:nth-child(2) {
            width: 15%;
            text-align: left;
            font-family: 'Consolas', 'Courier New', monospace;
            min-width: 100px;
        }  /* 字段类型 */
        .html-preview-container .three-line-table th:nth-child(3), .html-preview-container .three-line-table td:nth-child(3) {
            width: 10%;
            text-align: center;
            min-width: 60px;
        }  /* 长度 */
        .html-preview-container .three-line-table th:nth-child(4), .html-preview-container .three-line-table td:nth-child(4) {
            width: 10%;
            text-align: center;
            min-width: 60px;
        }  /* 允许空 */
        .html-preview-container .three-line-table th:nth-child(5), .html-preview-container .three-line-table td:nth-child(5) {
            width: 8%;
            text-align: center;
            min-width: 50px;
        }  /* 主键 */
        .html-preview-container .three-line-table th:nth-child(6), .html-preview-container .three-line-table td:nth-child(6) {
            width: 15%;
            text-align: center;
            min-width: 100px;
        }  /* 默认值 */
        .html-preview-container .three-line-table th:nth-child(7), .html-preview-container .three-line-table td:nth-child(7) {
            width: 26%;
            text-align: left;
            padding-left: 12px;
            min-width: 150px;
        }  /* 说明 */

        .html-preview-container .field-name {
            font-weight: 500;
            color: #2c3e50;
        }
        .html-preview-container .field-type {
            font-family: 'Consolas', 'Courier New', monospace;
            color: #8e44ad;
            font-weight: 500;
        }
        .html-preview-container .pk-yes {
            color: #e74c3c;
            font-weight: bold;
        }
        .html-preview-container .nullable-no {
            color: #e67e22;
            font-weight: bold;
        }
        .html-preview-container .default-value {
            font-family: 'Consolas', 'Courier New', monospace;
            color: #27ae60;
            font-style: italic;
        }
        .html-preview-container .document-footer {
            text-align: center;
            margin-top: 3rem;
            padding-top: 1rem;
            border-top: 1px solid #ccc;
            font-size: 9pt;
            color: #666;
        }
        .html-preview-container .note {
            font-size: 9pt;
            color: #666;
            margin-top: 0.5rem;
            margin-left: 1cm;
        }
        /* 响应式设计 */
        @media (max-width: 1200px) {
            .html-preview-container .document-container {
                min-width: 800px;
                padding: 1rem;
            }
            .html-preview-container .three-line-table {
                font-size: 10pt;
            }
            .html-preview-container .three-line-table th, .html-preview-container .three-line-table td {
                padding: 8px 6px;
            }
        }

        @media (max-width: 800px) {
            .html-preview-container .document-container {
                min-width: 600px;
                padding: 0.5rem;
            }
            .html-preview-container .table-section {
                padding: 1rem;
            }
            .html-preview-container .three-line-table {
                font-size: 9pt;
            }
        }

        @media print {
            .html-preview-container { padding: 0; }
            .html-preview-container .table-section { page-break-inside: avoid; }
            .html-preview-container .document-container {
                max-width: 210mm;
                min-width: auto;
            }
        }
    </style>
    """
    
    # 统计信息
    total_tables = len(tables_data)
    total_columns = sum(len(data["columns"]) for data in tables_data.values())
    
    html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <title>数据库结构文档</title>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        {css}
    </head>
    <body>
        <div class="html-preview-container">
            <div class="document-container">
            <div class="document-header">
                <h1 class="document-title">数据库结构设计文档</h1>
                <p class="document-info">
                    生成时间：{__import__('datetime').datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}<br>
                    数据表数量：{total_tables} 个 · 总字段数：{total_columns} 个
                </p>
            </div>
            
            <div class="document-body">
    """
    
    # 生成每个表的文档
    for idx, (table_name, data) in enumerate(tables_data.items()):
        # 表标题格式：表序号：表说明(英文名)
        table_title = f"表{idx + 1}：{data.get('comment', table_name)}({table_name})"
        
        html += f"""
                <div class="table-section">
                    <h2 class="table-title">{table_title}</h2>
                    <table class="three-line-table">
                        <thead>
                            <tr>
                                <th>字段名</th>
                                <th>字段类型</th>
                                <th>长度</th>
                                <th>允许空</th>
                                <th>主键</th>
                                <th>默认值</th>
                                <th>说明</th>
                            </tr>
                        </thead>
                        <tbody>"""
        
        # 获取主键列表
        primary_keys = data.get('primary_keys', [])
        
        for col in data["columns"]:
            # 字段名
            field_name = col.get('name', '')

            # 字段类型和长度处理
            data_type_full = col.get('type', '')
            data_type, length = _extract_type_and_length(data_type_full)

            # 允许空
            nullable = col.get('nullable', True)
            nullable_text = '是' if nullable else '否'
            nullable_class = '' if nullable else ' class="nullable-no"'

            # 主键
            is_pk = col.get('pk') or col.get('name') in primary_keys
            pk_text = '是' if is_pk else '否'
            pk_class = ' class="pk-yes"' if is_pk else ''

            # 默认值处理
            default_val = col.get('default')
            if default_val is None:
                default_text = 'NULL' if nullable else '-'
            else:
                default_text = str(default_val)
            default_class = ' class="default-value"' if default_val is not None else ''

            # 说明
            comment = col.get('comment', '-') if col.get('comment') else '-'

            html += f"""
                            <tr>
                                <td><span class="field-name">{field_name}</span></td>
                                <td><span class="field-type">{data_type}</span></td>
                                <td>{length}</td>
                                <td{nullable_class}>{nullable_text}</td>
                                <td{pk_class}>{pk_text}</td>
                                <td{default_class}>{default_text}</td>
                                <td>{comment}</td>
                            </tr>"""
        
        html += """
                        </tbody>
                    </table>"""
        
        # 添加表注释（外键信息等）
        notes = []
        fk_info = {}
        for fk in data.get('foreign_keys', []):
            if fk['column'] not in fk_info:
                fk_info[fk['column']] = f"{fk['column']} → {fk['ref']['table']}.{fk['ref']['column']}"
        
        if fk_info:
            notes.append("外键关系：" + "；".join(fk_info.values()))
        
        if notes:
            html += f'<div class="note">注：{"<br>".join(notes)}</div>'
        
        html += """
                </div>"""
    
    html += f"""
            </div>
            
            <div class="document-footer">
                <p>———————— 文档结尾 ————————</p>
                <p>本文档由 ER图在线编辑器 自动生成</p>
            </div>
            </div>
        </div>
    </body>
    </html>
    """
    return html


def generate_docx(tables_data: Dict[str, Any], filename: str):
    """
    生成包含三线表格式的.docx文件，用于数据库结构文档
    """
    doc = Document()
    
    # 设置文档标题
    title = doc.add_heading('数据库结构设计文档', level=0)
    title.alignment = 1  # 居中
    
    # 添加文档信息
    info_para = doc.add_paragraph()
    info_para.alignment = 1  # 居中
    info_para.add_run(f'生成时间：{__import__("datetime").datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")}\n').font.size = Pt(10)
    info_para.add_run(f'数据表数量：{len(tables_data)} 个\n').font.size = Pt(10)
    info_para.add_run(f'总字段数：{sum(len(data["columns"]) for data in tables_data.values())} 个').font.size = Pt(10)
    
    doc.add_paragraph()  # 空行
    
    # 设置全局字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(10.5)
    
    # 为每个表生成文档
    for idx, (table_name, data) in enumerate(tables_data.items()):
        # 添加表标题 - 按照规范：表序号：表名(英文名)
        table_title = f"表{idx + 1}：{data.get('comment', table_name)}({table_name})"
        table_heading = doc.add_paragraph(table_title)
        table_heading.alignment = 1  # 居中
        table_heading.runs[0].font.bold = True
        table_heading.runs[0].font.size = Pt(12)  # 小四号字
        
        # 创建表格 - 使用完整的列设计
        headers = ["字段名", "字段类型", "长度", "允许空", "主键", "默认值", "说明"]
        tbl = doc.add_table(rows=1, cols=len(headers))
        # 不使用任何预定义样式
        tbl.style = None
        
        # 设置表头
        hdr_cells = tbl.rows[0].cells
        for i, header_text in enumerate(headers):
            hdr_cells[i].text = header_text
            # 设置表头样式
            for paragraph in hdr_cells[i].paragraphs:
                paragraph.alignment = 1  # 居中
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)  # 五号字
        
        # 获取主键列表
        primary_keys = data.get('primary_keys', [])
        
        # 添加数据行
        for col in data["columns"]:
            row_cells = tbl.add_row().cells

            # 字段名
            row_cells[0].text = col.get('name', '')

            # 字段类型和长度处理
            data_type_full = col.get('type', '')
            data_type, length = _extract_type_and_length(data_type_full)
            row_cells[1].text = data_type
            row_cells[2].text = length

            # 允许空
            nullable = col.get('nullable', True)
            row_cells[3].text = '是' if nullable else '否'

            # 主键
            is_pk = col.get('pk') or col.get('name') in primary_keys
            row_cells[4].text = '是' if is_pk else '否'

            # 默认值
            default_val = col.get('default')
            if default_val is None:
                row_cells[5].text = 'NULL' if nullable else '-'
            elif default_val == 'CURRENT_TIMESTAMP':
                row_cells[5].text = 'CURRENT_TIMESTAMP'
            else:
                row_cells[5].text = str(default_val)

            # 说明
            row_cells[6].text = col.get('comment', '-') if col.get('comment') else '-'
            
            # 设置单元格对齐和字体
            for i, cell in enumerate(row_cells):
                for paragraph in cell.paragraphs:
                    if i in [0, 6]:  # 字段名和说明左对齐
                        paragraph.alignment = 0
                    elif i in [2, 3, 4, 5]:  # 长度、允许空、主键、默认值居中
                        paragraph.alignment = 1
                    else:  # 字段类型左对齐
                        paragraph.alignment = 0
                    for run in paragraph.runs:
                        run.font.name = '宋体'
                        run.font.size = Pt(10)  # 五号字
        
        # 应用三线表样式
        _set_triple_line_style(tbl)
        
        # 设置表格宽度
        # 调整列宽以防止挤压 - 7列布局
        # 字段名、字段类型、长度、允许空、主键、默认值、说明
        widths = [Cm(3.0), Cm(2.5), Cm(1.5), Cm(1.5), Cm(1.5), Cm(2.5), Cm(5.5)]  # 总宽度约17.5cm
        for i, width in enumerate(widths):
            for row in tbl.rows:
                row.cells[i].width = width
        
        # 设置表格属性，避免自动调整
        tbl.autofit = False
        tbl.allow_autofit = False
        
        # 添加表注（如果有外键或特殊说明）
        notes = []
        # 收集外键信息
        fk_info = {}
        for fk in data.get('foreign_keys', []):
            if fk['column'] not in fk_info:
                fk_info[fk['column']] = f"{fk['column']} → {fk['ref']['table']}.{fk['ref']['column']}"
        
        if fk_info:
            notes.append("外键关系：" + "；".join(fk_info.values()))
        
        if notes:
            note_para = doc.add_paragraph("注：" + "\n    ".join(notes))
            note_para.paragraph_format.left_indent = Cm(0.5)
            note_para.runs[0].font.size = Pt(9)
            note_para.runs[0].font.italic = True
        
        doc.add_paragraph()  # 表格之间添加空行
        
        # 每3个表后添加分页符
        if (idx + 1) % 3 == 0 and idx + 1 < len(tables_data):
            doc.add_page_break()
    
    # 添加文档尾部信息
    doc.add_paragraph()  # 空行
    footer = doc.add_paragraph()
    footer.alignment = 1  # 居中
    footer.add_run('\n——————————— 文档结尾 ———————————\n').font.size = Pt(10)
    footer.add_run('本文档由 ER图在线编辑器 自动生成').font.size = Pt(9)
    
    doc.save(filename)